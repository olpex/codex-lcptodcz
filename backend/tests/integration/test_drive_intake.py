from datetime import datetime, timedelta, timezone
from io import BytesIO
from types import SimpleNamespace
from urllib.error import HTTPError

from docx import Document as DocxDocument
from openpyxl import Workbook

from app.core.crypto import cipher
from app.api.routes import jobs as jobs_route
from app.models import (
    Document,
    DocumentType,
    Group,
    GroupStatus,
    ImportJob,
    JobStatus,
    JournalMonitorEntry,
    JournalMonitorSection,
    JournalWorkloadEntry,
    Room,
    ScheduleSlot,
    Subject,
    Teacher,
    Trainee,
)
from app.services import drive_intake
from app.services.import_export import collect_teacher_workload_summary
from app.tasks import worker as worker_tasks
from app.tasks.worker import process_drive_intake_auto_task, process_import_job_task


def _contracts_workbook_bytes() -> bytes:
    stream = BytesIO()
    workbook = Workbook()
    workbook.active.title = "Архів"
    workbook.active.append(["Службова вкладка"])
    sheet = workbook.create_sheet("Додаток")
    sheet.append(["Група 184-25 Цифровий світ"])
    sheet.append([])
    sheet.append(
        [
            "№",
            "ПІБ безробітного",
            "Дата народження",
            "№ Договору",
            "Ідентифікаційний номер",
            "Адреса",
            "Телефон",
        ]
    )
    sheet.append(
        [
            1,
            "Петренко Іван Іванович",
            "01.02.1990",
            "184-25/001",
            "1234567890",
            "м. Львів, вул. Зелена 1",
            "+380501112233",
        ]
    )
    workbook.save(stream)
    stream.seek(0)
    return stream.read()


def _contracts_workbook_without_group_context_bytes() -> bytes:
    stream = BytesIO()
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Додаток"
    sheet.append(["Реєстр слухачів"])
    sheet.append([])
    sheet.append(["№", "ПІБ безробітного", "Дата народження", "№ Договору", "Телефон"])
    sheet.append([1, "Кравченко Олена Іванівна", "05.03.1991", "80-26/001", "+380501112233"])
    workbook.save(stream)
    stream.seek(0)
    return stream.read()


def _schedule_docx_bytes(group_code: str = "46-26") -> bytes:
    stream = BytesIO()
    document = DocxDocument()
    document.add_paragraph("1 пара - 9.30 - 11.05")
    document.add_paragraph("за напрямом")
    document.add_paragraph("Цифрові навички")
    document.add_paragraph(f"Група № {group_code}")
    document.add_paragraph("з 12 травня 2026 року до 12 травня 2026 року")
    table = document.add_table(rows=3, cols=6)
    for idx, value in enumerate(
        [
            "№п/п",
            "Назва предмета",
            "К-сть год.",
            "12.05",
            "Прізвище, ім'я, по-батькові викладача",
            "Примітка",
        ]
    ):
        table.cell(0, idx).text = value
    for idx, value in enumerate(["1", "Цифрова грамотність", "2", "1п/2год", "Коваль Олена Петрівна", ""]):
        table.cell(1, idx).text = value
    for idx, value in enumerate(["", "Загальний обсяг навчального часу:", "2", "", "", ""]):
        table.cell(2, idx).text = value
    document.save(stream)
    stream.seek(0)
    return stream.read()


def _run_import_job(job_id: int) -> dict:
    return process_import_job_task.run(job_id)


def test_drive_intake_processes_one_contract_file_and_updates_existing_trainee(db_session):
    db_session.add(
        Trainee(
            branch_id="main",
            first_name="Іван Іванович",
            last_name="Петренко",
            status="active",
            group_code="184-25",
        )
    )
    db_session.commit()

    files = [
        {
            "id": "contracts-184-25",
            "name": "184-25 Договори Цифровий світ.xlsx",
            "mimeType": drive_intake.GOOGLE_DRIVE_XLSX_MIME,
            "modifiedTime": "2026-05-12T07:00:00Z",
            "webViewLink": "https://drive.google.com/file/d/contracts-184-25/view",
        },
        {
            "id": "schedule-46-26",
            "name": "46-26 Розклад.docx",
            "mimeType": drive_intake.GOOGLE_DRIVE_DOCX_MIME,
            "modifiedTime": "2026-05-12T07:01:00Z",
            "webViewLink": "https://drive.google.com/file/d/schedule-46-26/view",
        },
    ]

    result = drive_intake.process_next_drive_intake_file(
        db_session,
        folder_url="https://drive.google.com/drive/folders/intake-folder",
        branch_id="main",
        file_lister=lambda folder_id, service_account_json=None: files,
        downloader=lambda file_id, mime_type=None, service_account_json=None: _contracts_workbook_bytes(),
        import_job_runner=_run_import_job,
    )

    assert result["processed"] == 1
    assert result["skipped_already_processed"] == 0
    assert result["job_id"] is not None
    db_session.expire_all()

    assert db_session.query(ImportJob).count() == 1
    trainee = db_session.query(Trainee).filter(Trainee.last_name == "Петренко").one()
    assert trainee.contract_number == "184-25/001"
    assert cipher.decrypt(trainee.tax_id_encrypted) == "1234567890"
    assert cipher.decrypt(trainee.phone_encrypted) == "+380501112233"
    assert db_session.query(ScheduleSlot).count() == 0


def test_drive_intake_uses_filename_group_code_when_contract_file_has_no_group_context(db_session):
    result = drive_intake.process_next_drive_intake_file(
        db_session,
        folder_url="https://drive.google.com/drive/folders/intake-folder",
        branch_id="main",
        file_lister=lambda folder_id, service_account_json=None: [
            {
                "id": "contracts-80-26",
                "name": "Договори 80-26 Організація трудових відносин.xls",
                "mimeType": drive_intake.GOOGLE_DRIVE_XLS_MIME,
                "modifiedTime": "2026-05-12T07:00:00Z",
                "webViewLink": "https://drive.google.com/file/d/contracts-80-26/view",
            }
        ],
        downloader=lambda file_id, mime_type=None, service_account_json=None: _contracts_workbook_without_group_context_bytes(),
        import_job_runner=_run_import_job,
    )

    assert result["processed"] == 1
    db_session.expire_all()

    trainee = db_session.query(Trainee).filter(Trainee.last_name == "Кравченко").one()
    assert trainee.group_code == "80-26"

    group = db_session.query(Group).filter(Group.code == "80-26").one()
    assert group.name == "Група 80-26"


def test_drive_intake_schedule_creates_calendar_without_duplicating_journal_workload(db_session):
    group = Group(branch_id="main", code="46-26", name="Група 46-26", status=GroupStatus.ACTIVE)
    teacher = Teacher(branch_id="main", first_name="Олена Петрівна", last_name="Коваль", is_active=True)
    subject = Subject(branch_id="main", name="Цифрова грамотність", hours_total=2)
    room = Room(branch_id="main", name="Журнал 46-26", capacity=20)
    section = JournalMonitorSection(
        branch_id="main",
        name="Журнали 2026",
        folder_url="https://drive.google.com/drive/folders/journals",
        folder_id="journals",
    )
    db_session.add_all([group, teacher, subject, room, section])
    db_session.flush()
    entry = JournalMonitorEntry(
        section_id=section.id,
        branch_id="main",
        drive_file_id="journal-46-26",
        journal_name="46-26 Журнал",
        group_code="46-26",
        workload_status="processed",
        workload_year=2026,
        workload_hours=2,
    )
    db_session.add(entry)
    db_session.flush()
    db_session.add(
        JournalWorkloadEntry(
            journal_monitor_entry_id=entry.id,
            branch_id="main",
            teacher_id=teacher.id,
            subject_name="Цифрова грамотність",
            hours=2,
        )
    )
    db_session.commit()

    result = drive_intake.process_next_drive_intake_file(
        db_session,
        folder_url="https://drive.google.com/drive/folders/intake-folder",
        branch_id="main",
        file_lister=lambda folder_id, service_account_json=None: [
            {
                "id": "schedule-46-26",
                "name": "46-26 Розклад.docx",
                "mimeType": drive_intake.GOOGLE_DRIVE_DOCX_MIME,
                "modifiedTime": "2026-05-12T07:00:00Z",
                "webViewLink": "https://drive.google.com/file/d/schedule-46-26/view",
            }
        ],
        downloader=lambda file_id, mime_type=None, service_account_json=None: _schedule_docx_bytes(),
        import_job_runner=_run_import_job,
    )

    assert result["processed"] == 1
    db_session.expire_all()

    assert db_session.query(ScheduleSlot).count() == 1
    assert db_session.query(JournalWorkloadEntry).count() == 1
    workload_rows = collect_teacher_workload_summary(db_session, "main")
    assert workload_rows[0]["total_hours"] == 2
    assert workload_rows[0]["groups"][0]["group_code"] == "46-26"
    assert workload_rows[0]["groups"][0]["hours"] == 2.0


def test_drive_intake_import_invalidates_cross_section_caches(db_session, monkeypatch):
    invalidated: list[tuple[str, str]] = []
    monkeypatch.setattr(
        worker_tasks,
        "invalidate_schedule_list_cache",
        lambda branch_id: invalidated.append(("schedule", branch_id)),
        raising=False,
    )
    monkeypatch.setattr(
        worker_tasks,
        "invalidate_group_list_cache",
        lambda branch_id: invalidated.append(("groups", branch_id)),
        raising=False,
    )
    monkeypatch.setattr(
        worker_tasks,
        "invalidate_attention_cache",
        lambda branch_id: invalidated.append(("dashboard", branch_id)),
        raising=False,
    )

    result = drive_intake.process_next_drive_intake_file(
        db_session,
        folder_url="https://drive.google.com/drive/folders/intake-folder",
        branch_id="main",
        file_lister=lambda folder_id, service_account_json=None: [
            {
                "id": "schedule-46-26",
                "name": "46-26 Розклад.docx",
                "mimeType": drive_intake.GOOGLE_DRIVE_DOCX_MIME,
                "modifiedTime": "2026-05-12T07:00:00Z",
                "webViewLink": "https://drive.google.com/file/d/schedule-46-26/view",
            }
        ],
        downloader=lambda file_id, mime_type=None, service_account_json=None: _schedule_docx_bytes(),
        import_job_runner=_run_import_job,
    )

    assert result["processed"] == 1
    assert result["status"] == JobStatus.SUCCEEDED.value
    assert invalidated == [("schedule", "main"), ("groups", "main"), ("dashboard", "main")]


def test_drive_intake_skips_files_that_were_already_processed(db_session):
    file_payload = {
        "id": "contracts-184-25",
        "name": "184-25 Договори.xlsx",
        "mimeType": drive_intake.GOOGLE_DRIVE_XLSX_MIME,
        "modifiedTime": "2026-05-12T07:00:00Z",
    }

    first = drive_intake.process_next_drive_intake_file(
        db_session,
        folder_url="https://drive.google.com/drive/folders/intake-folder",
        branch_id="main",
        file_lister=lambda folder_id, service_account_json=None: [file_payload],
        downloader=lambda file_id, mime_type=None, service_account_json=None: _contracts_workbook_bytes(),
        import_job_runner=_run_import_job,
        processed_file_marker=None,
    )
    second = drive_intake.process_next_drive_intake_file(
        db_session,
        folder_url="https://drive.google.com/drive/folders/intake-folder",
        branch_id="main",
        file_lister=lambda folder_id, service_account_json=None: [file_payload],
        downloader=lambda file_id, mime_type=None, service_account_json=None: _contracts_workbook_bytes(),
        import_job_runner=_run_import_job,
        processed_file_marker=None,
    )

    assert first["processed"] == 1
    assert second == {"processed": 0, "skipped_already_processed": 1, "skipped_unsupported": 0}
    assert db_session.query(ImportJob).count() == 1


def test_drive_intake_skips_files_with_processed_marker(db_session):
    downloader_calls: list[str] = []

    result = drive_intake.process_next_drive_intake_file(
        db_session,
        folder_url="https://drive.google.com/drive/folders/intake-folder",
        branch_id="main",
        file_lister=lambda folder_id, service_account_json=None: [
            {
                "id": "contracts-processed",
                "name": "184-25 Contracts [processed].xlsx",
                "mimeType": drive_intake.GOOGLE_DRIVE_XLSX_MIME,
                "modifiedTime": "2026-05-12T07:00:00Z",
            }
        ],
        downloader=lambda file_id, mime_type=None, service_account_json=None: downloader_calls.append(file_id) or _contracts_workbook_bytes(),
        import_job_runner=_run_import_job,
    )

    assert result == {
        "processed": 0,
        "skipped_already_processed": 0,
        "skipped_unsupported": 0,
        "skipped_marked_processed": 1,
    }
    assert downloader_calls == []
    assert db_session.query(ImportJob).count() == 0


def test_drive_intake_marks_file_after_successful_import(db_session):
    marker_calls: list[tuple[str, str, str | None]] = []

    result = drive_intake.process_next_drive_intake_file(
        db_session,
        folder_url="https://drive.google.com/drive/folders/intake-folder",
        branch_id="main",
        service_account_json="service-account-json",
        file_lister=lambda folder_id, service_account_json=None: [
            {
                "id": "contracts-184-25",
                "name": "184-25 Contracts.xlsx",
                "mimeType": drive_intake.GOOGLE_DRIVE_XLSX_MIME,
                "modifiedTime": "2026-05-12T07:00:00Z",
            }
        ],
        downloader=lambda file_id, mime_type=None, service_account_json=None: _contracts_workbook_bytes(),
        import_job_runner=_run_import_job,
        processed_file_marker=lambda file_id, next_name, service_account_json=None: marker_calls.append(
            (file_id, next_name, service_account_json)
        ),
    )

    assert result["processed"] == 1
    assert result["status"] == JobStatus.SUCCEEDED.value
    assert result["marked_processed"] is True
    assert marker_calls == [("contracts-184-25", "184-25 Contracts [processed].xlsx", "service-account-json")]


def test_drive_intake_does_not_mark_file_after_failed_import(db_session):
    marker_calls: list[tuple[str, str]] = []

    def failing_runner(job_id: int) -> dict:
        job = db_session.get(ImportJob, job_id)
        job.status = JobStatus.FAILED
        db_session.add(job)
        db_session.commit()
        return {"status": "failed"}

    result = drive_intake.process_next_drive_intake_file(
        db_session,
        folder_url="https://drive.google.com/drive/folders/intake-folder",
        branch_id="main",
        file_lister=lambda folder_id, service_account_json=None: [
            {
                "id": "bad-contracts",
                "name": "Bad Contracts.xlsx",
                "mimeType": drive_intake.GOOGLE_DRIVE_XLSX_MIME,
                "modifiedTime": "2026-05-12T07:00:00Z",
            }
        ],
        downloader=lambda file_id, mime_type=None, service_account_json=None: _contracts_workbook_bytes(),
        import_job_runner=failing_runner,
        processed_file_marker=lambda file_id, next_name, service_account_json=None: marker_calls.append((file_id, next_name)),
    )

    assert result["processed"] == 1
    assert result["status"] == JobStatus.FAILED.value
    assert result.get("marked_processed") is not True
    assert marker_calls == []


def test_drive_intake_marker_explains_google_drive_permission_denial(monkeypatch):
    class FakeErrorResponse:
        def read(self):
            return b'{"error":{"message":"The user does not have sufficient permissions"}}'

        def close(self):
            return None

    def fake_drive_request_url(url: str, service_account_json: str | None = None):
        return drive_intake.Request(url, headers={"Authorization": "Bearer service-token"})

    def fake_urlopen(request, timeout):
        raise HTTPError(request.full_url, 403, "Forbidden", {}, FakeErrorResponse())

    monkeypatch.setattr(drive_intake, "_drive_request_url", fake_drive_request_url)
    monkeypatch.setattr(drive_intake, "urlopen", fake_urlopen)

    try:
        drive_intake.mark_drive_intake_file_processed("file-1", "File [processed].docx", "service-account-json")
    except RuntimeError as exc:
        assert "Editor" in str(exc)
        assert "перейменувати" in str(exc)
    else:
        raise AssertionError("Expected permission error")


def test_drive_intake_listing_uses_short_lived_cache(monkeypatch):
    cache: dict[str, list[dict]] = {}
    cache_sets: list[tuple[str, int]] = []
    calls = {"count": 0}

    monkeypatch.setattr(drive_intake, "cache_get_json", lambda key: cache.get(key))

    def fake_cache_set(key: str, payload: list[dict], ttl_seconds: int) -> None:
        cache[key] = payload
        cache_sets.append((key, ttl_seconds))

    monkeypatch.setattr(drive_intake, "cache_set_json", fake_cache_set)
    monkeypatch.setattr(drive_intake, "_drive_request_url", lambda url, service_account_json=None: url)

    class FakeResponse:
        def __enter__(self):
            return self

        def __exit__(self, *_args):
            return False

        def read(self):
            return b'{"files":[{"id":"file-1","name":"Contracts.xlsx","mimeType":"application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"}]}'

    def fake_urlopen(request, timeout):
        calls["count"] += 1
        return FakeResponse()

    monkeypatch.setattr(drive_intake, "urlopen", fake_urlopen)

    first = drive_intake.list_drive_intake_files("root-folder")
    second = drive_intake.list_drive_intake_files("root-folder")

    assert first == second
    assert calls["count"] == 1
    assert cache_sets[0][1] == 45


def test_drive_intake_reports_marking_error_for_existing_successful_unmarked_file(db_session, tmp_path):
    schedule_path = tmp_path / "successful-schedule.docx"
    schedule_path.write_bytes(_schedule_docx_bytes())
    document = Document(
        branch_id="main",
        file_name="46-26 Schedule.docx",
        file_path=str(schedule_path),
        file_type=DocumentType.DOCX,
        source="drive_intake",
        mime_type=drive_intake.GOOGLE_DRIVE_DOCX_MIME,
        hash_sha256="already-succeeded",
    )
    db_session.add(document)
    db_session.flush()
    group = Group(branch_id="main", code="46-26", name="Група 46-26", status=GroupStatus.ACTIVE)
    teacher = Teacher(branch_id="main", first_name="Олена Петрівна", last_name="Коваль", is_active=True)
    subject = Subject(branch_id="main", name="Цифрова грамотність", hours_total=2)
    room = Room(branch_id="main", name="Імпорт: 46-26", capacity=20)
    db_session.add_all([group, teacher, subject, room])
    db_session.flush()
    db_session.add(
        ScheduleSlot(
            group_id=group.id,
            teacher_id=teacher.id,
            subject_id=subject.id,
            room_id=room.id,
            starts_at=datetime(2026, 5, 12, 9, 30, tzinfo=timezone.utc),
            ends_at=datetime(2026, 5, 12, 11, 5, tzinfo=timezone.utc),
            pair_number=1,
            academic_hours=2,
        )
    )
    db_session.add(
        ImportJob(
            branch_id="main",
            idempotency_key=drive_intake._idempotency_key(
                "main",
                "schedule-46-26",
                "2026-05-12T07:00:00Z",
                "46-26 Schedule.docx",
            ),
            document_id=document.id,
            status=JobStatus.SUCCEEDED,
            message="already imported",
            result_payload={"source": "drive_intake"},
        )
    )
    db_session.commit()

    result = drive_intake.process_next_drive_intake_file(
        db_session,
        folder_url="https://drive.google.com/drive/folders/intake-folder",
        branch_id="main",
        file_lister=lambda folder_id, service_account_json=None: [
            {
                "id": "schedule-46-26",
                "name": "46-26 Schedule.docx",
                "mimeType": drive_intake.GOOGLE_DRIVE_DOCX_MIME,
                "modifiedTime": "2026-05-12T07:00:00Z",
            }
        ],
        downloader=lambda file_id, mime_type=None, service_account_json=None: b"",
        import_job_runner=_run_import_job,
        processed_file_marker=lambda file_id, next_name, service_account_json=None: (_ for _ in ()).throw(
            RuntimeError("Google Drive denied rename")
        ),
    )

    assert result["processed"] == 0
    assert result["skipped_already_processed"] == 1
    assert result["processed_drive_file_name"] == "46-26 Schedule [processed].docx"
    assert result["marking_error"] == "Google Drive denied rename"
    assert db_session.query(ImportJob).count() == 1


def test_drive_intake_reprocesses_when_processed_marker_is_removed(db_session):
    processed_name = "184-25 Contracts [processed].xlsx"
    db_session.add(
        Document(
            branch_id="main",
            file_name=processed_name,
            file_path="already-imported.xlsx",
            file_type=DocumentType.XLSX,
            source="drive_intake",
            mime_type=drive_intake.GOOGLE_DRIVE_XLSX_MIME,
            hash_sha256="already-imported",
        )
    )
    db_session.flush()
    db_session.add(
        ImportJob(
            branch_id="main",
            idempotency_key=drive_intake._idempotency_key(
                "main",
                "contracts-184-25",
                "2026-05-12T07:00:00Z",
                processed_name,
            ),
            document_id=db_session.query(Document).filter(Document.file_name == processed_name).one().id,
            status=JobStatus.SUCCEEDED,
            message="already imported",
            result_payload={"source": "drive_intake"},
        )
    )
    db_session.commit()

    result = drive_intake.process_next_drive_intake_file(
        db_session,
        folder_url="https://drive.google.com/drive/folders/intake-folder",
        branch_id="main",
        file_lister=lambda folder_id, service_account_json=None: [
            {
                "id": "contracts-184-25",
                "name": "184-25 Contracts.xlsx",
                "mimeType": drive_intake.GOOGLE_DRIVE_XLSX_MIME,
                "modifiedTime": "2026-05-12T07:00:00Z",
            }
        ],
        downloader=lambda file_id, mime_type=None, service_account_json=None: _contracts_workbook_bytes(),
        import_job_runner=_run_import_job,
        processed_file_marker=lambda file_id, next_name, service_account_json=None: None,
    )

    assert result["processed"] == 1
    assert result["status"] == JobStatus.SUCCEEDED.value
    assert db_session.query(ImportJob).count() == 2


def test_drive_intake_reprocesses_unmarked_file_with_legacy_success_job(db_session):
    original_name = "46-26 Schedule.docx"
    db_session.add(
        Document(
            branch_id="main",
            file_name=original_name,
            file_path="legacy-schedule.docx",
            file_type=DocumentType.DOCX,
            source="drive_intake",
            mime_type=drive_intake.GOOGLE_DRIVE_DOCX_MIME,
            hash_sha256="legacy-success",
        )
    )
    db_session.flush()
    db_session.add(
        ImportJob(
            branch_id="main",
            idempotency_key=drive_intake._idempotency_key("main", "schedule-46-26", "2026-05-12T07:00:00Z"),
            document_id=db_session.query(Document).filter(Document.hash_sha256 == "legacy-success").one().id,
            status=JobStatus.SUCCEEDED,
            message="imported before processed markers existed",
            result_payload={"source": "drive_intake"},
        )
    )
    db_session.commit()

    marker_calls: list[tuple[str, str]] = []

    result = drive_intake.process_next_drive_intake_file(
        db_session,
        folder_url="https://drive.google.com/drive/folders/intake-folder",
        branch_id="main",
        file_lister=lambda folder_id, service_account_json=None: [
            {
                "id": "schedule-46-26",
                "name": original_name,
                "mimeType": drive_intake.GOOGLE_DRIVE_DOCX_MIME,
                "modifiedTime": "2026-05-12T07:00:00Z",
                "webViewLink": "https://drive.google.com/file/d/schedule-46-26/view",
            }
        ],
        downloader=lambda file_id, mime_type=None, service_account_json=None: _schedule_docx_bytes(),
        import_job_runner=_run_import_job,
        processed_file_marker=lambda file_id, next_name, service_account_json=None: marker_calls.append((file_id, next_name)),
    )

    assert result["processed"] == 1
    assert result["status"] == JobStatus.SUCCEEDED.value
    assert result["marked_processed"] is True
    assert result["reprocessed_legacy_success_job"] is True
    assert marker_calls == [("schedule-46-26", "46-26 Schedule [processed].docx")]
    assert db_session.query(ImportJob).count() == 2
    assert db_session.query(ScheduleSlot).count() == 1


def test_drive_intake_resyncs_schedule_when_processed_drive_file_remains_after_slots_deleted(db_session):
    marker_calls: list[tuple[str, str]] = []
    file_payload = {
        "id": "schedule-46-26",
        "name": "46-26 Schedule.docx",
        "mimeType": drive_intake.GOOGLE_DRIVE_DOCX_MIME,
        "modifiedTime": "2026-05-12T07:00:00Z",
        "webViewLink": "https://drive.google.com/file/d/schedule-46-26/view",
    }

    first = drive_intake.process_next_drive_intake_file(
        db_session,
        folder_url="https://drive.google.com/drive/folders/intake-folder",
        branch_id="main",
        file_lister=lambda folder_id, service_account_json=None: [file_payload],
        downloader=lambda file_id, mime_type=None, service_account_json=None: _schedule_docx_bytes(),
        import_job_runner=_run_import_job,
        processed_file_marker=lambda file_id, next_name, service_account_json=None: marker_calls.append((file_id, next_name)),
    )
    assert first["processed"] == 1
    assert first["marked_processed"] is True
    assert db_session.query(ScheduleSlot).count() == 1

    db_session.query(ScheduleSlot).delete(synchronize_session=False)
    db_session.commit()

    resync = drive_intake.process_next_drive_intake_file(
        db_session,
        folder_url="https://drive.google.com/drive/folders/intake-folder",
        branch_id="main",
        file_lister=lambda folder_id, service_account_json=None: [
            {**file_payload, "name": "46-26 Schedule [processed].docx"}
        ],
        downloader=lambda file_id, mime_type=None, service_account_json=None: b"",
        import_job_runner=_run_import_job,
        processed_file_marker=lambda file_id, next_name, service_account_json=None: marker_calls.append((file_id, next_name)),
    )

    assert resync["processed"] == 1
    assert resync["resynced_schedule"] is True
    assert resync["job_id"] == first["job_id"]
    assert db_session.query(ScheduleSlot).count() == 1
    assert marker_calls == [("schedule-46-26", "46-26 Schedule [processed].docx")]


def test_drive_intake_redownloads_processed_schedule_when_local_copy_is_missing(db_session, tmp_path):
    missing_path = tmp_path / "missing-schedule.docx"
    document = Document(
        branch_id="main",
        file_name="46-26 Schedule [processed].docx",
        file_path=str(missing_path),
        file_type=DocumentType.DOCX,
        source="drive_intake",
        mime_type=drive_intake.GOOGLE_DRIVE_DOCX_MIME,
        hash_sha256="missing-local-copy",
    )
    db_session.add(document)
    db_session.flush()
    job = ImportJob(
        branch_id="main",
        idempotency_key=drive_intake._idempotency_key(
            "main",
            "schedule-46-26",
            "2026-05-12T07:00:00Z",
            "46-26 Schedule [processed].docx",
        ),
        document_id=document.id,
        status=JobStatus.SUCCEEDED,
        message="previously imported",
        result_payload={
            "source": "drive_intake",
            "drive_file_id": "schedule-46-26",
            "drive_file_name": "46-26 Schedule [processed].docx",
            "drive_modified_time": "2026-05-12T07:00:00Z",
            "import_result": {"created_slots": 1, "group_code": "46-26"},
        },
    )
    db_session.add(job)
    db_session.commit()

    downloader_calls: list[str] = []

    result = drive_intake.process_next_drive_intake_file(
        db_session,
        folder_url="https://drive.google.com/drive/folders/intake-folder",
        branch_id="main",
        file_lister=lambda folder_id, service_account_json=None: [
            {
                "id": "schedule-46-26",
                "name": "46-26 Schedule [processed].docx",
                "mimeType": drive_intake.GOOGLE_DRIVE_DOCX_MIME,
                "modifiedTime": "2026-05-12T07:00:00Z",
                "webViewLink": "https://drive.google.com/file/d/schedule-46-26/view",
            }
        ],
        downloader=lambda file_id, mime_type=None, service_account_json=None: downloader_calls.append(file_id)
        or _schedule_docx_bytes(),
        import_job_runner=_run_import_job,
        processed_file_marker=lambda file_id, next_name, service_account_json=None: None,
    )

    assert result["processed"] == 1
    assert result["resynced_schedule"] is True
    assert result["job_id"] == job.id
    assert downloader_calls == ["schedule-46-26"]
    db_session.expire_all()
    assert db_session.get(ImportJob, job.id).status == JobStatus.SUCCEEDED
    assert db_session.query(ScheduleSlot).count() == 1


def test_import_worker_redownloads_missing_drive_intake_document(db_session, tmp_path, monkeypatch):
    missing_path = tmp_path / "missing-worker-schedule.docx"
    document = Document(
        branch_id="main",
        file_name="46-26 Розклад.docx",
        file_path=str(missing_path),
        file_type=DocumentType.DOCX,
        source="drive_intake",
        mime_type=drive_intake.GOOGLE_DRIVE_DOCX_MIME,
        hash_sha256="missing-worker-copy",
    )
    db_session.add(document)
    db_session.flush()
    job = ImportJob(
        branch_id="main",
        idempotency_key="main:drive-worker-missing-doc",
        document_id=document.id,
        status=JobStatus.QUEUED,
        message="queued",
        result_payload={
            "source": "drive_intake",
            "channel": "google_drive_folder",
            "drive_file_id": "schedule-46-26",
            "drive_file_name": "46-26 Розклад.docx",
            "drive_modified_time": "2026-05-12T07:00:00Z",
            "import_mode": "overwrite",
        },
    )
    db_session.add(job)
    db_session.commit()

    download_calls: list[tuple[str, str | None, str | None]] = []

    monkeypatch.setattr(
        "app.tasks.worker.resolve_drive_intake_service_account_json",
        lambda db, branch_id: "section-service-account-json",
    )
    monkeypatch.setattr(
        "app.tasks.worker.download_drive_file_bytes",
        lambda file_id, mime_type=None, service_account_json=None: download_calls.append(
            (file_id, mime_type, service_account_json)
        )
        or _schedule_docx_bytes(),
        raising=False,
    )

    result = process_import_job_task.run(job.id)

    assert result["status"] == "ok"
    assert download_calls == [("schedule-46-26", drive_intake.GOOGLE_DRIVE_DOCX_MIME, "section-service-account-json")]
    db_session.expire_all()
    refreshed_document = db_session.get(Document, document.id)
    assert refreshed_document is not None
    assert refreshed_document.file_path != str(missing_path)
    assert db_session.get(ImportJob, job.id).status == JobStatus.SUCCEEDED
    assert db_session.query(ScheduleSlot).count() == 1


def test_manual_reprocess_redownloads_missing_drive_intake_document(
    client,
    auth_headers,
    db_session,
    tmp_path,
    monkeypatch,
):
    missing_path = tmp_path / "missing-reprocess-schedule.docx"
    document = Document(
        branch_id="main",
        file_name="46-26 Розклад.docx",
        file_path=str(missing_path),
        file_type=DocumentType.DOCX,
        source="drive_intake",
        mime_type=drive_intake.GOOGLE_DRIVE_DOCX_MIME,
        hash_sha256="missing-reprocess-copy",
    )
    db_session.add(document)
    db_session.flush()
    source_job = ImportJob(
        branch_id="main",
        idempotency_key="main:drive-manual-reprocess-source",
        document_id=document.id,
        status=JobStatus.FAILED,
        message="Package not found at old path",
        result_payload={
            "source": "drive_intake",
            "channel": "google_drive_folder",
            "drive_file_id": "schedule-46-26",
            "drive_file_name": "46-26 Розклад.docx",
            "drive_modified_time": "2026-05-12T07:00:00Z",
            "drive_url": "https://drive.google.com/file/d/schedule-46-26/view",
            "group_code_hint": "46-26",
            "import_mode": "overwrite",
        },
    )
    db_session.add(source_job)
    db_session.commit()

    download_calls: list[str] = []

    monkeypatch.setattr(jobs_route.process_import_job_task, "delay", lambda job_id: (_ for _ in ()).throw(RuntimeError("queue off")))
    monkeypatch.setattr(
        "app.tasks.worker.resolve_drive_intake_service_account_json",
        lambda db, branch_id: "section-service-account-json",
    )
    monkeypatch.setattr(
        "app.tasks.worker.download_drive_file_bytes",
        lambda file_id, mime_type=None, service_account_json=None: download_calls.append(file_id)
        or _schedule_docx_bytes(),
        raising=False,
    )

    response = client.post(f"/api/v1/jobs/{source_job.id}/reprocess-import", headers=auth_headers)

    assert response.status_code == 200
    payload = response.json()
    assert payload["job"]["status"] == JobStatus.SUCCEEDED.value
    assert download_calls == ["schedule-46-26"]
    db_session.expire_all()
    reprocess_job = db_session.get(ImportJob, payload["job"]["id"])
    assert reprocess_job is not None
    assert reprocess_job.result_payload["drive_file_id"] == "schedule-46-26"
    assert reprocess_job.result_payload["source"] == "manual_reprocess"
    assert db_session.query(ScheduleSlot).count() == 1


def test_drive_intake_resyncs_unmarked_existing_schedule_before_marking_processed(db_session):
    file_payload = {
        "id": "schedule-46-26",
        "name": "46-26 Schedule.docx",
        "mimeType": drive_intake.GOOGLE_DRIVE_DOCX_MIME,
        "modifiedTime": "2026-05-12T07:00:00Z",
        "webViewLink": "https://drive.google.com/file/d/schedule-46-26/view",
    }

    first = drive_intake.process_next_drive_intake_file(
        db_session,
        folder_url="https://drive.google.com/drive/folders/intake-folder",
        branch_id="main",
        file_lister=lambda folder_id, service_account_json=None: [file_payload],
        downloader=lambda file_id, mime_type=None, service_account_json=None: _schedule_docx_bytes(),
        import_job_runner=_run_import_job,
        processed_file_marker=None,
    )
    assert first["processed"] == 1
    assert db_session.query(ScheduleSlot).count() == 1

    db_session.query(ScheduleSlot).delete(synchronize_session=False)
    db_session.commit()
    marker_calls: list[tuple[str, str]] = []

    resync = drive_intake.process_next_drive_intake_file(
        db_session,
        folder_url="https://drive.google.com/drive/folders/intake-folder",
        branch_id="main",
        file_lister=lambda folder_id, service_account_json=None: [file_payload],
        downloader=lambda file_id, mime_type=None, service_account_json=None: b"",
        import_job_runner=_run_import_job,
        processed_file_marker=lambda file_id, next_name, service_account_json=None: marker_calls.append((file_id, next_name)),
    )

    assert resync["processed"] == 1
    assert resync["resynced_schedule"] is True
    assert resync["marked_processed"] is True
    assert resync["job_id"] == first["job_id"]
    assert db_session.query(ScheduleSlot).count() == 1
    assert marker_calls == [("schedule-46-26", "46-26 Schedule [processed].docx")]


def test_drive_intake_defaults_xlsx_import_to_overwrite_for_corrections():
    assert drive_intake._default_import_mode(DocumentType.XLSX) == "overwrite"


def test_drive_intake_retries_failed_existing_schedule_job(db_session, tmp_path):
    failed_path = tmp_path / "failed-schedule.docx"
    failed_path.write_bytes(_schedule_docx_bytes())
    document = Document(
        branch_id="main",
        file_name="46-26 Розклад.docx",
        file_path=str(failed_path),
        file_type=DocumentType.DOCX,
        source="drive_intake",
        mime_type=drive_intake.GOOGLE_DRIVE_DOCX_MIME,
        hash_sha256="failed-once",
    )
    db_session.add(document)
    db_session.flush()
    failed_job = ImportJob(
        branch_id="main",
        idempotency_key=drive_intake._idempotency_key("main", "schedule-46-26", "2026-05-12T07:00:00Z"),
        document_id=document.id,
        status=JobStatus.FAILED,
        message="Попередня обробка впала",
        result_payload={
            "source": "drive_intake",
            "drive_file_id": "schedule-46-26",
            "drive_file_name": "46-26 Розклад.docx",
            "drive_modified_time": "2026-05-12T07:00:00Z",
            "import_mode": "overwrite",
        },
    )
    db_session.add(failed_job)
    db_session.commit()

    runner_calls: list[int] = []

    def rerun_existing_job(job_id: int) -> dict:
        runner_calls.append(job_id)
        return process_import_job_task.run(job_id)

    result = drive_intake.process_next_drive_intake_file(
        db_session,
        folder_url="https://drive.google.com/drive/folders/intake-folder",
        branch_id="main",
        file_lister=lambda folder_id, service_account_json=None: [
            {
                "id": "schedule-46-26",
                "name": "46-26 Розклад.docx",
                "mimeType": drive_intake.GOOGLE_DRIVE_DOCX_MIME,
                "modifiedTime": "2026-05-12T07:00:00Z",
                "webViewLink": "https://drive.google.com/file/d/schedule-46-26/view",
            }
        ],
        downloader=lambda file_id, mime_type=None, service_account_json=None: b"",
        import_job_runner=rerun_existing_job,
    )

    assert result["processed"] == 1
    assert result["retried_failed_job"] is True
    assert result["job_id"] == failed_job.id
    assert runner_calls == [failed_job.id]
    db_session.expire_all()
    assert db_session.get(ImportJob, failed_job.id).status == JobStatus.SUCCEEDED
    assert db_session.query(ScheduleSlot).count() == 1


def test_drive_intake_worker_reuses_active_journal_section_credentials(db_session, monkeypatch):
    captured: dict[str, object] = {}
    monkeypatch.setattr("app.services.drive_intake.settings.google_drive_service_account_json", "")
    section = JournalMonitorSection(
        branch_id="main",
        name="Журнали 2026",
        folder_url="https://drive.google.com/drive/folders/journals",
        folder_id="journals",
        service_account_json_encrypted=cipher.encrypt("section-service-account-json"),
    )
    db_session.add(section)
    db_session.commit()

    def fake_process_next_drive_intake_file(db, **kwargs):
        captured["service_account_json"] = kwargs.get("service_account_json")
        return {"processed": 0, "skipped_already_processed": 0, "skipped_unsupported": 0}

    monkeypatch.setattr("app.tasks.worker.process_next_drive_intake_file", fake_process_next_drive_intake_file)

    result = process_drive_intake_auto_task.run()

    assert result["processed"] == 0
    assert captured["service_account_json"] == "section-service-account-json"


def test_drive_intake_worker_processes_configured_batch_size(db_session, monkeypatch):
    calls: list[int] = []

    monkeypatch.setattr(
        "app.tasks.worker.settings",
        SimpleNamespace(
            google_drive_intake_auto_enabled=True,
            google_drive_intake_batch_size=2,
            imap_branch_id="main",
        ),
    )
    monkeypatch.setattr(
        "app.tasks.worker.resolve_drive_intake_service_account_json",
        lambda db, branch_id: None,
    )

    def fake_process_next_drive_intake_file(db, **kwargs):
        calls.append(len(calls) + 1)
        return {
            "processed": 1,
            "skipped_already_processed": 0,
            "skipped_unsupported": 0,
            "job_id": calls[-1],
            "filename": f"file-{calls[-1]}.xlsx",
        }

    monkeypatch.setattr("app.tasks.worker.process_next_drive_intake_file", fake_process_next_drive_intake_file)

    result = process_drive_intake_auto_task.run()

    assert calls == [1, 2]
    assert result["processed"] == 2
    assert result["batch_size"] == 2
    assert [item["job_id"] for item in result["items"]] == [1, 2]
    assert result["job_id"] == 2
    assert result["filename"] == "file-2.xlsx"


def test_drive_intake_failed_file_does_not_block_next_supported_file(db_session):
    files = [
        {
            "id": "bad-schedule-46-26",
            "name": "46-26 Bad Schedule.docx",
            "mimeType": drive_intake.GOOGLE_DRIVE_DOCX_MIME,
            "modifiedTime": "2026-05-12T07:00:00Z",
        },
        {
            "id": "good-schedule-90-26",
            "name": "90-26 Good Schedule.docx",
            "mimeType": drive_intake.GOOGLE_DRIVE_DOCX_MIME,
            "modifiedTime": "2026-05-12T07:01:00Z",
        },
    ]

    def downloader(file_id, mime_type=None, service_account_json=None):
        if file_id == "bad-schedule-46-26":
            return b"not-a-docx-package"
        return _schedule_docx_bytes("90-26")

    excluded_job_ids: set[int] = set()
    first = drive_intake.process_next_drive_intake_file(
        db_session,
        folder_url="https://drive.google.com/drive/folders/intake-folder",
        branch_id="main",
        file_lister=lambda folder_id, service_account_json=None: files,
        downloader=downloader,
        import_job_runner=_run_import_job,
        processed_file_marker=None,
        excluded_job_ids=excluded_job_ids,
    )
    assert first["processed"] == 1
    assert first["failed"] == 1
    excluded_job_ids.add(first["job_id"])

    second = drive_intake.process_next_drive_intake_file(
        db_session,
        folder_url="https://drive.google.com/drive/folders/intake-folder",
        branch_id="main",
        file_lister=lambda folder_id, service_account_json=None: files,
        downloader=downloader,
        import_job_runner=_run_import_job,
        processed_file_marker=None,
        excluded_job_ids=excluded_job_ids,
    )

    assert second["processed"] == 1
    assert second["status"] == JobStatus.SUCCEEDED.value
    assert second["filename"] == "90-26 Good Schedule.docx"
    assert db_session.query(ImportJob).filter(ImportJob.status == JobStatus.FAILED).count() == 1
    assert db_session.query(ScheduleSlot).count() == 1
