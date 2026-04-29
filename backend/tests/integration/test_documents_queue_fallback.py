import io

from docx import Document as DocxDocument
from openpyxl import Workbook

from app.api.routes import documents as documents_route
from app.models import Trainee


def _docx_bytes(text: str) -> bytes:
    stream = io.BytesIO()
    doc = DocxDocument()
    doc.add_paragraph(text)
    doc.save(stream)
    stream.seek(0)
    return stream.read()


def _contracts_xlsx_bytes() -> bytes:
    stream = io.BytesIO()
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Додаток"
    sheet.append(["Група 46-26 UX тест"])
    sheet.append([])
    sheet.append(["№", "ПІБ безробітного", "Дата народження", "№ Договору"])
    sheet.append([1, "Іваненко Іван Іванович", "01.02.2000", "46-26/001"])
    workbook.save(stream)
    stream.seek(0)
    return stream.read()


def _schedule_docx_bytes(group_code: str = "46-26") -> bytes:
    stream = io.BytesIO()
    doc = DocxDocument()
    doc.add_paragraph("1 пара - 9.30 – 11.05")
    doc.add_paragraph("за напрямом")
    doc.add_paragraph("UX тест")
    doc.add_paragraph(f"Група № {group_code}")
    doc.add_paragraph("з 21 жовтня 2025 року до 21 жовтня 2025 року")
    table = doc.add_table(rows=3, cols=6)
    for idx, value in enumerate(
        ["№п/п", "Назва предмета", "К-сть год.", "21.10", "22.10", "Прізвище, ім'я, по-батькові викладача"]
    ):
        table.cell(0, idx).text = value
    for idx, value in enumerate(["1", "Тема", "2", "1п/2год", "", "Коваль Іван Петрович"]):
        table.cell(1, idx).text = value
    for idx, value in enumerate(["", "Загальний обсяг навчального часу:", "2", "", "", ""]):
        table.cell(2, idx).text = value
    doc.save(stream)
    stream.seek(0)
    return stream.read()


def test_import_preview_summarizes_contract_xlsx(client, auth_headers):
    response = client.post(
        "/api/v1/documents/import/preview",
        headers=auth_headers,
        files={"file": ("contracts.xlsx", _contracts_xlsx_bytes(), "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")},
    )

    assert response.status_code == 200
    payload = response.json()
    assert payload["import_kind"] == "contracts"
    assert payload["rows"] == 1
    assert payload["sheet_name"] == "Додаток"
    assert payload["default_group_code"] == "46-26"
    assert payload["preview"][0]["ПІБ безробітного"] == "Іваненко Іван Іванович"


def test_import_preview_reports_contract_duplicates(client, auth_headers, db_session):
    db_session.add(
        Trainee(
            branch_id="main",
            first_name="Іван Іванович",
            last_name="Іваненко",
            contract_number="46-26/001",
            status="active",
        )
    )
    db_session.commit()

    response = client.post(
        "/api/v1/documents/import/preview",
        headers=auth_headers,
        files={"file": ("contracts.xlsx", _contracts_xlsx_bytes(), "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")},
    )

    assert response.status_code == 200
    payload = response.json()
    assert payload["duplicate_count"] == 1
    assert payload["new_count"] == 0
    assert payload["duplicate_preview"][0]["incoming_name"] == "Іваненко Іван Іванович"


def test_import_preview_summarizes_schedule_docx(client, auth_headers):
    response = client.post(
        "/api/v1/documents/import/preview",
        headers=auth_headers,
        files={"file": ("schedule.docx", _schedule_docx_bytes(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
    )

    assert response.status_code == 200
    payload = response.json()
    assert payload["import_kind"] == "schedule"
    assert payload["rows"] == 1
    assert payload["groups"][0]["code"] == "46-26"
    assert payload["groups"][0]["lessons"] == 1
    assert payload["groups"][0]["teachers"] == 1


def test_import_works_when_queue_is_unavailable(client, auth_headers, monkeypatch):
    def _raise_queue_error(job_id: int):
        raise RuntimeError("redis unavailable")

    monkeypatch.setattr(documents_route.process_import_job_task, "delay", _raise_queue_error)

    file_bytes = _docx_bytes("Тестовий DOCX документ")
    response = client.post(
        "/api/v1/documents/import",
        headers=auth_headers,
        files={
            "file": (
                "167-25 Організація трудових відносин – копія.docx",
                file_bytes,
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    assert response.status_code == 202
    payload = response.json()
    assert payload["status"] in {"running", "succeeded", "failed"}
    assert payload["status"] != "queued"


def test_import_accepts_overwrite_mode_form_field(client, auth_headers, monkeypatch):
    def _raise_queue_error(job_id: int):
        raise RuntimeError("redis unavailable")

    monkeypatch.setattr(documents_route.process_import_job_task, "delay", _raise_queue_error)

    file_bytes = _docx_bytes("Тестовий DOCX документ")
    response = client.post(
        "/api/v1/documents/import",
        headers=auth_headers,
        data={"update_existing_mode": "overwrite"},
        files={
            "file": (
                "167-25 Організація трудових відносин – копія.docx",
                file_bytes,
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    assert response.status_code == 202
    payload = response.json()
    assert payload["status"] in {"running", "succeeded", "failed"}


def test_export_works_when_queue_is_unavailable(client, auth_headers, monkeypatch):
    def _raise_queue_error(job_id: int):
        raise RuntimeError("redis unavailable")

    monkeypatch.setattr(documents_route.process_export_job_task, "delay", _raise_queue_error)

    response = client.post(
        "/api/v1/documents/export",
        headers=auth_headers,
        json={"report_type": "kpi", "export_format": "xlsx"},
    )
    assert response.status_code == 202
    payload = response.json()
    assert payload["status"] in {"running", "succeeded", "failed"}
    assert payload["status"] != "queued"
