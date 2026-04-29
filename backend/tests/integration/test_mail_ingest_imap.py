from email.message import EmailMessage
from datetime import datetime, timezone
from pathlib import Path

from docx import Document as DocxDocument
from openpyxl import Workbook

from app.models import (
    DraftStatus,
    Group,
    ImportJob,
    JobStatus,
    MailMessage,
    MailStatus,
    OCRResult,
    Room,
    ScheduleSlot,
    Subject,
    Teacher,
    Trainee,
)
from app.services import mail_ingest


class FakeIMAP4SSL:
    def __init__(self, raw_message: bytes):
        self.raw_message = raw_message
        self.fetch_payloads: list[str] = []
        self.logged_out = False

    def login(self, user: str, password: str):
        return "OK", [b"LOGIN"]

    def select(self, mailbox: str):
        return "OK", [b"1"]

    def search(self, charset, criteria):
        return "OK", [b"1"]

    def fetch(self, msg_id: bytes, payload: str):
        self.fetch_payloads.append(payload)
        return "OK", [(b"1 (UID 1001 RFC822 {256})", self.raw_message)]

    def logout(self):
        self.logged_out = True
        return "BYE", [b"LOGOUT"]


class FakeIMAP4SSLMulti:
    def __init__(self, raw_messages: list[bytes]):
        self.raw_messages = raw_messages
        self.fetch_payloads: list[str] = []
        self.logged_out = False

    def login(self, user: str, password: str):
        return "OK", [b"LOGIN"]

    def select(self, mailbox: str):
        return "OK", [b"1"]

    def search(self, charset, criteria):
        ids = b" ".join(str(index + 1).encode("ascii") for index in range(len(self.raw_messages)))
        return "OK", [ids]

    def fetch(self, msg_id: bytes, payload: str):
        self.fetch_payloads.append(payload)
        index = int(msg_id.decode("ascii")) - 1
        uid = 1001 + index
        raw_message = self.raw_messages[index]
        return "OK", [(f"{index + 1} (UID {uid} RFC822 {{256}})".encode("ascii"), raw_message)]

    def logout(self):
        self.logged_out = True
        return "BYE", [b"LOGOUT"]


def _build_message_with_docx_attachment(tmp_path: Path) -> bytes:
    docx_path = tmp_path / "attachment.docx"
    document = DocxDocument()
    document.add_paragraph("Іван Петренко")
    document.add_paragraph("Заява на навчання")
    document.save(docx_path)

    msg = EmailMessage()
    msg["Subject"] = "Тестовий лист"
    msg["From"] = "sender@example.com"
    msg["To"] = "inbox@example.com"
    msg["Message-ID"] = "<test-message-1@example.com>"
    msg.set_content("Вітаю, дивіться вкладення")

    msg.add_attachment(
        docx_path.read_bytes(),
        maintype="application",
        subtype="vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename="trainee.docx",
    )
    return msg.as_bytes()


def _build_contract_registry_xlsx(tmp_path: Path) -> Path:
    xlsx_path = tmp_path / "contracts.xlsx"
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Додаток"
    sheet.append(["Група 73-26 Штучний інтелект"])
    sheet.append([])
    sheet.append(
        [
            "№",
            "Центр зайнятості, який направив безробітного  на професійне навчання",
            "ПІБ безробітного",
            "Дата народження",
            "№ Договору",
            "Телефон",
        ]
    )
    sheet.append([1, "Львівський ОЦЗ", "Іваненко Іван Іванович", "01.02.2000", "73-26/001", "+380501112233"])
    workbook.save(xlsx_path)
    return xlsx_path


def _build_schedule_docx(tmp_path: Path) -> Path:
    docx_path = tmp_path / "schedule.docx"
    document = DocxDocument()
    document.add_paragraph("1 пара - 9.30 – 11.05")
    document.add_paragraph("2 пара – 11.10 – 12.45")
    document.add_paragraph("за напрямом")
    document.add_paragraph("Штучний інтелект")
    document.add_paragraph("Група № 162-25")
    document.add_paragraph("з 23 червня 2025 року до 24 червня 2025 року")

    table = document.add_table(rows=3, cols=6)
    for idx, value in enumerate(
        ["№п/п", "Назва предмета", "К-сть год.", "23.06", "24.06", "Прізвище, ім'я, по-батькові викладача"]
    ):
        table.cell(0, idx).text = value
    for idx, value in enumerate(["1", "Тема групи 162", "2", "1п/1год", "2п/1год", "Паращук Світлана Зеновіївна"]):
        table.cell(1, idx).text = value
    for idx, value in enumerate(["", "Загальний обсяг навчального часу:", "2", "", "", ""]):
        table.cell(2, idx).text = value

    document.save(docx_path)
    return docx_path


def _build_message_with_attachment(
    sender: str,
    filename: str,
    payload: bytes,
    subtype: str,
    message_id: str,
) -> bytes:
    msg = EmailMessage()
    msg["Subject"] = "Тестовий лист"
    msg["From"] = sender
    msg["To"] = "inbox@example.com"
    msg["Message-ID"] = message_id
    msg.set_content("Вітаю, дивіться вкладення")
    msg.add_attachment(
        payload,
        maintype="application",
        subtype=subtype,
        filename=filename,
    )
    return msg.as_bytes()


def _build_message_with_explicit_id_and_docx(sender: str, message_id: str, tmp_path: Path, filename: str) -> bytes:
    docx_path = tmp_path / filename
    document = DocxDocument()
    document.add_paragraph("Тестовий текст для обробки")
    document.save(docx_path)

    msg = EmailMessage()
    msg["Subject"] = "Лист з вкладенням"
    msg["From"] = sender
    msg["To"] = "inbox@example.com"
    msg["Message-ID"] = message_id
    msg.set_content("Перевірка обробки")
    msg.add_attachment(
        docx_path.read_bytes(),
        maintype="application",
        subtype="vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=filename,
    )
    return msg.as_bytes()


def test_ingest_mailbox_creates_draft_for_docx_attachment(db_session, monkeypatch, tmp_path: Path):
    raw_message = _build_message_with_docx_attachment(tmp_path)
    fake_client = FakeIMAP4SSL(raw_message)

    monkeypatch.setattr(mail_ingest.settings, "imap_host", "imap.example.com")
    monkeypatch.setattr(mail_ingest.settings, "imap_port", 993)
    monkeypatch.setattr(mail_ingest.settings, "imap_user", "inbox@example.com")
    monkeypatch.setattr(mail_ingest.settings, "imap_password", "secret")
    monkeypatch.setattr(mail_ingest.settings, "imap_mailbox", "INBOX")
    monkeypatch.setattr(mail_ingest.imaplib, "IMAP4_SSL", lambda host, port: fake_client)

    result = mail_ingest.ingest_mailbox(db_session)
    assert result["processed"] == 1
    assert fake_client.fetch_payloads == ["(UID BODY.PEEK[])"]

    message = db_session.query(MailMessage).filter(MailMessage.message_id == "<test-message-1@example.com>").one()
    assert message.status == MailStatus.PROCESSED
    assert message.subject == "Тестовий лист"
    assert message.raw_document_id is not None

    import_job = db_session.query(ImportJob).order_by(ImportJob.id.desc()).first()
    assert import_job is None
    draft = db_session.query(OCRResult).order_by(OCRResult.id.desc()).first()
    assert draft is not None
    assert draft.status == DraftStatus.PENDING
    assert draft.confidence > 0.5
    assert "Іван" in draft.extracted_text


def test_ingest_mailbox_auto_imports_contracts_for_configured_sender(db_session, monkeypatch, tmp_path: Path):
    xlsx_path = _build_contract_registry_xlsx(tmp_path)
    raw_message = _build_message_with_attachment(
        sender="Львівський центр ПТО ДСЗ <lcptodcz@gmail.com>",
        filename="Договори 73-26 Штучний інтелект.xlsx",
        payload=xlsx_path.read_bytes(),
        subtype="vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        message_id="<test-message-contracts@example.com>",
    )
    fake_client = FakeIMAP4SSL(raw_message)

    monkeypatch.setattr(mail_ingest.settings, "imap_host", "imap.example.com")
    monkeypatch.setattr(mail_ingest.settings, "imap_port", 993)
    monkeypatch.setattr(mail_ingest.settings, "imap_user", "inbox@example.com")
    monkeypatch.setattr(mail_ingest.settings, "imap_password", "secret")
    monkeypatch.setattr(mail_ingest.settings, "imap_mailbox", "INBOX")
    monkeypatch.setattr(mail_ingest.settings, "imap_contract_sender_name", "Львівський центр ПТО ДСЗ")
    monkeypatch.setattr(mail_ingest.settings, "imap_contract_sender_email", "lcptodcz@gmail.com")
    monkeypatch.setattr(mail_ingest.settings, "imap_contract_attachment_prefix", "Договори")
    monkeypatch.setattr(mail_ingest.settings, "imap_contract_update_mode", "overwrite")
    monkeypatch.setattr(mail_ingest.imaplib, "IMAP4_SSL", lambda host, port: fake_client)

    result = mail_ingest.ingest_mailbox(db_session)
    assert result["processed"] == 1

    import_job = db_session.query(ImportJob).order_by(ImportJob.id.desc()).first()
    assert import_job is not None
    assert import_job.status == JobStatus.SUCCEEDED
    assert import_job.result_payload is not None
    assert import_job.result_payload.get("source") == "mail_auto_contracts"
    assert import_job.result_payload.get("group_code_hint") == "73-26"

    trainee = db_session.query(Trainee).filter(Trainee.contract_number == "73-26/001").first()
    assert trainee is not None


def test_ingest_mailbox_auto_imports_when_group_code_is_before_keyword(db_session, monkeypatch, tmp_path: Path):
    xlsx_path = _build_contract_registry_xlsx(tmp_path)
    raw_message = _build_message_with_attachment(
        sender="Львівський центр ПТО ДСЗ (навчання) <lcptodcz@gmail.com>",
        filename="73–26 Штучний інтелект Договори.xlsx",
        payload=xlsx_path.read_bytes(),
        subtype="vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        message_id="<test-message-contracts-before-keyword@example.com>",
    )
    fake_client = FakeIMAP4SSL(raw_message)

    monkeypatch.setattr(mail_ingest.settings, "imap_host", "imap.example.com")
    monkeypatch.setattr(mail_ingest.settings, "imap_port", 993)
    monkeypatch.setattr(mail_ingest.settings, "imap_user", "inbox@example.com")
    monkeypatch.setattr(mail_ingest.settings, "imap_password", "secret")
    monkeypatch.setattr(mail_ingest.settings, "imap_mailbox", "INBOX")
    monkeypatch.setattr(mail_ingest.settings, "imap_contract_sender_name", "Львівський центр ПТО ДСЗ")
    monkeypatch.setattr(mail_ingest.settings, "imap_contract_sender_email", "lcptodcz@gmail.com")
    monkeypatch.setattr(mail_ingest.settings, "imap_contract_attachment_prefix", "Договори")
    monkeypatch.setattr(mail_ingest.imaplib, "IMAP4_SSL", lambda host, port: fake_client)

    result = mail_ingest.ingest_mailbox(db_session)
    assert result["processed"] == 1

    import_job = db_session.query(ImportJob).order_by(ImportJob.id.desc()).first()
    assert import_job is not None
    assert import_job.status == JobStatus.SUCCEEDED
    assert import_job.result_payload is not None
    assert import_job.result_payload.get("source") == "mail_auto_contracts"
    assert import_job.result_payload.get("group_code_hint") == "73-26"


def test_ingest_mailbox_skips_non_matching_excel_attachment(db_session, monkeypatch, tmp_path: Path):
    xlsx_path = _build_contract_registry_xlsx(tmp_path)
    raw_message = _build_message_with_attachment(
        sender="Львівський центр ПТО ДСЗ <lcptodcz@gmail.com>",
        filename="Слухачі 73-26.xlsx",
        payload=xlsx_path.read_bytes(),
        subtype="vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        message_id="<test-message-skip-excel@example.com>",
    )
    fake_client = FakeIMAP4SSL(raw_message)

    monkeypatch.setattr(mail_ingest.settings, "imap_host", "imap.example.com")
    monkeypatch.setattr(mail_ingest.settings, "imap_port", 993)
    monkeypatch.setattr(mail_ingest.settings, "imap_user", "inbox@example.com")
    monkeypatch.setattr(mail_ingest.settings, "imap_password", "secret")
    monkeypatch.setattr(mail_ingest.settings, "imap_mailbox", "INBOX")
    monkeypatch.setattr(mail_ingest.settings, "imap_contract_sender_name", "Львівський центр ПТО ДСЗ")
    monkeypatch.setattr(mail_ingest.settings, "imap_contract_sender_email", "lcptodcz@gmail.com")
    monkeypatch.setattr(mail_ingest.settings, "imap_contract_attachment_prefix", "Договори")
    monkeypatch.setattr(mail_ingest.imaplib, "IMAP4_SSL", lambda host, port: fake_client)

    result = mail_ingest.ingest_mailbox(db_session)
    assert result["processed"] == 1

    import_job = db_session.query(ImportJob).order_by(ImportJob.id.desc()).first()
    assert import_job is None
    message = db_session.query(MailMessage).filter(MailMessage.message_id == "<test-message-skip-excel@example.com>").one()
    assert message.status == MailStatus.PROCESSED
    assert message.snippet is not None
    assert "пропущено" in message.snippet.lower()


def test_ingest_mailbox_reprocesses_unread_schedule_even_when_group_has_other_slots(
    db_session, monkeypatch, tmp_path: Path
):
    group = Group(branch_id="main", code="162-25", name="Existing group")
    teacher = Teacher(branch_id="main", last_name="Існуючий", first_name="Викладач", hourly_rate=0)
    subject = Subject(branch_id="main", name="Старий предмет", hours_total=1)
    room = Room(branch_id="main", name="Стара аудиторія", capacity=20)
    db_session.add_all([group, teacher, subject, room])
    db_session.flush()
    db_session.add(
        ScheduleSlot(
            group_id=group.id,
            teacher_id=teacher.id,
            subject_id=subject.id,
            room_id=room.id,
            starts_at=datetime(2025, 1, 10, 9, 0, tzinfo=timezone.utc),
            ends_at=datetime(2025, 1, 10, 10, 0, tzinfo=timezone.utc),
            pair_number=1,
            academic_hours=1,
        )
    )
    db_session.commit()

    schedule_path = _build_schedule_docx(tmp_path)
    raw_message = _build_message_with_attachment(
        sender="Львівський центр ПТО ДСЗ <lcptodcz@gmail.com>",
        filename="Розклад 162-25.docx",
        payload=schedule_path.read_bytes(),
        subtype="vnd.openxmlformats-officedocument.wordprocessingml.document",
        message_id="<test-message-schedule-retry@example.com>",
    )
    fake_client = FakeIMAP4SSL(raw_message)

    monkeypatch.setattr(mail_ingest.settings, "imap_host", "imap.example.com")
    monkeypatch.setattr(mail_ingest.settings, "imap_port", 993)
    monkeypatch.setattr(mail_ingest.settings, "imap_user", "inbox@example.com")
    monkeypatch.setattr(mail_ingest.settings, "imap_password", "secret")
    monkeypatch.setattr(mail_ingest.settings, "imap_mailbox", "INBOX")
    monkeypatch.setattr(mail_ingest.imaplib, "IMAP4_SSL", lambda host, port: fake_client)

    result = mail_ingest.ingest_mailbox(db_session)
    assert result["processed"] == 1

    import_job = db_session.query(ImportJob).order_by(ImportJob.id.desc()).first()
    assert import_job is not None
    assert import_job.status == JobStatus.SUCCEEDED
    assert import_job.result_payload is not None
    assert import_job.result_payload.get("source") == "mail_auto_schedules"
    assert import_job.result_payload["import_result"]["created_slots"] == 2

    imported_slots = (
        db_session.query(ScheduleSlot)
        .filter(ScheduleSlot.group_id == group.id, ScheduleSlot.starts_at >= datetime(2025, 6, 23, tzinfo=timezone.utc))
        .all()
    )
    assert len(imported_slots) == 2


def test_ingest_mailbox_processes_two_messages_with_same_message_id(db_session, monkeypatch, tmp_path: Path):
    duplicated_message_id = "<same-message-id@example.com>"
    raw_message_one = _build_message_with_explicit_id_and_docx(
        sender="sender@example.com",
        message_id=duplicated_message_id,
        tmp_path=tmp_path,
        filename="one.docx",
    )
    raw_message_two = _build_message_with_explicit_id_and_docx(
        sender="sender@example.com",
        message_id=duplicated_message_id,
        tmp_path=tmp_path,
        filename="two.docx",
    )
    fake_client = FakeIMAP4SSLMulti([raw_message_one, raw_message_two])

    monkeypatch.setattr(mail_ingest.settings, "imap_host", "imap.example.com")
    monkeypatch.setattr(mail_ingest.settings, "imap_port", 993)
    monkeypatch.setattr(mail_ingest.settings, "imap_user", "inbox@example.com")
    monkeypatch.setattr(mail_ingest.settings, "imap_password", "secret")
    monkeypatch.setattr(mail_ingest.settings, "imap_mailbox", "INBOX")
    monkeypatch.setattr(mail_ingest.imaplib, "IMAP4_SSL", lambda host, port: fake_client)

    result1 = mail_ingest.ingest_mailbox(db_session)
    assert result1["processed"] == 1

    result2 = mail_ingest.ingest_mailbox(db_session)
    assert result2["processed"] == 1

    messages = db_session.query(MailMessage).filter(MailMessage.subject == "Лист з вкладенням").all()
    assert len(messages) == 2
