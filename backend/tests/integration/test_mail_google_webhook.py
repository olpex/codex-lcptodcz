import base64
import io

from docx import Document as DocxDocument
from openpyxl import Workbook

from app.api.routes import mail as mail_routes
from app.models import Group, ImportJob, JobStatus, ScheduleSlot, Trainee


def _contracts_xlsx_bytes() -> bytes:
    stream = io.BytesIO()
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Додаток"
    sheet.append(["Група 73-26 Штучний інтелект"])
    sheet.append([])
    sheet.append(
        [
            "№",
            "Центр зайнятості, який направив безробітного  на професійне навчання",
            "ПІБ безробітного",
            "Дата народження",
            "№ Договору",
            "Телефон",
        ]
    )
    sheet.append([1, "Львівський ОЦЗ", "Іваненко Іван Іванович", "01.02.2000", "73-26/001", "+380501112233"])
    workbook.save(stream)
    stream.seek(0)
    return stream.read()


def _schedule_docx_bytes() -> bytes:
    stream = io.BytesIO()
    document = DocxDocument()
    document.add_paragraph("1 пара - 9.30 – 11.05")
    document.add_paragraph("2 пара – 11.10 – 12.45")
    document.add_paragraph("за напрямом")
    document.add_paragraph("Організація трудових відносин")
    document.add_paragraph("Група № 167-25")
    document.add_paragraph("з 21 жовтня 2025 року до 22 жовтня 2025 року")

    table = document.add_table(rows=3, cols=6)
    for idx, value in enumerate(
        ["№п/п", "Назва предмета", "К-сть год.", "21.10", "22.10", "Прізвище, ім'я, по-батькові викладача"]
    ):
        table.cell(0, idx).text = value
    for idx, value in enumerate(["1", "Трудовий договір", "2", "1п/1год", "2п/1год", "Штогрин Лілія Володимирівна"]):
        table.cell(1, idx).text = value
    for idx, value in enumerate(["", "Загальний обсяг навчального часу:", "2", "", "", ""]):
        table.cell(2, idx).text = value

    document.save(stream)
    stream.seek(0)
    return stream.read()


def test_google_webhook_imports_contract_file(client, db_session, monkeypatch):
    def _raise_queue_error(job_id: int):
        raise RuntimeError("redis unavailable")

    monkeypatch.setattr(mail_routes.process_import_job_task, "delay", _raise_queue_error)
    monkeypatch.setattr(mail_routes.settings, "mail_webhook_secret", "mail-webhook-secret")
    monkeypatch.setattr(mail_routes.settings, "imap_branch_id", "main")
    monkeypatch.setattr(mail_routes.settings, "imap_contract_sender_name", "Львівський центр ПТО ДСЗ")
    monkeypatch.setattr(mail_routes.settings, "imap_contract_sender_email", "lcptodcz@gmail.com")
    monkeypatch.setattr(mail_routes.settings, "imap_contract_attachment_prefix", "Договори")

    response = client.post(
        "/api/v1/mail/google-webhook/contracts",
        headers={"Authorization": "Bearer mail-webhook-secret"},
        data={
            "sender_email": "lcptodcz@gmail.com",
            "sender_name": "Львівський центр ПТО ДСЗ",
            "subject": "Договори групи",
            "message_id": "<google-webhook-test-1@example.com>",
            "update_existing_mode": "overwrite",
        },
        files={
            "file": (
                "73-26 Договори Штучний інтелект.xlsx",
                _contracts_xlsx_bytes(),
                "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            )
        },
    )
    assert response.status_code == 202
    payload = response.json()
    assert payload["status"] in {"running", "succeeded", "failed"}
    assert payload["result_payload"]["source"] == "mail_google_script"
    assert payload["result_payload"]["group_code_hint"] == "73-26"

    job = db_session.query(ImportJob).order_by(ImportJob.id.desc()).first()
    assert job is not None
    trainee = db_session.query(Trainee).filter(Trainee.contract_number == "73-26/001").first()
    assert trainee is not None


def test_google_webhook_rejects_sender_mismatch(client, monkeypatch):
    monkeypatch.setattr(mail_routes.settings, "mail_webhook_secret", "mail-webhook-secret")
    monkeypatch.setattr(mail_routes.settings, "imap_contract_sender_name", "Львівський центр ПТО ДСЗ")
    monkeypatch.setattr(mail_routes.settings, "imap_contract_sender_email", "lcptodcz@gmail.com")
    monkeypatch.setattr(mail_routes.settings, "imap_contract_attachment_prefix", "Договори")

    response = client.post(
        "/api/v1/mail/google-webhook/contracts",
        headers={"Authorization": "Bearer mail-webhook-secret"},
        data={
            "sender_email": "other@gmail.com",
            "sender_name": "Інший відправник",
            "subject": "Договори групи",
        },
        files={
            "file": (
                "73-26 Договори.xlsx",
                _contracts_xlsx_bytes(),
                "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            )
        },
    )
    assert response.status_code == 400
    assert "Відправник" in response.json()["detail"]


def test_google_webhook_imports_docx_inline_without_queue(client, db_session, monkeypatch):
    def _fail_if_queue_used(job_id: int):
        raise AssertionError("DOCX schedule imports must not depend on Celery queue")

    monkeypatch.setattr(mail_routes.settings, "mail_webhook_secret", "mail-webhook-secret")
    monkeypatch.setattr(mail_routes.settings, "imap_contract_sender_name", "Львівський центр ПТО ДСЗ")
    monkeypatch.setattr(mail_routes.settings, "imap_contract_sender_email", "lcptodcz@gmail.com")
    monkeypatch.setattr(mail_routes.settings, "imap_contract_attachment_prefix", "Договори")
    monkeypatch.setattr(mail_routes.process_import_job_task, "delay", _fail_if_queue_used)

    response = client.post(
        "/api/v1/mail/google-webhook/contracts",
        headers={"Authorization": "Bearer mail-webhook-secret"},
        data={
            "sender_email": "lcptodcz@gmail.com",
            "sender_name": "Львівський центр ПТО ДСЗ",
            "subject": "Файл групи",
            "message_id": "<google-webhook-docx-test-1@example.com>",
            "update_existing_mode": "overwrite",
        },
        files={
            "file": (
                "167-26 Осінній модуль.docx",
                _schedule_docx_bytes(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )

    assert response.status_code == 202
    assert response.json()["status"] == JobStatus.SUCCEEDED.value

    group = db_session.query(Group).filter(Group.code == "167-25").one()
    assert db_session.query(ScheduleSlot).filter(ScheduleSlot.group_id == group.id).count() == 2


def test_google_webhook_accepts_docx_with_schedule_keyword_and_fwd_prefix(client, monkeypatch):
    monkeypatch.setattr(mail_routes.settings, "mail_webhook_secret", "mail-webhook-secret")
    monkeypatch.setattr(mail_routes.settings, "imap_contract_sender_name", "Львівський центр ПТО ДСЗ")
    monkeypatch.setattr(mail_routes.settings, "imap_contract_sender_email", "lcptodcz@gmail.com")
    monkeypatch.setattr(mail_routes.settings, "imap_contract_attachment_prefix", "Договори")

    response = client.post(
        "/api/v1/mail/google-webhook/contracts",
        headers={"Authorization": "Bearer mail-webhook-secret"},
        data={
            "sender_email": "lcptodcz@gmail.com",
            "sender_name": "Львівський центр ПТО ДСЗ",
            "subject": "Fwd: Розклад",
            "message_id": "<google-webhook-docx-test-3@example.com>",
            "update_existing_mode": "overwrite",
        },
        files={
            "file": (
                "162-25 Штучний інтелект – копія.docx",
                _schedule_docx_bytes(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    assert response.status_code == 202


def test_google_webhook_accepts_docx_with_schedule_keyword(client, monkeypatch):
    monkeypatch.setattr(mail_routes.settings, "mail_webhook_secret", "mail-webhook-secret")
    monkeypatch.setattr(mail_routes.settings, "imap_contract_sender_name", "Львівський центр ПТО ДСЗ")
    monkeypatch.setattr(mail_routes.settings, "imap_contract_sender_email", "lcptodcz@gmail.com")
    monkeypatch.setattr(mail_routes.settings, "imap_contract_attachment_prefix", "Договори")

    response = client.post(
        "/api/v1/mail/google-webhook/contracts",
        headers={"Authorization": "Bearer mail-webhook-secret"},
        data={
            "sender_email": "lcptodcz@gmail.com",
            "sender_name": "Львівський центр ПТО ДСЗ",
            "subject": "Розклад групи",
            "message_id": "<google-webhook-docx-test-2@example.com>",
            "update_existing_mode": "overwrite",
        },
        files={
            "file": (
                "167-26 Розклад осінній модуль.docx",
                _schedule_docx_bytes(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    assert response.status_code == 202


def test_google_webhook_accepts_docx_when_subject_is_empty(client, monkeypatch):
    monkeypatch.setattr(mail_routes.settings, "mail_webhook_secret", "mail-webhook-secret")
    monkeypatch.setattr(mail_routes.settings, "imap_contract_sender_name", "Львівський центр ПТО ДСЗ")
    monkeypatch.setattr(mail_routes.settings, "imap_contract_sender_email", "lcptodcz@gmail.com")
    monkeypatch.setattr(mail_routes.settings, "imap_contract_attachment_prefix", "Договори")

    response = client.post(
        "/api/v1/mail/google-webhook/contracts",
        headers={"Authorization": "Bearer mail-webhook-secret"},
        data={
            "sender_email": "lcptodcz@gmail.com",
            "sender_name": "Львівський центр ПТО ДСЗ",
            "subject": "",
            "message_id": "<google-webhook-docx-empty-subject-test@example.com>",
            "update_existing_mode": "overwrite",
        },
        files={
            "file": (
                "Розклад 162-25 Штучний інтелект.docx",
                _schedule_docx_bytes(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    assert response.status_code == 202


def test_google_webhook_reimports_docx_when_data_missing_even_with_same_message_id(client, db_session, monkeypatch):
    monkeypatch.setattr(mail_routes.settings, "mail_webhook_secret", "mail-webhook-secret")
    monkeypatch.setattr(mail_routes.settings, "imap_contract_sender_name", "Львівський центр ПТО ДСЗ")
    monkeypatch.setattr(mail_routes.settings, "imap_contract_sender_email", "lcptodcz@gmail.com")
    monkeypatch.setattr(mail_routes.settings, "imap_contract_attachment_prefix", "Договори")

    payload = {
        "sender_email": "lcptodcz@gmail.com",
        "sender_name": "Львівський центр ПТО ДСЗ",
        "subject": "Розклад групи",
        "message_id": "<google-webhook-docx-reimport-test@example.com>",
        "update_existing_mode": "overwrite",
    }
    files = {
        "file": (
            "167-26 Розклад осінній модуль.docx",
            _schedule_docx_bytes(),
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
    }

    first = client.post(
        "/api/v1/mail/google-webhook/contracts",
        headers={"Authorization": "Bearer mail-webhook-secret"},
        data=payload,
        files=files,
    )
    assert first.status_code == 202
    first_job_id = first.json()["id"]

    group = db_session.query(Group).filter(Group.code == "167-25").one()
    db_session.query(ScheduleSlot).filter(ScheduleSlot.group_id == group.id).delete(synchronize_session=False)
    db_session.commit()

    second = client.post(
        "/api/v1/mail/google-webhook/contracts",
        headers={"Authorization": "Bearer mail-webhook-secret"},
        data=payload,
        files=files,
    )
    assert second.status_code == 202
    second_job_id = second.json()["id"]

    assert second_job_id != first_job_id
    assert db_session.query(ImportJob).count() == 2
    assert db_session.query(ScheduleSlot).filter(ScheduleSlot.group_id == group.id).count() == 2


def test_gmail_api_webhook_accepts_docx_when_subject_is_missing(client, monkeypatch):
    monkeypatch.setattr(mail_routes.settings, "mail_webhook_secret", "mail-webhook-secret")
    monkeypatch.setattr(mail_routes.settings, "imap_contract_sender_name", "Львівський центр ПТО ДСЗ")
    monkeypatch.setattr(mail_routes.settings, "imap_contract_sender_email", "lcptodcz@gmail.com")

    file_base64 = base64.urlsafe_b64encode(_schedule_docx_bytes()).decode("ascii")
    response = client.post(
        "/api/v1/mail/gmail-api-webhook/contracts",
        headers={"Authorization": "Bearer mail-webhook-secret"},
        json={
            "filename": "розклад 162-25 Штучний інтелект.docx",
            "messageId": "<gmail-api-docx-test-1@example.com>",
            "fileBase64": file_base64,
            "subject": None,
        },
    )

    assert response.status_code == 202
