from pathlib import Path

import pytest
from docx import Document as DocxDocument

from app.models import Group, ScheduleSlot, Subject, Teacher
from app.services.schedule_import import import_schedule_docx, parse_schedule_docx


def _build_schedule_docx(path: Path) -> None:
    doc = DocxDocument()
    doc.add_paragraph("1 пара - 9.30 – 11.05")
    doc.add_paragraph("2 пара – 11.10 – 12.45")
    doc.add_paragraph("3 пара – 13.05 – 14.40")
    doc.add_paragraph("4 пара – 14.45– 16.20")
    doc.add_paragraph("за напрямом")
    doc.add_paragraph("Організація трудових відносин")
    doc.add_paragraph("Група № 167-25")
    doc.add_paragraph("з    21 жовтня – 24 жовтня 2025 р.")

    table = doc.add_table(rows=4, cols=8)
    headers = [
        "№п/п",
        "Назва предмета",
        "К-сть год.",
        "21.10",
        "22.10",
        "23.10",
        "24.10",
        "Прізвище, ім'я, по-батькові викладача",
    ]
    for idx, value in enumerate(headers):
        table.cell(0, idx).text = value

    row1 = ["1", "Трудовий договір", "3", "", "1п/2год", "", "2п/1год", "Штогрин Лілія Володимирівна"]
    row2 = ["2", "Охорона праці", "2", "3п/2год", "", "", "", "Костів Артур Романович"]
    total = ["", "Загальний обсяг навчального часу:", "5", "2", "2", "0", "1", ""]

    for idx, value in enumerate(row1):
        table.cell(1, idx).text = value
    for idx, value in enumerate(row2):
        table.cell(2, idx).text = value
    for idx, value in enumerate(total):
        table.cell(3, idx).text = value

    doc.save(path)


def _build_schedule_docx_with_short_year_and_merged_teacher(path: Path) -> None:
    doc = DocxDocument()
    doc.add_paragraph("1 пара - 9.30 – 11.05")
    doc.add_paragraph("2 пара – 11.10 – 12.45")
    doc.add_paragraph("за напрямом")
    doc.add_paragraph("Організація трудових відносин")
    doc.add_paragraph("Група № 167-26")
    doc.add_paragraph("з 21 жовтня – 24 жовтня 25-го року")

    table = doc.add_table(rows=4, cols=8)
    headers = [
        "№п/п",
        "Назва предмета",
        "К-сть год.",
        "21.10",
        "22.10",
        "23.10",
        "24.10",
        "Прізвище, ім'я, по-батькові викладача",
    ]
    for idx, value in enumerate(headers):
        table.cell(0, idx).text = value

    row1 = ["1", "Трудовий договір", "4", "1п/2год", "", "", "", "Войтехівська Галина Михайлівна"]
    row2 = ["2", "Охорона праці", "2", "", "2п/2год", "", "", ""]
    total = ["", "Загальний обсяг навчального часу:", "6", "2", "2", "0", "0", ""]

    for idx, value in enumerate(row1):
        table.cell(1, idx).text = value
    for idx, value in enumerate(row2):
        table.cell(2, idx).text = value
    for idx, value in enumerate(total):
        table.cell(3, idx).text = value

    doc.save(path)


def _build_schedule_docx_with_two_groups(path: Path) -> None:
    doc = DocxDocument()
    doc.add_paragraph("1 пара - 9.30 – 11.05")
    doc.add_paragraph("2 пара – 11.10 – 12.45")
    doc.add_paragraph("за напрямом")
    doc.add_paragraph("Штучний інтелект")
    doc.add_paragraph("Група № 162-25")
    doc.add_paragraph("з 23 червня 2025 року до 26 червня 2025 року")

    table1 = doc.add_table(rows=3, cols=6)
    for idx, value in enumerate(
        ["№п/п", "Назва предмета", "К-сть год.", "23.06", "24.06", "Прізвище, ім'я, по-батькові викладача"]
    ):
        table1.cell(0, idx).text = value
    for idx, value in enumerate(["1", "Тема групи 162", "2", "1п/1год", "2п/1год", "Паращук Світлана Зеновіївна"]):
        table1.cell(1, idx).text = value
    for idx, value in enumerate(["", "Загальний обсяг навчального часу:", "2", "", "", ""]):
        table1.cell(2, idx).text = value

    doc.add_paragraph("за напрямом")
    doc.add_paragraph("Токар")
    doc.add_paragraph("Група № 47п-25")
    doc.add_paragraph("з 23 червня 2025 року до 26 червня 2025 року")

    table2 = doc.add_table(rows=3, cols=6)
    for idx, value in enumerate(
        ["№п/п", "Назва предмета", "К-сть год.", "23.06", "24.06", "Прізвище, ім'я, по-батькові викладача"]
    ):
        table2.cell(0, idx).text = value
    for idx, value in enumerate(["1", "Тема групи 47п", "2", "1п/1год", "2п/1год", "Паращук Світлана Зеновіївна"]):
        table2.cell(1, idx).text = value
    for idx, value in enumerate(["", "Загальний обсяг навчального часу:", "2", "", "", ""]):
        table2.cell(2, idx).text = value

    doc.save(path)


def _build_schedule_docx_with_two_groups_in_table_headers(path: Path) -> None:
    doc = DocxDocument()
    doc.add_paragraph("1 пара - 9.30 – 11.05")
    doc.add_paragraph("2 пара – 11.10 – 12.45")

    table1 = doc.add_table(rows=4, cols=6)
    table1.cell(0, 0).text = "Розклад занять. Група № 162-25"
    for idx, value in enumerate(
        ["№п/п", "Назва предмета", "К-сть год.", "23.06", "24.06", "Прізвище, ім'я, по-батькові викладача"]
    ):
        table1.cell(1, idx).text = value
    for idx, value in enumerate(["1", "Тема групи 162", "2", "1п/1год", "2п/1год", "Паращук Світлана Зеновіївна"]):
        table1.cell(2, idx).text = value
    for idx, value in enumerate(["", "Загальний обсяг навчального часу:", "2", "", "", ""]):
        table1.cell(3, idx).text = value

    table2 = doc.add_table(rows=4, cols=6)
    table2.cell(0, 0).text = "Розклад занять. Група № 47п-25"
    for idx, value in enumerate(
        ["№п/п", "Назва предмета", "К-сть год.", "23.06", "24.06", "Прізвище, ім'я, по-батькові викладача"]
    ):
        table2.cell(1, idx).text = value
    for idx, value in enumerate(["1", "Тема групи 47п", "2", "1п/1год", "2п/1год", "Паращук Світлана Зеновіївна"]):
        table2.cell(2, idx).text = value
    for idx, value in enumerate(["", "Загальний обсяг навчального часу:", "2", "", "", ""]):
        table2.cell(3, idx).text = value

    doc.save(path)


def _build_schedule_docx_with_dot_separator_group_code(path: Path) -> None:
    doc = DocxDocument()
    doc.add_paragraph("1 пара - 9.30 – 11.05")
    doc.add_paragraph("2 пара – 11.10 – 12.45")
    doc.add_paragraph("за напрямом")
    doc.add_paragraph("Штучний інтелект: розвиток кар'єри")
    doc.add_paragraph("Група № 162.25")
    doc.add_paragraph("з 23 червня 2025 року до 26 червня 2025 року")

    table1 = doc.add_table(rows=3, cols=6)
    for idx, value in enumerate(
        ["№п/п", "Назва предмета", "К-сть год.", "23.06", "24.06", "Прізвище, ім'я, по-батькові викладача"]
    ):
        table1.cell(0, idx).text = value
    for idx, value in enumerate(["1", "Тема групи 162", "2", "1п/1год", "2п/1год", "Паращук Світлана Зеновіївна"]):
        table1.cell(1, idx).text = value
    for idx, value in enumerate(["", "Загальний обсяг навчального часу:", "2", "", "", ""]):
        table1.cell(2, idx).text = value

    doc.add_paragraph("за напрямом")
    doc.add_paragraph("Токар")
    doc.add_paragraph("Група № 47п-25")
    doc.add_paragraph("з 23 червня 2025 року до 26 червня 2025 року")

    table2 = doc.add_table(rows=3, cols=6)
    for idx, value in enumerate(
        ["№п/п", "Назва предмета", "К-сть год.", "23.06", "24.06", "Прізвище, ім'я, по-батькові викладача"]
    ):
        table2.cell(0, idx).text = value
    for idx, value in enumerate(["1", "Тема групи 47п", "2", "1п/1год", "2п/1год", "Паращук Світлана Зеновіївна"]):
        table2.cell(1, idx).text = value
    for idx, value in enumerate(["", "Загальний обсяг навчального часу:", "2", "", "", ""]):
        table2.cell(2, idx).text = value

    doc.save(path)


def _build_schedule_docx_with_duplicated_course_title(path: Path) -> None:
    title = "Штучний інтелект: розвиток кар'єри та професійне зростання"
    doc = DocxDocument()
    doc.add_paragraph("1 пара - 9.30 – 11.05")
    doc.add_paragraph("за напрямом")
    doc.add_paragraph(f"{title} ” „,,„{title}{title}”")
    doc.add_paragraph("Група № 47п-25")
    doc.add_paragraph("з 23 червня 2025 року до 24 червня 2025 року")

    table = doc.add_table(rows=3, cols=6)
    for idx, value in enumerate(
        ["№п/п", "Назва предмета", "К-сть год.", "23.06", "24.06", "Прізвище, ім'я, по-батькові викладача"]
    ):
        table.cell(0, idx).text = value
    for idx, value in enumerate(["1", "Тема групи", "2", "1п/1год", "2п/1год", "Паращук Світлана Зеновіївна"]):
        table.cell(1, idx).text = value
    for idx, value in enumerate(["", "Загальний обсяг навчального часу:", "2", "", "", ""]):
        table.cell(2, idx).text = value

    doc.save(path)


def _build_schedule_docx_without_index_column(path: Path) -> None:
    doc = DocxDocument()
    doc.add_paragraph("1 пара - 9.30 – 11.05")
    doc.add_paragraph("за напрямом")
    doc.add_paragraph("Штучний інтелект")
    doc.add_paragraph("Група № 162-25")
    doc.add_paragraph("з 21 жовтня 2025 року до 22 жовтня 2025 року")

    table = doc.add_table(rows=3, cols=5)
    for idx, value in enumerate(["Предмет", "Години", "21.10", "22.10", "Викладач"]):
        table.cell(0, idx).text = value
    for idx, value in enumerate(["Кар'єрний розвиток", "2", "1п/1год", "2п/1год", "Паращук Світлана Зеновіївна"]):
        table.cell(1, idx).text = value
    for idx, value in enumerate(["Загальний обсяг навчального часу:", "2", "", "", ""]):
        table.cell(2, idx).text = value

    doc.save(path)


def _build_schedule_docx_as_list_table(path: Path) -> None:
    doc = DocxDocument()
    doc.add_paragraph("1 пара - 9.30 – 11.05")
    doc.add_paragraph("2 пара – 11.10 – 12.45")
    doc.add_paragraph("за напрямом")
    doc.add_paragraph("Штучний інтелект")
    doc.add_paragraph("з 21 жовтня 2025 року до 22 жовтня 2025 року")

    table = doc.add_table(rows=4, cols=5)
    for idx, value in enumerate(["Дата", "Пара", "Предмет", "Години", "Викладач"]):
        table.cell(0, idx).text = value
    for idx, value in enumerate(["21.10.2025", "1", "Кар'єрний розвиток", "1", "Паращук Світлана Зеновіївна"]):
        table.cell(1, idx).text = value
    for idx, value in enumerate(["22.10.2025", "2", "Професійне зростання", "1", ""]):
        table.cell(2, idx).text = value
    for idx, value in enumerate(["", "", "Загальний обсяг навчального часу:", "2", ""]):
        table.cell(3, idx).text = value

    doc.save(path)


def _build_schedule_docx_with_teacher(path: Path, group_code: str, teacher_name: str) -> None:
    doc = DocxDocument()
    doc.add_paragraph("1 пара - 9.30 – 11.05")
    doc.add_paragraph("за напрямом")
    doc.add_paragraph("Тестовий курс")
    doc.add_paragraph(f"Група № {group_code}")
    doc.add_paragraph("з 21 жовтня 2025 року до 21 жовтня 2025 року")

    table = doc.add_table(rows=3, cols=6)
    for idx, value in enumerate(
        ["№п/п", "Назва предмета", "К-сть год.", "21.10", "22.10", "Прізвище, ім'я, по-батькові викладача"]
    ):
        table.cell(0, idx).text = value
    for idx, value in enumerate(["1", "Тема", "2", "1п/2год", "", teacher_name]):
        table.cell(1, idx).text = value
    for idx, value in enumerate(["", "Загальний обсяг навчального часу:", "2", "", "", ""]):
        table.cell(2, idx).text = value

    doc.save(path)


def _build_schedule_docx_with_vsoho_total_row(path: Path) -> None:
    doc = DocxDocument()
    doc.add_paragraph("1 пара - 9.30 – 11.05")
    doc.add_paragraph("2 пара – 11.10 – 12.45")
    doc.add_paragraph("за напрямом")
    doc.add_paragraph("Технології комп'ютерної обробки інформації")
    doc.add_paragraph("Група № 46-26")
    doc.add_paragraph("з 11 березня 2026 року до 12 березня 2026 року")

    table = doc.add_table(rows=4, cols=6)
    for idx, value in enumerate(
        ["№ з/п", "Назва предмета", "К-сть год.", "11.03", "12.03", "Прізвище, ім'я, по-батькові викладача"]
    ):
        table.cell(0, idx).text = value
    for idx, value in enumerate(["1", "Тема", "4", "1,2п/4год", "", "Паращук О.Л."]):
        table.cell(1, idx).text = value
    for idx, value in enumerate(["", "ВСЬОГО ГОДИН:", "4", "4", "0", ""]):
        table.cell(2, idx).text = value
    for idx, value in enumerate(["", "", "", "", "", ""]):
        table.cell(3, idx).text = value

    doc.save(path)


def test_parse_schedule_docx(tmp_path: Path):
    file_path = tmp_path / "schedule.docx"
    _build_schedule_docx(file_path)

    payload_list = parse_schedule_docx(str(file_path))
    assert len(payload_list) == 1
    payload = payload_list[0]
    assert payload["group_code"] == "167-25"
    assert payload["group_total_hours"] == 5
    assert payload["entries"]
    assert any(item["pair_number"] == 1 for item in payload["entries"])


def test_parse_schedule_docx_supports_short_year_and_teacher_carryover(tmp_path: Path):
    file_path = tmp_path / "schedule-short-year.docx"
    _build_schedule_docx_with_short_year_and_merged_teacher(file_path)

    payload_list = parse_schedule_docx(str(file_path))
    assert len(payload_list) == 1
    payload = payload_list[0]
    assert payload["group_code"] == "167-26"
    assert payload["start_date"] == "2025-10-21"
    assert payload["end_date"] == "2025-10-24"
    assert payload["group_total_hours"] == 6
    assert payload["entries"]
    assert all(item["teacher_name"] == "Войтехівська Галина Михайлівна" for item in payload["entries"])


def test_parse_schedule_docx_skips_vsoho_total_row(tmp_path: Path):
    file_path = tmp_path / "schedule-vsoho-total-row.docx"
    _build_schedule_docx_with_vsoho_total_row(file_path)

    payload = parse_schedule_docx(str(file_path))[0]

    assert payload["group_code"] == "46-26"
    assert payload["group_total_hours"] == 4
    assert len(payload["entries"]) == 2
    assert sum(item["academic_hours"] for item in payload["entries"]) == 4
    assert {item["subject_name"] for item in payload["entries"]} == {"Тема"}


def test_import_schedule_docx_with_conflict_detection(db_session, tmp_path: Path):
    file_path = tmp_path / "schedule.docx"
    _build_schedule_docx(file_path)

    summary = import_schedule_docx(db_session, str(file_path), branch_id="main", actor_user_id=1)
    db_session.commit()
    assert summary["created_slots"] == 3

    slots = db_session.query(ScheduleSlot).all()
    assert slots
    assert all(slot.pair_number is not None for slot in slots)
    assert all(slot.academic_hours > 0 for slot in slots)

    summary2 = import_schedule_docx(db_session, str(file_path), branch_id="main", actor_user_id=1)
    db_session.commit()
    assert summary2["created_slots"] == 3
    slots_after = db_session.query(ScheduleSlot).all()
    assert len(slots_after) == 3  # Old slots were deleted and new ones were inserted


def test_import_schedule_docx_skip_existing_keeps_current_schedule(db_session, tmp_path: Path):
    file_path = tmp_path / "schedule-skip-existing.docx"
    _build_schedule_docx(file_path)

    summary = import_schedule_docx(db_session, str(file_path), branch_id="main", actor_user_id=1)
    db_session.commit()
    assert summary["created_slots"] == 3

    summary2 = import_schedule_docx(
        db_session,
        str(file_path),
        branch_id="main",
        actor_user_id=1,
        update_existing_mode="skip_existing",
    )
    db_session.commit()

    assert summary2["created_slots"] == 0
    assert summary2["deleted_slots"] == 0
    assert summary2["skipped_existing_groups"] == 1
    assert db_session.query(ScheduleSlot).count() == 3


def test_import_schedule_docx_missing_only_restores_deleted_slot(db_session, tmp_path: Path):
    file_path = tmp_path / "schedule-missing-only.docx"
    _build_schedule_docx(file_path)

    import_schedule_docx(db_session, str(file_path), branch_id="main", actor_user_id=1)
    db_session.commit()
    first_slot = db_session.query(ScheduleSlot).order_by(ScheduleSlot.starts_at).first()
    db_session.delete(first_slot)
    db_session.commit()

    summary = import_schedule_docx(
        db_session,
        str(file_path),
        branch_id="main",
        actor_user_id=1,
        update_existing_mode="missing_only",
    )
    db_session.commit()

    assert summary["created_slots"] == 1
    assert summary["deleted_slots"] == 0
    assert summary["skipped_existing_slots"] == 2
    assert db_session.query(ScheduleSlot).count() == 3


def test_parse_schedule_docx_with_two_groups(tmp_path: Path):
    file_path = tmp_path / "schedule-two-groups.docx"
    _build_schedule_docx_with_two_groups(file_path)

    payload_list = parse_schedule_docx(str(file_path))
    codes = {payload["group_code"] for payload in payload_list}
    assert codes == {"162-25", "47п-25"}
    assert all(payload["entries"] for payload in payload_list)


def test_import_schedule_docx_with_two_groups_keeps_group_binding(db_session, tmp_path: Path):
    file_path = tmp_path / "schedule-two-groups.docx"
    _build_schedule_docx_with_two_groups(file_path)

    summary = import_schedule_docx(db_session, str(file_path), branch_id="main", actor_user_id=1)
    db_session.commit()

    assert summary["created_slots"] == 4
    slots = db_session.query(ScheduleSlot).all()
    assert len(slots) == 4
    group_ids = {slot.group_id for slot in slots}
    assert len(group_ids) == 2


def test_import_schedule_docx_clips_text_fields_to_db_limits(db_session, tmp_path: Path):
    file_path = tmp_path / "schedule-long-fields.docx"
    doc = DocxDocument()
    doc.add_paragraph("1 пара - 9.30 – 11.05")
    doc.add_paragraph("за напрямом")
    doc.add_paragraph("Дуже довга назва " * 40)
    doc.add_paragraph("Група № 167-25")
    doc.add_paragraph("з 21 жовтня 2025 року до 21 жовтня 2025 року")

    table = doc.add_table(rows=3, cols=6)
    for idx, value in enumerate(["№п/п", "Назва предмета", "К-сть год.", "21.10", "22.10", "Викладач"]):
        table.cell(0, idx).text = value
    for idx, value in enumerate(
        ["1", "Дуже довгий предмет " * 40, "2", "1п/2год", "", "Наддовгепрізвище" * 20]
    ):
        table.cell(1, idx).text = value
    for idx, value in enumerate(["", "Загальний обсяг навчального часу:", "2", "", "", ""]):
        table.cell(2, idx).text = value
    doc.save(file_path)

    import_schedule_docx(db_session, str(file_path), branch_id="main")
    db_session.commit()

    group = db_session.query(Group).filter(Group.code == "167-25").one()
    subject = db_session.query(Subject).filter(Subject.name.like("Дуже довгий предмет%")).one()
    teacher = db_session.query(Teacher).filter(Teacher.branch_id == "main").one()

    assert len(group.name) <= 255
    assert len(subject.name) <= 255
    assert len(teacher.last_name) <= 120
    assert len(teacher.first_name) <= 120


def test_parse_schedule_docx_with_two_groups_in_table_headers(tmp_path: Path):
    file_path = tmp_path / "schedule-two-groups-table-headers.docx"
    _build_schedule_docx_with_two_groups_in_table_headers(file_path)

    payload_list = parse_schedule_docx(str(file_path))
    codes = {payload["group_code"] for payload in payload_list}
    assert codes == {"162-25", "47п-25"}


def test_parse_schedule_docx_with_dot_separator_group_code(tmp_path: Path):
    file_path = tmp_path / "schedule-two-groups-dot-separator.docx"
    _build_schedule_docx_with_dot_separator_group_code(file_path)

    payload_list = parse_schedule_docx(str(file_path))
    codes = {payload["group_code"] for payload in payload_list}
    assert codes == {"162-25", "47п-25"}


def test_parse_schedule_docx_cleans_duplicated_course_title(tmp_path: Path):
    file_path = tmp_path / "schedule-duplicated-title.docx"
    _build_schedule_docx_with_duplicated_course_title(file_path)

    payload = parse_schedule_docx(str(file_path))[0]

    assert payload["group_code"] == "47п-25"
    assert payload["group_name"] == "Штучний інтелект: розвиток кар'єри та професійне зростання"


def test_parse_schedule_docx_without_index_column(tmp_path: Path):
    file_path = tmp_path / "schedule-no-index.docx"
    _build_schedule_docx_without_index_column(file_path)

    payload = parse_schedule_docx(str(file_path))[0]

    assert payload["group_code"] == "162-25"
    assert len(payload["entries"]) == 2
    assert {item["pair_number"] for item in payload["entries"]} == {1, 2}


def test_parse_schedule_docx_as_list_table_uses_filename_group_code(tmp_path: Path):
    file_path = tmp_path / "162-25.docx"
    _build_schedule_docx_as_list_table(file_path)

    payload = parse_schedule_docx(str(file_path))[0]

    assert payload["group_code"] == "162-25"
    assert len(payload["entries"]) == 2
    assert {item["subject_name"] for item in payload["entries"]} == {"Кар'єрний розвиток", "Професійне зростання"}


def test_import_schedule_docx_updates_existing_duplicated_group_name(db_session, tmp_path: Path):
    file_path = tmp_path / "schedule-duplicated-title.docx"
    _build_schedule_docx_with_duplicated_course_title(file_path)
    bad_name = "Штучний інтелект: розвиток кар'єри та професійне зростання " * 3
    db_session.add(Group(branch_id="main", code="47п-25", name=bad_name))
    db_session.commit()

    import_schedule_docx(db_session, str(file_path), branch_id="main")
    db_session.commit()

    group = db_session.query(Group).filter(Group.code == "47п-25").one()
    assert group.name == "Штучний інтелект: розвиток кар'єри та професійне зростання"


def test_import_schedule_docx_matches_teacher_initials_to_full_name(db_session, tmp_path: Path):
    file_path = tmp_path / "schedule-teacher-initials.docx"
    _build_schedule_docx_with_teacher(file_path, "46-26", "Коваль І. П.")
    existing_teacher = Teacher(
        branch_id="main",
        last_name="Коваль",
        first_name="Іван Петрович",
        hourly_rate=0,
        is_active=True,
    )
    db_session.add(existing_teacher)
    db_session.commit()

    import_schedule_docx(db_session, str(file_path), branch_id="main")
    db_session.commit()

    teachers = db_session.query(Teacher).filter(Teacher.branch_id == "main", Teacher.last_name == "Коваль").all()
    slot = db_session.query(ScheduleSlot).one()
    assert len(teachers) == 1
    assert slot.teacher_id == existing_teacher.id


def test_import_schedule_docx_keeps_same_surname_different_initials_separate(db_session, tmp_path: Path):
    file_path = tmp_path / "schedule-teacher-different-initials.docx"
    _build_schedule_docx_with_teacher(file_path, "46-26", "Коваль І. М.")
    existing_teacher = Teacher(
        branch_id="main",
        last_name="Коваль",
        first_name="Іван Петрович",
        hourly_rate=0,
        is_active=True,
    )
    db_session.add(existing_teacher)
    db_session.commit()

    import_schedule_docx(db_session, str(file_path), branch_id="main")
    db_session.commit()

    teachers = db_session.query(Teacher).filter(Teacher.branch_id == "main", Teacher.last_name == "Коваль").all()
    slot = db_session.query(ScheduleSlot).one()
    assert len(teachers) == 2
    assert slot.teacher_id != existing_teacher.id
    assert any(teacher.first_name == "І. М." for teacher in teachers)
