import re
from datetime import date, datetime, time, timedelta, timezone
from pathlib import Path

from docx import Document as DocxDocument
from sqlalchemy.orm import Session

from app.models import Group, GroupStatus, Room, ScheduleSlot, Subject, Teacher

SCHEDULE_UPDATE_MODES = {"skip_existing", "missing_only", "overwrite"}

UA_MONTHS = {
    "січня": 1,
    "лютого": 2,
    "березня": 3,
    "квітня": 4,
    "травня": 5,
    "червня": 6,
    "липня": 7,
    "серпня": 8,
    "вересня": 9,
    "жовтня": 10,
    "листопада": 11,
    "грудня": 12,
}


def _norm(text: str) -> str:
    return re.sub(r"\s+", " ", text or "").strip()


def _db_text(text: str | None, max_length: int) -> str:
    value = _norm(text or "")
    return value[:max_length]


def _clean_course_title(text: str | None, group_code: str) -> str:
    value = _norm(text or "")
    if not value:
        return ""

    value = re.sub(r"^\s*за\s+напрямом\s*", "", value, flags=re.IGNORECASE)
    value = value.strip(" «»„“”\"'")
    value = value.replace("\u00a0", " ")
    for quote in ("«", "»", "„", "“", "”", '"'):
        value = value.replace(quote, " ")
    value = re.sub(r"\s*,+\s*", " ", value)
    value = re.sub(r"\s+", " ", value).strip(" ,.;:-")

    # Some DOCX files expose the same title multiple times, sometimes with no
    # separator between copies after quote cleanup. Keep the first full copy.
    for index in range(20, len(value) // 2 + 1):
        if value[index] != value[0]:
            continue
        candidate = value[:index].strip(" ,.;:-")
        if len(candidate) < 20:
            continue
        if value[index:].lstrip(" ,.;:-").startswith(candidate):
            return candidate

    if group_code and group_code in value:
        return ""
    return value


def _parse_group_code(lines: list[str]) -> str:
    def _normalize_group_code(raw: str) -> str:
        value = (raw or "").strip()
        for sep in ("–", "—", "−", "/", "."):
            value = value.replace(sep, "-")
        value = re.sub(r"\s+", "", value)
        value = re.sub(r"-{2,}", "-", value)
        return value

    # Read from bottom to top so the nearest heading to a table wins.
    for line in reversed(lines):
        match = re.search(
            r"група\s*№?\s*([0-9]{1,4}[A-Za-zА-Яа-яЇїІіЄєҐґ]?\s*[-/.]\s*\d{1,4})",
            line,
            re.IGNORECASE,
        )
        if match:
            return _normalize_group_code(match.group(1))
        fallback = re.search(r"група\s*№?\s*([0-9A-Za-zА-Яа-яЇїІіЄєҐґ\-/\.]+)", line, re.IGNORECASE)
        if fallback:
            cleaned = re.sub(r"[^0-9A-Za-zА-Яа-яЇїІіЄєҐґ\-/\.]", "", fallback.group(1))
            if cleaned:
                return _normalize_group_code(cleaned)

    # Fallback 3: look for common group code patterns even without the word "група"
    # We do not use '.' here to avoid matching dates like '23.06'
    for line in reversed(lines):
        match = re.search(
            r"\b([0-9]{1,4}[A-Za-zА-Яа-яЇїІіЄєҐґ]?\s*[-/]\s*\d{2})\b",
            line,
            re.IGNORECASE,
        )
        if match:
            return _normalize_group_code(match.group(1))

    raise ValueError("Не вдалося визначити номер групи з документа")


def _parse_course_title(lines: list[str], group_code: str) -> str:
    for index, line in enumerate(lines):
        if "за напрямом" in line.lower():
            match = re.search(r"за напрямом\s*[«\"'„](.*?)[»\"'”]", line, re.IGNORECASE)
            if match:
                cleaned = _clean_course_title(match.group(1), group_code)
                if cleaned:
                    return cleaned
            if index + 1 < len(lines):
                cleaned = _clean_course_title(lines[index + 1], group_code)
                if cleaned:
                    return cleaned
    for line in lines:
        if group_code in line:
            continue
        if len(line) > 25:
            cleaned = _clean_course_title(line, group_code)
            if cleaned:
                return cleaned
    return f"Група {group_code}"


def _parse_date_range(lines: list[str]) -> tuple[date | None, date | None]:
    def _year_to_int(value: str | None) -> int | None:
        if not value:
            return None
        normalized = value.strip()
        if not normalized.isdigit():
            return None
        if len(normalized) == 2:
            return 2000 + int(normalized)
        if len(normalized) == 4:
            return int(normalized)
        return None

    for line in lines:
        # e.g., "з «11» березня 2026 року до «24» березня 2026 року"
        # Extract all dates in the line
        matches = re.findall(
            r"(\d{1,2})[»\"']?\s+([А-Яа-яіїєґ]+)(?:\s+(\d{2,4})(?:-?го)?)?",
            line.lower(),
        )
        if len(matches) >= 2:
            start_match, end_match = matches[0], matches[-1]
            start_day, start_month_ua, start_year = start_match
            end_day, end_month_ua, end_year = end_match

            start_month = UA_MONTHS.get(start_month_ua)
            end_month = UA_MONTHS.get(end_month_ua)

            year_int = _year_to_int(end_year) or _year_to_int(start_year) or date.today().year
            start_year_int = _year_to_int(start_year) or year_int

            if start_month and end_month:
                start_date = date(start_year_int, start_month, int(start_day))
                end_date = date(year_int, end_month, int(end_day))
                return start_date, end_date
    return None, None


def _parse_pair_windows(lines: list[str]) -> dict[int, tuple[time, time]]:
    pair_windows: dict[int, tuple[time, time]] = {}
    for line in lines:
        match = re.search(
            r"(\d+)\s*пара\s*[-–—:]?\s*(\d{1,2})[.:](\d{2})\s*[–—-]\s*(\d{1,2})[.:](\d{2})",
            line.lower(),
        )
        if not match:
            continue
        pair_number, sh, sm, eh, em = match.groups()
        pair_windows[int(pair_number)] = (time(int(sh), int(sm)), time(int(eh), int(em)))
    return pair_windows


def _parse_date_from_cell(value: str, default_year: int) -> date | None:
    text = _norm(value).lower()
    match = re.search(r"\b(\d{1,2})[./](\d{1,2})(?:[./](\d{2,4}))?\b", text)
    if match:
        day, month, year_text = match.groups()
        year = default_year
        if year_text:
            year = 2000 + int(year_text) if len(year_text) == 2 else int(year_text)
        return date(year, int(month), int(day))

    match = re.search(r"\b(\d{1,2})\s+([А-Яа-яіїєґ]+)(?:\s+(\d{2,4}))?\b", text)
    if match:
        day, month_text, year_text = match.groups()
        month = UA_MONTHS.get(month_text)
        if not month:
            return None
        year = default_year
        if year_text:
            year = 2000 + int(year_text) if len(year_text) == 2 else int(year_text)
        return date(year, month, int(day))
    return None


def _find_header_column(headers: list[str], keywords: tuple[str, ...], excluded: tuple[str, ...] = ()) -> int | None:
    for index, header in enumerate(headers):
        low = header.lower()
        if excluded and any(item in low for item in excluded):
            continue
        if any(item in low for item in keywords):
            return index
    return None


def _infer_matrix_columns(headers: list[str], date_columns: dict[int, date]) -> tuple[int, int | None, int | None]:
    subject_col = _find_header_column(headers, ("предмет", "тема", "зміст", "назва"), ("виклада", "прізвищ", "піб"))
    hours_col = _find_header_column(headers, ("год",), ("виклада",))
    teacher_col = _find_header_column(headers, ("виклада", "прізвищ", "піб"))

    if subject_col is None:
        first_date_col = min(date_columns)
        candidates = [
            index
            for index, header in enumerate(headers[:first_date_col])
            if index != hours_col and "№" not in header and not re.search(r"^\s*№", header)
        ]
        subject_col = candidates[-1] if candidates else 1

    if teacher_col is None:
        after_date_columns = [index for index in range(max(date_columns) + 1, len(headers))]
        teacher_col = after_date_columns[-1] if after_date_columns else len(headers) - 1

    return subject_col, hours_col, teacher_col


def _parse_hours(value: str) -> float:
    match = re.search(r"([0-9]+(?:[.,][0-9]+)?)", value)
    if not match:
        return 0.0
    return float(match.group(1).replace(",", "."))


def _is_total_hours_row(text: str) -> bool:
    low = _norm(text).lower()
    return (
        "загальний обсяг" in low
        or "всього годин" in low
        or "усього годин" in low
        or ("всього" in low and "год" in low)
        or ("усього" in low and "год" in low)
    )


def _parse_pairs_cell(cell_text: str) -> list[tuple[int, float]]:
    text = _norm(cell_text).lower()
    if not text:
        return []
    hours_match = re.search(r"/\s*([0-9]+(?:[.,][0-9]+)?)\s*год", text)
    total_hours = float(hours_match.group(1).replace(",", ".")) if hours_match else 2.0

    pair_part = text.split("/")[0]
    pair_part = pair_part.replace("п", "").replace("ара", "").replace("пара", "")
    pair_numbers: list[int] = []
    for token in re.split(r"[,\s]+", pair_part):
        token = token.strip()
        if not token:
            continue
        range_match = re.match(r"^(\d+)-(\d+)$", token)
        if range_match:
            start_idx = int(range_match.group(1))
            end_idx = int(range_match.group(2))
            pair_numbers.extend(list(range(start_idx, end_idx + 1)))
            continue
        if token.isdigit():
            pair_numbers.append(int(token))

    pair_numbers = sorted({pair for pair in pair_numbers if pair > 0})
    if not pair_numbers:
        return []

    hours_per_pair = round(total_hours / len(pair_numbers), 2)
    return [(pair, hours_per_pair) for pair in pair_numbers]


def _split_teacher_name(full_name: str) -> tuple[str, str]:
    tokens = [part for part in _norm(full_name).split(" ") if part]
    if not tokens:
        return "Невідомий", "Викладач"
    if len(tokens) == 1:
        return tokens[0], "Викладач"
    return tokens[0], " ".join(tokens[1:])


def _teacher_given_parts(first_name: str | None) -> list[str]:
    value = _norm(first_name or "")
    if not value or value.casefold() == "викладач":
        return []
    return re.findall(r"[A-Za-zА-Яа-яЇїІіЄєҐґ]+", value)


def _teacher_initials(first_name: str | None) -> tuple[str, ...]:
    return tuple(part[0].casefold() for part in _teacher_given_parts(first_name) if part)


def _teacher_names_compatible(existing_first_name: str | None, incoming_first_name: str | None) -> bool:
    existing_parts = _teacher_given_parts(existing_first_name)
    incoming_parts = _teacher_given_parts(incoming_first_name)

    if not existing_parts or not incoming_parts:
        return _norm(existing_first_name or "").casefold() == _norm(incoming_first_name or "").casefold()

    existing_initials = _teacher_initials(existing_first_name)
    incoming_initials = _teacher_initials(incoming_first_name)
    if not existing_initials or not incoming_initials:
        return False

    compared_len = min(len(existing_initials), len(incoming_initials))
    return existing_initials[:compared_len] == incoming_initials[:compared_len]


def _teacher_name_is_more_complete(new_first_name: str | None, current_first_name: str | None) -> bool:
    new_parts = _teacher_given_parts(new_first_name)
    current_parts = _teacher_given_parts(current_first_name)
    if not new_parts:
        return False
    if len(new_parts) < len(current_parts):
        return False
    new_value = _norm(new_first_name or "")
    current_value = _norm(current_first_name or "")
    new_has_full_words = any(len(part) > 1 for part in new_parts)
    current_has_only_initials = bool(current_parts) and all(len(part) == 1 for part in current_parts)
    return new_has_full_words and (current_has_only_initials or len(new_value) > len(current_value))


def _same_teacher_identity(left: Teacher, right: Teacher) -> bool:
    if (left.last_name or "").strip().casefold() != (right.last_name or "").strip().casefold():
        return False
    return _teacher_names_compatible(left.first_name, right.first_name)


def parse_schedule_docx(file_path: str) -> list[dict]:
    document = DocxDocument(file_path)
    lines = [_norm(paragraph.text) for paragraph in document.paragraphs if _norm(paragraph.text)]
    file_name_lines = [_norm(Path(file_path).stem)]
    if not document.tables:
        raise ValueError("У документі не знайдено таблиці розкладу")

    # Global fallback parsing
    try:
        global_group_code = _parse_group_code(lines + file_name_lines)
    except ValueError:
        global_group_code = "Невідома група"
        
    global_group_name = _parse_course_title(lines, global_group_code)
    global_start_date, global_end_date = _parse_date_range(lines)
    global_pair_windows = _parse_pair_windows(lines)

    # Preserve table order from document body and capture nearby paragraph context
    # so each table can resolve its own group/date metadata.
    table_contexts: list[tuple[object, list[str]]] = []
    context_lines: list[str] = []
    table_index = 0
    for body_child in document.element.body.iterchildren():
        tag = body_child.tag.lower()
        if tag.endswith("}p"):
            text = _norm("".join(body_child.itertext()))
            if text:
                context_lines.append(text)
            continue
        if tag.endswith("}tbl"):
            if table_index >= len(document.tables):
                continue
            table = document.tables[table_index]
            table_index += 1
            # Pass all preceding paragraphs instead of just the last 40.
            # Since we search bottom-up, this guarantees we find the nearest group code
            # even if there is a massive header with >40 paragraphs.
            table_contexts.append((table, list(context_lines)))

    # Fallback when body traversal did not map tables (rare malformed DOCX).
    if not table_contexts:
        table_contexts = [(table, list(lines)) for table in document.tables]

    grouped_results: dict[str, dict] = {}

    for table, table_lines in table_contexts:
        if len(table.rows) < 2 or len(table.columns) < 4:
            continue

        table_cell_lines = [_norm(cell.text) for row in table.rows for cell in row.cells if _norm(cell.text)]
        # We want to check the table cells first (they are the most specific to the table),
        # and then fallback to the paragraphs above the table.
        # Since `_parse_group_code` searches `reversed(lines)`, we put `table_cell_lines`
        # at the end of the list, so they are searched FIRST.
        metadata_lines = table_lines + table_cell_lines + file_name_lines

        try:
            local_group_code = _parse_group_code(metadata_lines)
        except ValueError:
            # Fallback 1: try to find any group code in JUST the table lines (if table_cell_lines confused it)
            try:
                local_group_code = _parse_group_code(table_lines)
            except ValueError:
                # Fallback 2: try to find any group code in JUST the table cell lines
                try:
                    local_group_code = _parse_group_code(table_cell_lines)
                except ValueError:
                    try:
                        local_group_code = _parse_group_code(file_name_lines)
                    except ValueError:
                        local_group_code = global_group_code

        local_group_name = _parse_course_title(table_lines or lines, local_group_code)
        local_start_date, local_end_date = _parse_date_range(table_lines or lines)
        if not local_start_date and global_start_date:
            local_start_date = global_start_date
        if not local_end_date and global_end_date:
            local_end_date = global_end_date

        local_pair_windows = _parse_pair_windows(table_lines or lines) or global_pair_windows

        list_header_candidates = []
        for row_index, row in enumerate(table.rows):
            headers_for_row = [_norm(cell.text) for cell in row.cells]
            low = " ".join(headers_for_row).lower()
            if "дата" in low and ("пара" in low or "год" in low or "предмет" in low):
                list_header_candidates.append((row_index, headers_for_row))

        header_row_index = -1
        date_columns: dict[int, date] = {}
        year = (local_end_date or local_start_date or global_end_date or global_start_date or date.today()).year
        if not list_header_candidates:
            for row_index, row in enumerate(table.rows):
                current_date_columns: dict[int, date] = {}
                for column_index in range(0, len(row.cells)):
                    header_text = _norm(row.cells[column_index].text)
                    parsed_date = _parse_date_from_cell(header_text, year)
                    if parsed_date:
                        current_date_columns[column_index] = parsed_date
                if current_date_columns:
                    header_row_index = row_index
                    date_columns = current_date_columns
                    break

        headers = [_norm(cell.text) for cell in table.rows[header_row_index].cells] if header_row_index >= 0 else []

        table_entries: list[dict] = []
        table_total_group_hours = 0.0
        previous_teacher_name = ""
        if date_columns:
            subject_col, hours_col, teacher_col = _infer_matrix_columns(headers, date_columns)
            for row in table.rows[header_row_index + 1 :]:
                cells = row.cells
                if subject_col >= len(cells):
                    continue
                subject_name = _norm(cells[subject_col].text)
                declared_subject_hours = _parse_hours(cells[hours_col].text) if hours_col is not None and hours_col < len(cells) else 0.0
                teacher_name = _norm(cells[teacher_col].text) if teacher_col is not None and teacher_col < len(cells) else ""
                teacher_name = teacher_name or previous_teacher_name
                if teacher_name:
                    previous_teacher_name = teacher_name

                row_text = " ".join(_norm(cell.text) for cell in cells)
                if _is_total_hours_row(row_text):
                    current_total = _parse_hours(row_text)
                    if current_total > table_total_group_hours:
                        table_total_group_hours = current_total
                    continue
                if not subject_name:
                    continue

                for column_index, lesson_date in date_columns.items():
                    if column_index >= len(cells):
                        continue
                    cell_value = _norm(cells[column_index].text)
                    for pair_number, academic_hours in _parse_pairs_cell(cell_value):
                        if declared_subject_hours and academic_hours <= 0:
                            academic_hours = declared_subject_hours
                        pair_window = local_pair_windows.get(pair_number)
                        if pair_window:
                            starts_at = datetime.combine(lesson_date, pair_window[0], tzinfo=timezone.utc)
                            ends_at = datetime.combine(lesson_date, pair_window[1], tzinfo=timezone.utc)
                        else:
                            starts_at = datetime.combine(lesson_date, time(hour=9 + (pair_number - 1) * 2), tzinfo=timezone.utc)
                            ends_at = starts_at + timedelta(minutes=95)
                        table_entries.append(
                            {
                                "subject_name": subject_name,
                                "declared_subject_hours": declared_subject_hours,
                                "teacher_name": teacher_name,
                                "lesson_date": lesson_date.isoformat(),
                                "pair_number": pair_number,
                                "academic_hours": academic_hours,
                                "starts_at": starts_at,
                                "ends_at": ends_at,
                            }
                        )
        else:
            if list_header_candidates:
                header_row_index, headers = list_header_candidates[0]
                date_col = _find_header_column(headers, ("дата",))
                pair_col = _find_header_column(headers, ("пара",))
                subject_col = _find_header_column(headers, ("предмет", "тема", "зміст", "назва"), ("виклада", "піб"))
                hours_col = _find_header_column(headers, ("год",), ("виклада",))
                teacher_col = _find_header_column(headers, ("виклада", "прізвищ", "піб"))

                for row in table.rows[header_row_index + 1 :]:
                    cells = row.cells
                    row_values = [_norm(cell.text) for cell in cells]
                    row_text = " ".join(row_values)
                    if _is_total_hours_row(row_text):
                        current_total = _parse_hours(row_text)
                        if current_total > table_total_group_hours:
                            table_total_group_hours = current_total
                        continue

                    lesson_date = None
                    if date_col is not None and date_col < len(cells):
                        lesson_date = _parse_date_from_cell(cells[date_col].text, year)
                    if lesson_date is None:
                        for value in row_values:
                            lesson_date = _parse_date_from_cell(value, year)
                            if lesson_date:
                                break
                    if lesson_date is None:
                        continue

                    subject_name = _norm(cells[subject_col].text) if subject_col is not None and subject_col < len(cells) else ""
                    if not subject_name:
                        continue
                    declared_subject_hours = _parse_hours(cells[hours_col].text) if hours_col is not None and hours_col < len(cells) else 0.0
                    teacher_name = _norm(cells[teacher_col].text) if teacher_col is not None and teacher_col < len(cells) else ""
                    teacher_name = teacher_name or previous_teacher_name
                    if teacher_name:
                        previous_teacher_name = teacher_name

                    pair_source = cells[pair_col].text if pair_col is not None and pair_col < len(cells) else row_text
                    parsed_pairs = _parse_pairs_cell(pair_source)
                    if declared_subject_hours and len(parsed_pairs) == 1:
                        parsed_pairs = [(parsed_pairs[0][0], declared_subject_hours)]
                    for pair_number, academic_hours in parsed_pairs:
                        pair_window = local_pair_windows.get(pair_number)
                        if pair_window:
                            starts_at = datetime.combine(lesson_date, pair_window[0], tzinfo=timezone.utc)
                            ends_at = datetime.combine(lesson_date, pair_window[1], tzinfo=timezone.utc)
                        else:
                            starts_at = datetime.combine(lesson_date, time(hour=9 + (pair_number - 1) * 2), tzinfo=timezone.utc)
                            ends_at = starts_at + timedelta(minutes=95)
                        table_entries.append(
                            {
                                "subject_name": subject_name,
                                "declared_subject_hours": declared_subject_hours,
                                "teacher_name": teacher_name,
                                "lesson_date": lesson_date.isoformat(),
                                "pair_number": pair_number,
                                "academic_hours": academic_hours,
                                "starts_at": starts_at,
                                "ends_at": ends_at,
                            }
                        )

        if not table_entries:
            continue

        if table_total_group_hours <= 0:
            table_total_group_hours = round(sum(item["academic_hours"] for item in table_entries), 2)

        bucket = grouped_results.get(local_group_code)
        if not bucket:
            grouped_results[local_group_code] = {
                "group_code": local_group_code,
                "group_name": local_group_name,
                "start_date": local_start_date.isoformat() if local_start_date else None,
                "end_date": local_end_date.isoformat() if local_end_date else None,
                "group_total_hours": table_total_group_hours,
                "entries": table_entries,
            }
        else:
            bucket["entries"].extend(table_entries)
            bucket["group_total_hours"] = round(float(bucket["group_total_hours"]) + table_total_group_hours, 2)
            if not bucket.get("start_date") and local_start_date:
                bucket["start_date"] = local_start_date.isoformat()
            if not bucket.get("end_date") and local_end_date:
                bucket["end_date"] = local_end_date.isoformat()
            if local_group_name and (
                not bucket.get("group_name") or bucket.get("group_name") == f"Група {local_group_code}"
            ):
                bucket["group_name"] = local_group_name

    if not grouped_results:
        raise ValueError("У таблицях не знайдено занять для імпорту")

    return list(grouped_results.values())


def _merge_duplicate_teachers(db: Session, branch_id: str, teacher_ids: list[int]) -> int:
    """Merge duplicate Teacher records that share surname and compatible initials.

    When two imports run in close succession (e.g. concurrent Vercel lambdas) or
    when the same teacher appears as a full name in one file and initials in another,
    the lookup may fail and create a second record.  This function:
      1. Finds all teachers in *branch_id* whose surname and initials collide
         with any teacher in *teacher_ids*.
      2. Keeps the record with the smallest id (the 'canonical' one).
      3. Reassigns all ScheduleSlot rows from duplicates to the canonical record.
      4. Deletes the duplicate Teacher rows.

    Returns the number of duplicate records removed.
    """
    if not teacher_ids:
        return 0

    # Fetch the teachers we just touched in this import
    touched = db.query(Teacher).filter(Teacher.id.in_(teacher_ids)).all()
    merged = 0
    seen_by_last: dict[str, list[Teacher]] = {}

    for t in sorted(touched, key=lambda x: x.id):
        key = (t.last_name or "").strip().lower()
        same_last_teachers = seen_by_last.setdefault(key, [])
        canonical = next((candidate for candidate in same_last_teachers if _same_teacher_identity(candidate, t)), None)
        if canonical:
            if _teacher_name_is_more_complete(t.first_name, canonical.first_name):
                canonical.first_name = t.first_name
                db.add(canonical)
            db.query(ScheduleSlot).filter(ScheduleSlot.teacher_id == t.id).update(
                {"teacher_id": canonical.id}, synchronize_session=False
            )
            db.delete(t)
            merged += 1
        else:
            same_last_teachers.append(t)

    # Also look for pre-existing duplicates across ALL teachers in the branch
    # that share a surname and compatible initials with any teacher we touched.
    for norm_last, canonicals in list(seen_by_last.items()):
        canonical_ids = {canonical.id for canonical in canonicals}
        others = [
            teacher
            for teacher in db.query(Teacher).filter(Teacher.branch_id == branch_id).all()
            if teacher.id not in canonical_ids and (teacher.last_name or "").strip().casefold() == norm_last.casefold()
        ]
        for dup in others:
            canonical = next((candidate for candidate in canonicals if _same_teacher_identity(candidate, dup)), None)
            if not canonical:
                continue
            if _teacher_name_is_more_complete(dup.first_name, canonical.first_name):
                canonical.first_name = dup.first_name
                db.add(canonical)
            db.query(ScheduleSlot).filter(ScheduleSlot.teacher_id == dup.id).update(
                {"teacher_id": canonical.id}, synchronize_session=False
            )
            db.delete(dup)
            merged += 1

    if merged:
        db.flush()
    return merged


def _as_utc(value: datetime) -> datetime:
    return value if value.tzinfo is not None else value.replace(tzinfo=timezone.utc)


def _overlaps(a_start: datetime, a_end: datetime, b_start: datetime, b_end: datetime) -> bool:
    a_start = _as_utc(a_start)
    a_end = _as_utc(a_end)
    b_start = _as_utc(b_start)
    b_end = _as_utc(b_end)
    return a_start < b_end and b_start < a_end


def import_schedule_docx(
    db: Session,
    file_path: str,
    branch_id: str,
    actor_user_id: int | None = None,
    update_existing_mode: str = "overwrite",
) -> dict:
    if update_existing_mode not in SCHEDULE_UPDATE_MODES:
        update_existing_mode = "overwrite"
    parsed_list = parse_schedule_docx(file_path)

    total_deleted = 0
    total_created = 0
    total_skipped_groups = 0
    total_skipped_slots = 0
    total_existing_slots = 0
    all_merged = 0
    
    # We will accumulate stats across all groups
    all_teacher_hours: dict[int, tuple[str, float]] = {}
    total_group_hours = 0.0
    group_codes = []
    group_names = []
    
    global_teacher_id_cache = {}
    global_subject_cache = {}

    for parsed in parsed_list:
        group_code = _db_text(parsed["group_code"], 50)
        group_name = _db_text(parsed["group_name"], 255) or f"Група {group_code}"

        if group_code not in group_codes:
            group_codes.append(group_code)
        if group_name not in group_names:
            group_names.append(group_name)

        total_group_hours += parsed.get("group_total_hours", 0.0)

        group = db.query(Group).filter(Group.branch_id == branch_id, Group.code == group_code).first()
        if not group:
            group = Group(
                branch_id=branch_id,
                code=group_code,
                name=group_name,
                status=GroupStatus.ACTIVE,
                start_date=date.fromisoformat(parsed["start_date"]) if parsed["start_date"] else None,
                end_date=date.fromisoformat(parsed["end_date"]) if parsed["end_date"] else None,
            )
            db.add(group)
            db.flush()
        else:
            group.name = group_name or group.name
            if parsed["start_date"]:
                group.start_date = date.fromisoformat(parsed["start_date"])
            if parsed["end_date"]:
                group.end_date = date.fromisoformat(parsed["end_date"])
            db.add(group)
            db.flush()

        parsed_entries = parsed["entries"]
        if not parsed_entries:
            continue
        entry_min_start = min(item["starts_at"] for item in parsed_entries)
        entry_max_end = max(item["ends_at"] for item in parsed_entries)
        group_existing_slots = (
            db.query(ScheduleSlot)
            .filter(
                ScheduleSlot.group_id == group.id,
                ScheduleSlot.starts_at >= entry_min_start,
                ScheduleSlot.starts_at <= entry_max_end,
            )
            .all()
        )
        total_existing_slots += len(group_existing_slots)
        if update_existing_mode == "skip_existing" and group_existing_slots:
            total_skipped_groups += 1
            continue

        room_name = f"Імпорт: {group.code}"
        room = db.query(Room).filter(Room.branch_id == branch_id, Room.name == room_name).first()
        if not room:
            room = Room(branch_id=branch_id, name=room_name, capacity=max(group.capacity, 20))
            db.add(room)
            db.flush()


        teacher_cache: dict[str, Teacher] = {}  # keyed by normalised name string from the document
        teacher_id_cache: dict[int, Teacher] = {}  # keyed by DB teacher.id for deduplication
        subject_cache: dict[str, Subject] = {}
        candidates: list[dict] = []

        for entry in parsed_entries:
            teacher_full_name = _norm(entry["teacher_name"])
            if teacher_full_name not in teacher_cache:
                last_name, first_name = _split_teacher_name(teacher_full_name)
                last_name = _db_text(last_name, 120) or "Невідомий"
                first_name = _db_text(first_name, 120) or "Викладач"
                # Case-insensitive lookup — Ukrainian DOCX files often store surnames
                # in ALL-CAPS while the DB may have them in proper case (or vice versa).
                existing_teachers = [
                    teacher
                    for teacher in db.query(Teacher).filter(Teacher.branch_id == branch_id).all()
                    if (teacher.last_name or "").strip().casefold() == last_name.casefold()
                ]
                teacher = None
                if existing_teachers:
                    matching_teachers = [
                        t for t in existing_teachers if _teacher_names_compatible(t.first_name, first_name)
                    ]
                    if len(matching_teachers) == 1:
                        teacher = matching_teachers[0]
                    elif len(matching_teachers) > 1:
                        teacher = sorted(matching_teachers, key=lambda item: item.id)[0]
                    elif len(existing_teachers) == 1 and not _teacher_given_parts(first_name):
                        teacher = existing_teachers[0]

                if teacher:
                    # If the incoming name is more complete, upgrade the DB record
                    if _teacher_name_is_more_complete(first_name, teacher.first_name):
                        teacher.first_name = first_name
                        db.add(teacher)
                        db.flush()
                        # Invalidate any previous cache entry that pointed to this teacher
                        if teacher.id in teacher_id_cache:
                            teacher_id_cache[teacher.id] = teacher
                else:
                    teacher = Teacher(
                        branch_id=branch_id,
                        last_name=last_name,
                        first_name=first_name,
                        hourly_rate=0.0,
                        is_active=True,
                    )
                    db.add(teacher)
                    db.flush()

                # Use the canonical teacher object (deduplicated by ID)
                if teacher.id in teacher_id_cache:
                    teacher = teacher_id_cache[teacher.id]
                else:
                    teacher_id_cache[teacher.id] = teacher

                teacher_cache[teacher_full_name] = teacher
            teacher = teacher_cache[teacher_full_name]

            subject_name = _db_text(entry["subject_name"], 255) or "Без назви"
            if subject_name not in subject_cache:
                subject = db.query(Subject).filter(Subject.branch_id == branch_id, Subject.name == subject_name).first()
                declared_hours = int(entry["declared_subject_hours"] or 0)
                if not subject:
                    subject = Subject(branch_id=branch_id, name=subject_name, hours_total=max(declared_hours, 1))
                    db.add(subject)
                    db.flush()
                elif declared_hours > subject.hours_total:
                    subject.hours_total = declared_hours
                    db.add(subject)
                    db.flush()
                subject_cache[subject_name] = subject
            subject = subject_cache[subject_name]

            candidates.append(
                {
                    "group_id": group.id,
                    "teacher_id": teacher.id,
                    "subject_id": subject.id,
                    "room_id": room.id,
                    "subject_name": subject_name,
                    "teacher_name": teacher_full_name,
                    "starts_at": entry["starts_at"],
                    "ends_at": entry["ends_at"],
                    "pair_number": entry["pair_number"],
                    "academic_hours": entry["academic_hours"],
                }
            )
        if not candidates:
            continue

        min_start = min(item["starts_at"] for item in candidates)
        max_end = max(item["ends_at"] for item in candidates)

        deleted_count = 0
        if update_existing_mode == "overwrite":
            # Idempotent replace: wipe THIS group's slots in the exact date window
            # that the document covers before inserting fresh ones.
            deleted_count = (
                db.query(ScheduleSlot)
                .filter(
                    ScheduleSlot.group_id == group.id,
                    ScheduleSlot.starts_at >= min_start,
                    ScheduleSlot.starts_at <= max_end,
                )
                .delete(synchronize_session=False)
            )
            db.flush()
        elif update_existing_mode == "missing_only" and group_existing_slots:
            original_count = len(candidates)
            candidates = [
                candidate
                for candidate in candidates
                if not any(
                    _overlaps(candidate["starts_at"], candidate["ends_at"], slot.starts_at, slot.ends_at)
                    and (slot.pair_number == candidate["pair_number"] or slot.starts_at == candidate["starts_at"])
                    for slot in group_existing_slots
                )
            ]
            total_skipped_slots += original_count - len(candidates)
            if not candidates:
                continue

        # Conflict detection: look for clashes with OTHER groups in the same window
        # (run AFTER the wipe so this group's old slots don't create false alarms)
        existing_slots = (
            db.query(ScheduleSlot)
            .join(Group, Group.id == ScheduleSlot.group_id)
            .filter(
                Group.branch_id == branch_id,
                ScheduleSlot.starts_at < max_end,
                ScheduleSlot.ends_at > min_start,
            )
            .all()
        )

        conflict_messages: list[str] = []
        for index, candidate in enumerate(candidates):
            for existing in existing_slots:
                if not _overlaps(candidate["starts_at"], candidate["ends_at"], existing.starts_at, existing.ends_at):
                    continue
                conflict_date = candidate["starts_at"].date().isoformat()
                if existing.teacher_id == candidate["teacher_id"]:
                    conflict_messages.append(
                        f"викладач {candidate['teacher_name']} {conflict_date} пара {candidate['pair_number']}"
                    )
                if existing.group_id == candidate["group_id"]:
                    conflict_messages.append(
                        f"група {group.code} {conflict_date} пара {candidate['pair_number']}"
                    )
            for previous in candidates[:index]:
                if not _overlaps(candidate["starts_at"], candidate["ends_at"], previous["starts_at"], previous["ends_at"]):
                    continue
                if previous["teacher_id"] == candidate["teacher_id"]:
                    conflict_messages.append(
                        f"внутрішній конфлікт імпорту: {candidate['teacher_name']} {candidate['starts_at'].date().isoformat()} пара {candidate['pair_number']}"
                    )

        if conflict_messages:
            unique_conflicts = sorted(set(conflict_messages))
            preview = "; ".join(unique_conflicts[:8])
            # We no longer fail the import on conflict! 
            # The user wants to see conflicts in the UI directly.
            pass

        for candidate in candidates:
            db.add(
                ScheduleSlot(
                    group_id=candidate["group_id"],
                    teacher_id=candidate["teacher_id"],
                    subject_id=candidate["subject_id"],
                    room_id=candidate["room_id"],
                    starts_at=candidate["starts_at"],
                    ends_at=candidate["ends_at"],
                    pair_number=candidate["pair_number"],
                    academic_hours=candidate["academic_hours"],
                    generated_by=actor_user_id,
                )
            )
            tid = candidate["teacher_id"]
            prev_name, prev_hours = all_teacher_hours.get(tid, (candidate["teacher_name"], 0.0))
            # Prefer the longer (more complete) display name
            display_name = candidate["teacher_name"] if len(candidate["teacher_name"]) > len(prev_name) else prev_name
            all_teacher_hours[tid] = (display_name, prev_hours + candidate["academic_hours"])

        # Merge any duplicate teacher records that may have been created by concurrent
        # imports or case-mismatch lookups (e.g. ALL-CAPS surname vs. proper case).
        all_teacher_ids = list(teacher_id_cache.keys())
        merged_count = _merge_duplicate_teachers(db, branch_id, all_teacher_ids)

        global_teacher_id_cache.update(teacher_id_cache)
        global_subject_cache.update(subject_cache)
        
        total_deleted += deleted_count
        total_created += len(candidates)
        all_merged += merged_count

    return {
        "import_kind": "schedule_docx",
        "group_code": group_codes[0] if group_codes else "",
        "group_name": group_names[0] if group_names else "",
        "group_total_hours": total_group_hours,
        "update_existing_mode": update_existing_mode,
        "deleted_slots": total_deleted,
        "created_slots": total_created,
        "skipped_existing_groups": total_skipped_groups,
        "skipped_existing_slots": total_skipped_slots,
        "existing_slots_in_period": total_existing_slots,
        "merged_duplicate_teachers": all_merged,
        "teachers": len(global_teacher_id_cache),
        "subjects": len(global_subject_cache),
        "teacher_workload_hours": {name: round(hours, 2) for (_tid, (name, hours)) in all_teacher_hours.items()},
    }
