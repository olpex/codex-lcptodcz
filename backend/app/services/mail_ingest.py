import email
import hashlib
import imaplib
import re
from datetime import datetime, timezone
from email.header import decode_header
from email.utils import parseaddr
from pathlib import Path
from uuid import uuid4

from docx import Document as DocxDocument
from pypdf import PdfReader
from sqlalchemy.orm import Session

from app.core.config import settings
from app.models import (
    Document,
    DocumentType,
    DraftStatus,
    ImportJob,
    JobStatus,
    MailMessage,
    MailStatus,
    OCRResult,
)
from app.services.import_export import IMPORT_UPDATE_MODES, parse_document_content, try_import_trainees
from app.services.ocr import guess_draft_from_text, ocr_image_file
from app.services.storage import detect_document_type, storage_path
from app.services.schedule_import import parse_schedule_docx, import_schedule_docx
from app.models import ScheduleSlot, Group, Trainee

GROUP_CODE_PATTERN = re.compile(r"(\d{1,4}\s*[-/]\s*\d{1,4})")
CONTRACT_KEYWORD_FALLBACK = "договор"
DASH_VARIANTS = "–—‑‒−﹘﹣"


def _normalize_compact(value: str | None) -> str:
    return " ".join((value or "").strip().lower().split())


def is_contract_sender(sender_name: str, sender_email: str) -> bool:
    expected_name = settings.imap_contract_sender_name_normalized
    expected_email = settings.imap_contract_sender_email_normalized
    if not expected_email:
        return False
    if _normalize_compact(sender_email) != expected_email:
        return False
    if expected_name and _normalize_compact(sender_name) != expected_name:
        actual_name = _normalize_compact(sender_name)
        if expected_name not in actual_name:
            return False
    return True


def extract_contract_group_code(filename: str | None) -> str | None:
    if not is_contract_attachment_filename(filename):
        return None

    stem = _normalized_contract_filename_stem(filename)
    if stem is None:
        return None

    # Accept both "Договори 73-26 ..." and "73-26 ... Договори ...".
    match = GROUP_CODE_PATTERN.search(stem)
    if not match:
        return None
    raw_group = "".join(match.group(1).split())
    return raw_group.replace("–", "-").replace("—", "-")


def _normalized_contract_filename_stem(filename: str | None) -> str | None:
    if not filename:
        return None
    lower = filename.strip().lower()
    if "." not in lower:
        return None
    stem, ext = lower.rsplit(".", 1)
    if ext not in {"xlsx", "xls"}:
        return None

    normalized_stem = stem.replace("_", " ").replace("\u00a0", " ")
    for dash in DASH_VARIANTS:
        normalized_stem = normalized_stem.replace(dash, "-")
    return _normalize_compact(normalized_stem)


def is_contract_attachment_filename(filename: str | None) -> bool:
    stem_compact = _normalized_contract_filename_stem(filename)
    if not stem_compact:
        return False

    keyword = _normalize_compact(settings.imap_contract_attachment_prefix)
    keyword_matched = (keyword and keyword in stem_compact) or (CONTRACT_KEYWORD_FALLBACK in stem_compact)
    if not keyword_matched:
        return False
    return True


def _decode_header(value: str | None) -> str:
    if not value:
        return ""
    fragments = decode_header(value)
    decoded: list[str] = []
    for fragment, encoding in fragments:
        if isinstance(fragment, bytes):
            decoded.append(fragment.decode(encoding or "utf-8", errors="ignore"))
        else:
            decoded.append(fragment)
    return "".join(decoded)


def _message_id_with_uid(message_id: str, uid_token: str) -> str:
    digest = hashlib.sha1(f"{message_id}|uid:{uid_token}".encode("utf-8")).hexdigest()[:12]
    max_base_len = 255 - len(digest) - 1
    base = message_id[:max_base_len]
    return f"{base}#{digest}"


def _extract_text_from_file(path: str, doc_type: DocumentType) -> str:
    if doc_type == DocumentType.DOCX:
        document = DocxDocument(path)
        return "\n".join(para.text for para in document.paragraphs if para.text.strip())
    if doc_type == DocumentType.PDF:
        reader = PdfReader(path)
        return "\n".join((page.extract_text() or "") for page in reader.pages)
    return ocr_image_file(path)


def is_duplicate_attachment(db: Session, branch_id: str, filename: str, file_path: str, doc_type: DocumentType, sender_email: str = "", subject: str = "") -> bool:
    if doc_type == DocumentType.XLSX and is_contract_attachment_filename(filename):
        group_code = extract_contract_group_code(filename)
        if group_code:
            exists = db.query(Trainee).filter(
                Trainee.branch_id == branch_id,
                Trainee.group_code == group_code,
                Trainee.is_deleted == False
            ).first()
            if exists:
                return True
    elif doc_type == DocumentType.DOCX and sender_email == "lcptodcz@gmail.com":
        import logging
        logger = logging.getLogger("mail.ingest")
        try:
            parsed_list = parse_schedule_docx(file_path)
        except Exception as e:
            logger.info(f"is_duplicate_attachment: Помилка парсингу {filename}: {e}. Вважаємо не дублікатом.")
            # If we can't parse it, we don't know what groups it contains, so we process it to let the job fail or process properly.
            return False

        if not parsed_list:
            logger.info(f"is_duplicate_attachment: У файлі {filename} не знайдено груп. Вважаємо не дублікатом.")
            return False

        for parsed in parsed_list:
            group_code = parsed.get("group_code")
            if not group_code:
                continue
            
            group = db.query(Group).filter(Group.branch_id == branch_id, Group.code == group_code).first()
            # If any group from the file doesn't exist in DB, it's NOT a duplicate (we must process it)
            if not group:
                logger.info(f"is_duplicate_attachment: Група {group_code} відсутня в БД. Файл {filename} не дублікат.")
                return False
            
            has_slots = db.query(ScheduleSlot.id).filter(ScheduleSlot.group_id == group.id).first() is not None
            # If any group exists but has no schedule slots, it's NOT a duplicate (we must process it)
            if not has_slots:
                logger.info(f"is_duplicate_attachment: Для групи {group_code} відсутні слоти. Файл {filename} не дублікат.")
                return False

        # Only skip as duplicate if ALL groups from the file exist AND have schedule slots.
        logger.info(f"is_duplicate_attachment: Всі групи з файлу {filename} присутні та мають слоти. Вважаємо дублікатом.")
        return True
    return False


def ingest_mailbox(db: Session) -> dict:
    if not settings.imap_host or not settings.imap_user or not settings.imap_password:
        return {"processed": 0, "message": "IMAP не налаштовано"}

    branch_id = settings.imap_branch_id or "main"
    allowed_senders = settings.imap_allowed_senders_list
    processed = 0
    mailbox = imaplib.IMAP4_SSL(settings.imap_host, settings.imap_port)
    mailbox.login(settings.imap_user, settings.imap_password)
    mailbox.select(settings.imap_mailbox)
    status, data = mailbox.search(None, "UNSEEN")
    if status != "OK":
        mailbox.logout()
        return {"processed": 0, "message": "Не вдалося отримати список листів"}

    ids = data[0].split()
    for msg_id in ids:
        status, message_data = mailbox.fetch(msg_id, "(UID RFC822)")
        if status != "OK" or not message_data:
            continue

        fetch_meta = b""
        raw = b""
        for part in message_data:
            if isinstance(part, tuple):
                if isinstance(part[0], bytes):
                    fetch_meta = part[0]
                if isinstance(part[1], bytes):
                    raw = part[1]
                break
        if not raw:
            continue

        uid_match = re.search(rb"UID\s+(\d+)", fetch_meta)
        uid_token = uid_match.group(1).decode("ascii") if uid_match else ""

        parsed = email.message_from_bytes(raw)
        message_id = (parsed.get("Message-ID") or "").strip() or f"local-{uuid4().hex}"

        existing = (
            db.query(MailMessage)
            .filter(MailMessage.branch_id == branch_id, MailMessage.message_id == message_id)
            .first()
        )
        if existing and uid_token:
            # Some providers may reuse Message-ID; append stable IMAP UID hash to avoid dropping newer messages.
            message_id = _message_id_with_uid(message_id, uid_token)
            existing = (
                db.query(MailMessage)
                .filter(MailMessage.branch_id == branch_id, MailMessage.message_id == message_id)
                .first()
            )
        if existing:
            # If the email is UNSEEN but already exists in DB, it means the user manually marked it as unread
            # to force reprocessing. We append a random suffix to bypass the unique constraint and process it as new.
            message_id = f"{message_id}:re:{uuid4().hex[:8]}"

        subject = _decode_header(parsed.get("Subject"))
        sender = _decode_header(parsed.get("From"))
        sender_lower = sender.lower()
        sender_name, sender_email = parseaddr(sender)
        sender_name = _decode_header(sender_name)
        sender_email = sender_email.strip().lower()
        sender_is_contract_source = is_contract_sender(sender_name, sender_email)
        received_at = datetime.now(timezone.utc)

        snippet = ""
        for part in parsed.walk():
            content_type = part.get_content_type()
            content_disposition = str(part.get("Content-Disposition", ""))
            if content_type == "text/plain" and "attachment" not in content_disposition:
                payload = part.get_payload(decode=True)
                if payload:
                    snippet = payload.decode(errors="ignore")[:500]
                    break

        record = MailMessage(
            branch_id=branch_id,
            message_id=message_id,
            sender=sender,
            subject=subject or "(без теми)",
            received_at=received_at,
            snippet=snippet,
            status=MailStatus.NEW,
        )
        db.add(record)
        db.flush()

        if allowed_senders and not any(token in sender_lower for token in allowed_senders):
            record.status = MailStatus.PROCESSED
            record.snippet = (record.snippet or "") + " [Пропущено: відправник не у списку дозволених]"
            processed += 1
            break

        attachment_notes: list[str] = []
        for part in parsed.walk():
            content_disposition = str(part.get("Content-Disposition", ""))
            if "attachment" not in content_disposition:
                continue

            filename = _decode_header(part.get_filename()) or f"attachment_{uuid4().hex}.bin"
            payload = part.get_payload(decode=True)
            if not payload:
                continue

            doc_type = detect_document_type(filename)
            out_path = storage_path() / f"{uuid4().hex}_{filename}"
            with Path(out_path).open("wb") as handle:
                handle.write(payload)

            if is_duplicate_attachment(db, branch_id, filename, str(out_path), doc_type, sender_email, subject):
                out_path.unlink(missing_ok=True)
                attachment_notes.append(f"Пропущено дублікат ({filename})")
                continue

            document = Document(
                branch_id=branch_id,
                file_name=filename,
                file_path=str(out_path),
                file_type=doc_type,
                source="mail",
                mime_type=part.get_content_type(),
            )
            db.add(document)
            db.flush()
            if record.raw_document_id is None:
                record.raw_document_id = document.id

            contract_group_code = extract_contract_group_code(filename)
            is_contract_attachment = is_contract_attachment_filename(filename)
            if sender_is_contract_source and is_contract_attachment and doc_type == DocumentType.XLSX:
                import_mode = settings.imap_contract_update_mode if settings.imap_contract_update_mode in IMPORT_UPDATE_MODES else "overwrite"
                job = ImportJob(
                    branch_id=branch_id,
                    idempotency_key=f"{branch_id}:mail-contracts:{uuid4().hex}",
                    document_id=document.id,
                    status=JobStatus.RUNNING,
                    message="Автоімпорт договорів із пошти виконується",
                    started_at=datetime.now(timezone.utc),
                    result_payload={
                        "source": "mail_auto_contracts",
                        "message_id": message_id,
                        "sender_name": sender_name,
                        "sender_email": sender_email,
                        "group_code_hint": contract_group_code,
                        "import_mode": import_mode,
                    },
                )
                db.add(job)
                db.flush()
                try:
                    parsed_content = parse_document_content(str(out_path), doc_type)
                    import_result = try_import_trainees(db, parsed_content, branch_id, update_existing_mode=import_mode)
                    job.status = JobStatus.SUCCEEDED
                    job.result_payload = {
                        "source": "mail_auto_contracts",
                        "message_id": message_id,
                        "sender_name": sender_name,
                        "sender_email": sender_email,
                        "group_code_hint": contract_group_code,
                        "import_mode": import_mode,
                        "parsed": {
                            key: value for key, value in parsed_content.items() if key != "data"
                        },
                        "import_result": import_result,
                    }
                    job.message = "Автоімпорт договорів із пошти виконано"
                    job.finished_at = datetime.now(timezone.utc)
                    attachment_notes.append(f"Імпорт договорів виконано ({filename})")
                except Exception as exc:
                    job.status = JobStatus.FAILED
                    job.message = f"Помилка автоімпорту договорів: {exc}"
                    job.finished_at = datetime.now(timezone.utc)
                    db.add(job)
                    attachment_notes.append(f"Помилка імпорту договорів ({filename})")
                db.flush()
                continue

            if doc_type in {DocumentType.XLSX, DocumentType.CSV}:
                attachment_notes.append(
                    f"Excel-вкладення пропущено ({filename}): не відповідає правилу 'Договори' або відправнику"
                )
                continue

            # Check if this DOCX is a schedule attachment
            is_schedule_attachment = False
            if doc_type == DocumentType.DOCX and sender_email == "lcptodcz@gmail.com":
                is_schedule_attachment = True

            if is_schedule_attachment:
                job = ImportJob(
                    branch_id=branch_id,
                    idempotency_key=f"{branch_id}:mail-schedules:{uuid4().hex}",
                    document_id=document.id,
                    status=JobStatus.RUNNING,
                    message="Автоімпорт розкладу із пошти виконується",
                    started_at=datetime.now(timezone.utc),
                    result_payload={
                        "source": "mail_auto_schedules",
                        "message_id": message_id,
                        "sender_name": sender_name,
                        "sender_email": sender_email,
                        "filename": filename,
                    },
                )
                db.add(job)
                db.flush()
                try:
                    import_result = import_schedule_docx(db, str(out_path), branch_id)
                    job.status = JobStatus.SUCCEEDED
                    job.result_payload["import_result"] = import_result
                    job.message = "Автоімпорт розкладу із пошти виконано"
                    job.finished_at = datetime.now(timezone.utc)
                    attachment_notes.append(f"Імпорт розкладу виконано ({filename})")
                except Exception as exc:
                    job.status = JobStatus.FAILED
                    job.message = f"Помилка автоімпорту розкладу: {exc}"
                    job.finished_at = datetime.now(timezone.utc)
                    attachment_notes.append(f"Помилка імпорту розкладу ({filename})")
                db.add(job)
                db.flush()
                continue

            text = _extract_text_from_file(str(out_path), doc_type)
            draft_type, payload_guess = guess_draft_from_text(text)
            ocr_result = OCRResult(
                branch_id=branch_id,
                document_id=document.id,
                extracted_text=text,
                structured_payload=payload_guess,
                draft_type=draft_type,
                status=DraftStatus.PENDING,
                confidence=0.75 if text else 0.1,
            )
            db.add(ocr_result)
            attachment_notes.append(f"Створено OCR-чернетку ({filename})")

        if attachment_notes:
            note = " | ".join(attachment_notes)
            if record.snippet:
                record.snippet = f"{record.snippet} [{note}]"
            else:
                record.snippet = note

        record.status = MailStatus.PROCESSED
        processed += 1
        break

    db.commit()
    mailbox.logout()
    return {"processed": processed}
