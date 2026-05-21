import csv
import base64
import json
import re
import socket
import tempfile
import time
from datetime import datetime, time as datetime_time, timezone
from io import BytesIO
from pathlib import Path
from typing import Any
from urllib.parse import quote, unquote, urlencode, urlparse, parse_qs
from urllib.error import HTTPError, URLError
from urllib.request import Request, urlopen
from zoneinfo import ZoneInfo

from docx import Document as DocxDocument
from jose import jwt
from openpyxl import Workbook, load_workbook
from sqlalchemy import func
from sqlalchemy.exc import DBAPIError
from sqlalchemy.orm import Session

from app.core.config import settings
from app.models import Group, GroupStatus, JournalMonitorEntry, JournalMonitorEvent, JournalMonitorSection, JournalWorkloadEntry, ScheduleSlot, Teacher, Trainee
from app.services.cache import cache_get_json, cache_set_json, hashed_cache_part
from app.services.import_export import save_report_file, try_import_trainees

GOOGLE_DRIVE_FOLDER_MIME = "application/vnd.google-apps.folder"
GOOGLE_DRIVE_SHEETS_MIME = "application/vnd.google-apps.spreadsheet"
GOOGLE_DRIVE_DOCS_MIME = "application/vnd.google-apps.document"
GOOGLE_DRIVE_XLSX_MIME = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
GOOGLE_DRIVE_XLS_MIME = "application/vnd.ms-excel"
GOOGLE_DRIVE_DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
GOOGLE_DRIVE_SCOPE = "https://www.googleapis.com/auth/drive"
GOOGLE_TOKEN_URI = "https://oauth2.googleapis.com/token"
_POSTGRES_TRANSIENT_LOCK_ERROR_CODES = {"40P01", "40001", "55P03"}
TRANSIENT_SYNC_LOCK_MESSAGE = "Синхронізацію Drive тимчасово відкладено: база даних зайнята. Спробуйте оновити ще раз за хвилину."
SERVICE_ACCOUNT_SETUP_MESSAGE = (
    "Для приватної Google Drive папки надайте доступ Editor для email service account "
    "suptc-drive-journal-monitor@gen-lang-client-0242013668.iam.gserviceaccount.com "
    "і задайте на backend змінну GOOGLE_DRIVE_SERVICE_ACCOUNT_JSON з JSON-ключем цього service account. "
    "GOOGLE_DRIVE_API_KEY потрібен тільки для публічних папок."
)
GROUP_CODE_PATTERN = re.compile(r"^\s*([0-9]{1,4}\s*[A-Za-zА-Яа-яІіЇїЄєҐґ]?\s*[-–—]\s*[0-9]{2,4})")
EXPORT_FORMATS = {"xlsx", "pdf", "docx", "csv"}
JOURNAL_WORKLOAD_START_YEAR = 2026
JOURNAL_MONITOR_MESSAGE_LIMIT = 500
GOOGLE_DRIVE_REQUEST_TIMEOUT_SECONDS = 30
GOOGLE_DRIVE_REQUEST_RETRY_ATTEMPTS = 3
GOOGLE_DRIVE_REQUEST_RETRY_DELAY_SECONDS = 0.75
GOOGLE_DRIVE_LIST_CACHE_TTL_SECONDS = 45
SUBJECTLESS_WORKLOAD_SUBJECT_NAME = "Без назви предмета"
JOURNAL_DAILY_ACTIVITY_ZONE = ZoneInfo("Europe/Kyiv")
JOURNAL_DAILY_ACTIVITY_START_HOUR = 8
AUTO_QUEUE_SIGNAL_MARKERS = (
    "після змін у google drive",
    "повторної обробки",
    "ручного опрацювання",
    "новий журнал у google drive",
)
_service_account_token_cache: dict[str, Any] = {"access_token": None, "expires_at": 0.0}


class JournalNoDataError(ValueError):
    def __init__(self, message: str, *, workload_hours: float | None = None):
        super().__init__(message)
        self.workload_hours = round(float(workload_hours), 2) if workload_hours is not None else None


class JournalMissingWorkbookError(JournalNoDataError):
    pass


def normalize_group_code(value: str | None) -> str:
    raw = (value or "").strip()
    raw = raw.replace("–", "-").replace("—", "-")
    raw = re.sub(r"\s*-\s*", "-", raw)
    return raw.casefold()


def display_group_code(value: str | None) -> str | None:
    raw = (value or "").strip()
    if not raw:
        return None
    raw = raw.replace("–", "-").replace("—", "-")
    return re.sub(r"\s*-\s*", "-", raw)


def _journal_group_name(journal_name: str, group_code: str | None) -> str:
    value = _norm(journal_name)
    if group_code:
        escaped = re.escape(group_code).replace("\\-", r"[-–—]")
        value = re.sub(rf"^\s*{escaped}\s*[-–—:]?\s*", "", value, flags=re.IGNORECASE)
    else:
        value = re.sub(GROUP_CODE_PATTERN, "", value, count=1).strip(" -–—:")
    return (value or journal_name or group_code or "Група")[:255]


def ensure_groups_for_journal_entries(db: Session, section: JournalMonitorSection) -> int:
    existing = {
        normalize_group_code(group.code): group
        for group in db.query(Group).filter(Group.branch_id == section.branch_id).all()
    }
    created = 0
    for entry in section.entries:
        group_code = display_group_code(entry.group_code)
        if not group_code:
            continue
        cache_key = normalize_group_code(group_code)
        existing_group = existing.get(cache_key)
        if existing_group:
            next_name = _journal_group_name(entry.journal_name, group_code)
            next_capacity = max(int(entry.trainee_count or 0), existing_group.capacity or 30, 30)
            if existing_group.name != next_name:
                existing_group.name = next_name
            if existing_group.capacity < next_capacity:
                existing_group.capacity = next_capacity
            if existing_group.hidden_from_registry:
                existing_group.hidden_from_registry = False
                created += 1
            db.add(existing_group)
            continue
        group = Group(
            branch_id=section.branch_id,
            code=group_code,
            name=_journal_group_name(entry.journal_name, group_code),
            capacity=max(int(entry.trainee_count or 0), 30),
            status=GroupStatus.ACTIVE,
        )
        db.add(group)
        db.flush()
        existing[cache_key] = group
        created += 1
    return created


def hide_groups_for_deleted_journal_entries(db: Session, entries: list[JournalMonitorEntry]) -> int:
    hidden = 0
    deleted_ids = {entry.id for entry in entries if entry.id is not None}
    seen_codes: set[str] = set()
    for entry in entries:
        group_code = display_group_code(entry.group_code)
        if not group_code:
            continue
        normalized_code = normalize_group_code(group_code)
        if normalized_code in seen_codes:
            continue
        seen_codes.add(normalized_code)
        remaining_entry = (
            db.query(JournalMonitorEntry)
            .filter(
                JournalMonitorEntry.branch_id == entry.branch_id,
                JournalMonitorEntry.group_code == group_code,
                JournalMonitorEntry.id.notin_(deleted_ids),
            )
            .first()
        )
        if remaining_entry:
            continue
        group = (
            db.query(Group)
            .filter(Group.branch_id == entry.branch_id, Group.code == group_code)
            .first()
        )
        if group and not group.hidden_from_registry:
            group.hidden_from_registry = True
            db.add(group)
            hidden += 1
    return hidden


def archive_trainees_for_deleted_journal_entries(db: Session, entries: list[JournalMonitorEntry]) -> int:
    archived = 0
    deleted_ids = {entry.id for entry in entries if entry.id is not None}
    seen_codes: set[str] = set()
    now = datetime.now(timezone.utc)
    for entry in entries:
        group_code = display_group_code(entry.group_code)
        if not group_code:
            continue
        normalized_code = normalize_group_code(group_code)
        if normalized_code in seen_codes:
            continue
        seen_codes.add(normalized_code)
        remaining_entry = (
            db.query(JournalMonitorEntry)
            .filter(
                JournalMonitorEntry.branch_id == entry.branch_id,
                JournalMonitorEntry.group_code == group_code,
                JournalMonitorEntry.id.notin_(deleted_ids),
            )
            .first()
        )
        if remaining_entry:
            continue
        trainees = (
            db.query(Trainee)
            .filter(
                Trainee.branch_id == entry.branch_id,
                Trainee.group_code == group_code,
                Trainee.is_deleted.is_(False),
            )
            .all()
        )
        for trainee in trainees:
            trainee.is_deleted = True
            trainee.deleted_at = now
            trainee.group_code = None
            db.add(trainee)
            archived += 1
    return archived


def delete_workload_for_journal_entries(db: Session, entries: list[JournalMonitorEntry]) -> int:
    entry_ids = [entry.id for entry in entries if entry.id is not None]
    if not entry_ids:
        return 0
    deleted = (
        db.query(JournalWorkloadEntry)
        .filter(JournalWorkloadEntry.journal_monitor_entry_id.in_(entry_ids))
        .delete(synchronize_session=False)
    )
    for entry in entries:
        entry.workload_status = "pending"
        entry.workload_message = "Педнавантаження видалено разом із журналом"
        entry.workload_processed_at = None
        entry.workload_year = None
        entry.workload_hours = 0.0
        entry.workload_source_names = None
        db.add(entry)
    db.flush()
    return int(deleted or 0)


def remove_journal_entries_from_project(db: Session, entries: list[JournalMonitorEntry]) -> dict[str, int]:
    entries = [entry for entry in entries if entry is not None]
    hidden_groups = hide_groups_for_deleted_journal_entries(db, entries)
    archived_trainees = archive_trainees_for_deleted_journal_entries(db, entries)
    deleted_workload = delete_workload_for_journal_entries(db, entries)
    for entry in entries:
        db.delete(entry)
    db.flush()
    return {
        "entries": len(entries),
        "hidden_groups": hidden_groups,
        "archived_trainees": archived_trainees,
        "deleted_workload": deleted_workload,
    }


def _requeue_entry_after_drive_change(
    db: Session,
    entry: JournalMonitorEntry,
    *,
    requeue_workload: bool = True,
    requeue_trainees: bool = True,
) -> None:
    if requeue_workload and entry.workload_status != "pending":
        entry.workload_status = "pending"
        entry.workload_message = "Поставлено в чергу після змін у Google Drive"
        entry.workload_processed_at = None
        entry.workload_hours = 0.0
        entry.workload_source_names = None
        db.query(JournalWorkloadEntry).filter(JournalWorkloadEntry.journal_monitor_entry_id == entry.id).delete(
            synchronize_session=False
        )
    if requeue_trainees and entry.trainees_status != "pending":
        entry.trainees_status = "pending"
        entry.trainees_message = "Поставлено в чергу після змін у Google Drive"
        entry.trainees_processed_at = None
        entry.trainees_source_names = None
    db.add(entry)


def _journal_workbooks_modified_after(
    entry: JournalMonitorEntry,
    processed_at: datetime | None,
    service_account_json: str | None,
) -> bool:
    processed_at = _as_aware_utc(processed_at)
    if processed_at is None:
        return False
    if entry.drive_folder_id:
        modified_at = _as_aware_utc(entry.drive_modified_at)
        return bool(modified_at and modified_at > processed_at)
    try:
        files = list_drive_journal_workbook_files(entry.drive_file_id, service_account_json=service_account_json)
    except Exception:
        return False
    if not files:
        return bool(entry.workload_source_names or entry.trainees_source_names)
    return _workbook_files_modified_after(files, processed_at)


def _workbook_files_modified_after(workbook_files: list[dict[str, Any]] | None, processed_at: datetime | None) -> bool:
    processed_at = _as_aware_utc(processed_at)
    if processed_at is None or not workbook_files:
        return False
    for workbook_file in workbook_files:
        modified_at = _as_aware_utc(_parse_datetime(str(workbook_file.get("modifiedTime") or "")))
        if modified_at and modified_at > processed_at:
            return True
    return False


def _entry_workbook_file_payload(entry: JournalMonitorEntry) -> dict[str, Any]:
    payload: dict[str, Any] = {
        "id": entry.drive_file_id,
        "name": entry.journal_name,
        "mimeType": entry.drive_mime_type or "",
    }
    if entry.drive_modified_at:
        payload["modifiedTime"] = entry.drive_modified_at.isoformat()
    if entry.drive_created_at:
        payload["createdTime"] = entry.drive_created_at.isoformat()
    if entry.drive_url:
        payload["webViewLink"] = entry.drive_url
    return payload


def _entry_workbook_files(
    entry: JournalMonitorEntry,
    workbook_lister,
    service_account_json: str | None = None,
) -> list[dict[str, Any]]:
    if _entry_is_folder_audit_only(entry):
        return []
    if entry.drive_folder_id:
        return [_entry_workbook_file_payload(entry)]
    return workbook_lister(entry.drive_file_id, service_account_json=service_account_json)


def _entry_is_folder_audit_only(entry: JournalMonitorEntry) -> bool:
    return (
        bool(entry.drive_file_id)
        and bool(entry.drive_folder_id)
        and entry.drive_file_id == entry.drive_folder_id
        and entry.drive_mime_type == GOOGLE_DRIVE_FOLDER_MIME
    )


def extract_group_code(folder_name: str) -> str | None:
    match = GROUP_CODE_PATTERN.search(folder_name or "")
    if not match:
        return None
    return display_group_code(match.group(1))


def extract_drive_folder_id(folder_url: str) -> str:
    value = (folder_url or "").strip()
    if not value:
        raise ValueError("Вкажіть URL папки Google Drive")

    parsed = urlparse(value)
    query_id = parse_qs(parsed.query).get("id")
    if query_id and query_id[0]:
        return query_id[0]

    match = re.search(r"/folders/([^/?#]+)", parsed.path)
    if match:
        return unquote(match.group(1))

    if re.fullmatch(r"[A-Za-z0-9_-]{10,}", value):
        return value

    raise ValueError("Не вдалося визначити ID папки Google Drive з посилання")


def _parse_datetime(value: str | None) -> datetime | None:
    if not value:
        return None
    try:
        return datetime.fromisoformat(value.replace("Z", "+00:00"))
    except ValueError:
        return None


def _is_transient_postgres_lock_error(exc: BaseException) -> bool:
    if not isinstance(exc, DBAPIError):
        return False
    orig = getattr(exc, "orig", None)
    code = getattr(orig, "pgcode", None) or getattr(orig, "sqlstate", None)
    if code in _POSTGRES_TRANSIENT_LOCK_ERROR_CODES:
        return True
    message = f"{orig or ''} {exc}".casefold()
    return "deadlock detected" in message or "could not serialize access" in message or "lock not available" in message


def _mark_sync_temporarily_busy(section: JournalMonitorSection) -> None:
    section.last_sync_status = "busy"
    section.last_sync_message = _clip_monitor_message(TRANSIENT_SYNC_LOCK_MESSAGE)


def _as_aware_utc(value: datetime | None) -> datetime | None:
    if value is None:
        return None
    if value.tzinfo is None:
        return value.replace(tzinfo=timezone.utc)
    return value.astimezone(timezone.utc)


def _journal_daily_cutoff(now: datetime | None = None) -> datetime:
    now_utc = _as_aware_utc(now or datetime.now(timezone.utc)) or datetime.now(timezone.utc)
    local_now = now_utc.astimezone(JOURNAL_DAILY_ACTIVITY_ZONE)
    local_cutoff = datetime.combine(
        local_now.date(),
        datetime_time(hour=JOURNAL_DAILY_ACTIVITY_START_HOUR),
        tzinfo=JOURNAL_DAILY_ACTIVITY_ZONE,
    )
    return local_cutoff.astimezone(timezone.utc)


def _drive_activity_modified_at(folder: dict[str, Any], workbook_files: list[dict[str, Any]] | None) -> datetime | None:
    timestamps = [_as_aware_utc(_parse_datetime(str(folder.get("modified_time") or "")))]
    for workbook_file in workbook_files or []:
        timestamps.append(_as_aware_utc(_parse_datetime(str(workbook_file.get("modifiedTime") or ""))))
    valid_timestamps = [item for item in timestamps if item is not None]
    return max(valid_timestamps) if valid_timestamps else None


def _update_daily_drive_change_start(
    entry: JournalMonitorEntry,
    *,
    cutoff_at: datetime,
    drive_created_at: datetime | None,
    folder: dict[str, Any],
    workbook_files: list[dict[str, Any]] | None,
) -> None:
    timestamps = [_as_aware_utc(_parse_datetime(str(folder.get("modified_time") or "")))]
    for workbook_file in workbook_files or []:
        timestamps.append(_as_aware_utc(_parse_datetime(str(workbook_file.get("modifiedTime") or ""))))
    changed_after_cutoff = [item for item in timestamps if item is not None and item >= cutoff_at]
    if not changed_after_cutoff:
        return

    candidate = min(changed_after_cutoff)
    created_at = _as_aware_utc(drive_created_at)
    if created_at and candidate <= created_at and all(item <= created_at for item in changed_after_cutoff):
        return

    existing = _as_aware_utc(entry.drive_change_started_at)
    if existing and existing >= cutoff_at:
        entry.drive_change_started_at = min(existing, candidate)
    else:
        entry.drive_change_started_at = candidate


def _norm(text: Any) -> str:
    return re.sub(r"\s+", " ", str(text or "").replace("\u00a0", " ")).strip()


def _clip_monitor_message(value: str | None) -> str | None:
    if value is None:
        return None
    text = _norm(value)
    return text[:JOURNAL_MONITOR_MESSAGE_LIMIT] if text else None


def _unique_ints(values: list[int] | tuple[int, ...] | None) -> list[int]:
    unique: list[int] = []
    seen: set[int] = set()
    for raw_value in values or []:
        try:
            value = int(raw_value)
        except (TypeError, ValueError):
            continue
        if value in seen:
            continue
        seen.add(value)
        unique.append(value)
    return unique


def _section_priority_entry_ids(section: JournalMonitorSection) -> list[int]:
    return _unique_ints(section.priority_entry_ids)


def _set_section_priority_queue(
    section: JournalMonitorSection,
    entry_ids: list[int],
    *,
    queue_year: int | None,
    message: str | None = None,
) -> None:
    section.priority_entry_ids = _unique_ints(entry_ids) or None
    section.priority_queue_year = queue_year
    if message is not None:
        section.last_processing_message = _clip_monitor_message(message)


def _clear_section_priority_queue(section: JournalMonitorSection, *, message: str | None = None) -> None:
    section.priority_entry_ids = None
    section.priority_queue_year = None
    if message is not None:
        section.last_processing_message = _clip_monitor_message(message)


def _has_processing_signal(message: str | None) -> bool:
    text = _norm(message).casefold()
    if not text:
        return False
    return any(marker in text for marker in AUTO_QUEUE_SIGNAL_MARKERS)


def _parse_hours(value: Any) -> float:
    text = _norm(value)
    match = re.search(r"([0-9]+(?:[.,][0-9]+)?)", text)
    if not match:
        return 0.0
    return float(match.group(1).replace(",", "."))


def _is_total_hours_row(text: str) -> bool:
    value = _norm(text).casefold()
    return (
        "загальний обсяг" in value
        or "всього годин" in value
        or "усього годин" in value
        or ("всього" in value and "год" in value)
        or ("усього" in value and "год" in value)
    )


def _split_teacher_name(full_name: str) -> tuple[str, str]:
    repaired = re.sub(r"(?<=[a-zа-яіїєґ])(?=[A-ZА-ЯІЇЄҐ])", " ", _norm(full_name))
    tokens = [part for part in repaired.split(" ") if part]
    if not tokens:
        return "Невідомий", "Викладач"
    if len(tokens) == 1:
        return tokens[0], "Викладач"
    return tokens[0], " ".join(tokens[1:])


def _short_teacher_display_name(teacher: Teacher | None) -> str:
    if teacher is None:
        return "Невідомий викладач"
    last_name = _norm(teacher.last_name) or "Невідомий"
    initials = [
        f"{part[0].upper()}."
        for part in _norm(teacher.first_name).split(" ")
        if part
    ]
    return f"{last_name} {' '.join(initials)}".strip()


def _normalize_teacher_identity_text(value: str | None) -> str:
    return re.sub(r"[^0-9A-Za-zА-Яа-яЇїІіЄєҐґ]", "", value or "").casefold()


def _split_teacher_cell(value: str) -> list[str]:
    repaired = re.sub(r"(?<=[a-zа-яіїєґ])(?=[A-ZА-ЯІЇЄҐ])", " ", _norm(value))
    parts = [
        part.strip(" .;:,")
        for part in re.split(r"\s*(?:[,;]|\n|\r|\s+та\s+|\s+і\s+|\s+&\s+)\s*", repaired)
        if part.strip(" .;:,")
    ]
    return parts or ([repaired] if repaired else [])


def _find_header_column(headers: list[str], keywords: tuple[str, ...], excluded: tuple[str, ...] = ()) -> int | None:
    for index, header in enumerate(headers):
        normalized = header.casefold()
        if excluded and any(item in normalized for item in excluded):
            continue
        if any(item in normalized for item in keywords):
            return index
    return None


UKRAINIAN_MOBILE_PREFIXES = "39|50|63|66|67|68|73|91|92|93|94|95|96|97|98|99"
PHONE_PATTERN = re.compile(
    rf"(?<!\d)(?:\+?38[\s().-]*)?0(?:{UKRAINIAN_MOBILE_PREFIXES})(?:[\s().-]*\d){{7}}(?!\d)"
)


def _normalize_ukrainian_mobile(value: str) -> str:
    digits = re.sub(r"\D", "", value)
    if len(digits) == 10 and digits.startswith("0"):
        return f"+38{digits}"
    if len(digits) == 12 and digits.startswith("380"):
        return f"+{digits}"
    return value.strip(" ,;")


def _split_address_phone_cell(value: Any) -> tuple[Any, str | None]:
    text = _norm(value)
    if not text:
        return value, None
    for match in PHONE_PATTERN.finditer(text):
        digits = re.sub(r"\D", "", match.group(0))
        if (digits.startswith("380") and len(digits) == 12) or (digits.startswith("0") and len(digits) == 10):
            phone = _normalize_ukrainian_mobile(match.group(0))
            address = f"{text[:match.start()]} {text[match.end():]}".strip(" ,;")
            address = re.sub(r"\s{2,}", " ", address)
            return address or None, phone
    return value, None


def _infer_journal_year(entry: JournalMonitorEntry, section: JournalMonitorSection | None = None) -> int | None:
    for value in (entry.group_code, entry.journal_name, section.name if section else None):
        text = _norm(value)
        if not text:
            continue
        full_year = re.search(r"\b(20\d{2})\b", text)
        if full_year:
            return int(full_year.group(1))
        group_year = re.search(r"[-–—/]\s*(\d{2})(?:\D|$)", text)
        if group_year:
            return 2000 + int(group_year.group(1))
    return None


def _workbook_display_name(name: str | None) -> str:
    value = _norm(name)
    if not value:
        return ""
    value = re.sub(r"\.(?:xlsx|xlsm|xls|csv)$", "", value, flags=re.IGNORECASE)
    return value.strip(" \"'«»“”„`")


def _journal_name_from_workbook(workbook_file: dict[str, Any], folder_name: str, folder_id: str) -> str:
    workbook_name = _workbook_display_name(str(workbook_file.get("name") or ""))
    workbook_id = str(workbook_file.get("id") or "").strip()
    if not workbook_name or workbook_name.casefold() in {folder_id.casefold(), workbook_id.casefold()}:
        return folder_name
    workbook_group_code = display_group_code(extract_group_code(workbook_name))
    if workbook_group_code and workbook_name.casefold() == workbook_group_code.casefold():
        return folder_name
    return workbook_name


def _source_name_key(value: str | None, group_code: str | None = None) -> str:
    text = _workbook_display_name(value)
    if group_code:
        escaped = re.escape(display_group_code(group_code) or group_code).replace("\\-", r"[-–—]")
        text = re.sub(rf"^\s*{escaped}\s*[-–—:]?\s*", "", text, flags=re.IGNORECASE)
    text = re.sub(r"^\s*журнал\s*[-–—:]?\s*", "", text, flags=re.IGNORECASE)
    return re.sub(r"[^0-9A-Za-zА-Яа-яІіЇїЄєҐґ]+", "", text).casefold()


_SOURCE_NAME_STOP_WORDS = {
    "в",
    "вп",
    "за",
    "з",
    "і",
    "й",
    "кат",
    "категорія",
    "категорії",
    "під",
    "практика",
    "практичні",
    "практичний",
    "роботи",
    "робота",
    "теорія",
    "час",
    "виробн",
    "виробниче",
    "виробничий",
}


def _source_name_tokens(value: str | None, group_code: str | None = None) -> set[str]:
    text = _workbook_display_name(value)
    if group_code:
        escaped = re.escape(display_group_code(group_code) or group_code).replace("\\-", r"[-–—]")
        text = re.sub(rf"^\s*{escaped}\s*[-–—:]?\s*", "", text, flags=re.IGNORECASE)
    text = re.sub(r"^\s*журнал\s*[-–—:]?\s*", "", text, flags=re.IGNORECASE)
    tokens = {
        token.casefold()
        for token in re.findall(r"[0-9A-Za-zА-Яа-яІіЇїЄєҐґ]{2,}", text)
    }
    return {token for token in tokens if token not in _SOURCE_NAME_STOP_WORDS}


def _source_name_has_group_prefix(value: str | None, group_code: str | None = None) -> bool:
    if not group_code:
        return False
    text = _workbook_display_name(value)
    escaped = re.escape(display_group_code(group_code) or group_code).replace("\\-", r"[-–—]")
    return bool(re.match(rf"^\s*{escaped}(?:\s|[-–—:]|$)", text, flags=re.IGNORECASE))


def _source_name_repeats_journal(source_name: str, entry: JournalMonitorEntry, journal_key: str) -> bool:
    source_key = _source_name_key(source_name, entry.group_code)
    if not source_key:
        return True
    if journal_key and (source_key.startswith(journal_key) or journal_key.startswith(source_key)):
        return True

    source_tokens = _source_name_tokens(source_name, entry.group_code)
    journal_tokens = _source_name_tokens(entry.journal_name, entry.group_code)
    if not source_tokens or not journal_tokens:
        return False

    has_group_prefix = _source_name_has_group_prefix(source_name, entry.group_code)
    if not has_group_prefix and len(source_tokens) <= 2:
        return False

    shared_tokens = source_tokens & journal_tokens
    if source_tokens.issubset(journal_tokens) or journal_tokens.issubset(source_tokens):
        return True
    if len(source_tokens) <= 3 and len(shared_tokens) >= max(1, len(source_tokens) - 1):
        return True
    if len(shared_tokens) >= 3:
        return True
    prefix_matches = sum(
        1
        for source_token in source_tokens
        if any(source_token.startswith(journal_token[:5]) or journal_token.startswith(source_token[:5]) for journal_token in journal_tokens)
    )
    if len(source_tokens) <= 3 and prefix_matches >= max(1, len(source_tokens) - 1):
        return True
    if prefix_matches >= 2:
        return True
    return False


def _visible_source_names(entry: JournalMonitorEntry, source_names: list[str] | None) -> list[str]:
    journal_key = _source_name_key(entry.journal_name, entry.group_code)
    visible: list[str] = []
    seen: set[str] = set()
    for raw_name in source_names or []:
        display_name = _workbook_display_name(raw_name)
        source_key = _source_name_key(display_name, entry.group_code)
        if not display_name or not source_key or source_key in seen:
            continue
        if journal_key and _source_name_repeats_journal(display_name, entry, journal_key):
            continue
        visible.append(display_name)
        seen.add(source_key)
    return visible


def _clear_repeated_contract_numbers(rows: list[dict[str, Any]]) -> list[dict[str, Any]]:
    counts: dict[str, int] = {}
    for row in rows:
        contract = _norm(row.get("№ договору") or row.get("Номер в журналі З-СНН"))
        if contract:
            counts[contract.casefold()] = counts.get(contract.casefold(), 0) + 1
    repeated = {contract for contract, count in counts.items() if count > 1}
    if not repeated:
        return rows
    cleaned_rows: list[dict[str, Any]] = []
    for row in rows:
        cleaned = dict(row)
        contract = _norm(cleaned.get("№ договору") or cleaned.get("Номер в журналі З-СНН"))
        if contract.casefold() in repeated:
            cleaned["№ договору"] = None
            cleaned["Номер в журналі З-СНН"] = None
        cleaned_rows.append(cleaned)
    return cleaned_rows


def _trainee_date_key(value: Any) -> str:
    if isinstance(value, datetime):
        return value.date().isoformat()
    if hasattr(value, "isoformat") and not isinstance(value, str):
        return value.isoformat()
    text = _norm(value)
    if not text:
        return ""
    for fmt in ("%d.%m.%Y", "%Y-%m-%d", "%d/%m/%Y", "%d-%m-%Y", "%d.%m.%y"):
        try:
            return datetime.strptime(text, fmt).date().isoformat()
        except ValueError:
            continue
    return text.casefold()


def _trainee_identity(last_name: Any, first_name: Any, birth_date: Any) -> tuple[str, str, str]:
    return (_norm(last_name).casefold(), _norm(first_name).casefold(), _trainee_date_key(birth_date))


def _archive_missing_group_trainees(db: Session, branch_id: str, group_code: str | None, rows: list[dict[str, Any]]) -> int:
    display_code = display_group_code(group_code)
    if not display_code:
        return 0
    incoming = {
        _trainee_identity(
            row.get("Прізвище"),
            " ".join(part for part in [row.get("Ім'я"), row.get("По батькові")] if _norm(part)),
            row.get("Дата народження"),
        )
        for row in rows
    }
    incoming_loose = [
        {
            "last_name": _norm(row.get("Прізвище")).casefold(),
            "first_name": " ".join(part for part in [row.get("Ім'я"), row.get("По батькові")] if _norm(part)).casefold(),
            "birth_date": _trainee_date_key(row.get("Дата народження")),
        }
        for row in rows
    ]
    now = datetime.now(timezone.utc)
    archived = 0
    trainees = (
        db.query(Trainee)
        .filter(
            Trainee.branch_id == branch_id,
            Trainee.group_code == display_code,
            Trainee.is_deleted.is_(False),
        )
        .all()
    )
    for trainee in trainees:
        identity = _trainee_identity(trainee.last_name, trainee.first_name, trainee.birth_date)
        if identity in incoming:
            continue
        trainee_first = _norm(trainee.first_name).casefold()
        trainee_birth = _trainee_date_key(trainee.birth_date)
        matches_loose = any(
            item["last_name"] == _norm(trainee.last_name).casefold()
            and item["first_name"]
            and (item["first_name"].startswith(trainee_first) or trainee_first.startswith(item["first_name"].split(" ")[0]))
            and (not item["birth_date"] or not trainee_birth or item["birth_date"] == trainee_birth)
            for item in incoming_loose
        )
        if matches_loose:
            continue
        trainee.is_deleted = True
        trainee.deleted_at = now
        trainee.group_code = None
        db.add(trainee)
        archived += 1
    return archived


def _expected_trainee_count_from_message(message: str | None) -> int | None:
    match = re.search(r"слухачів із журналу:\s*(\d+)", message or "", flags=re.IGNORECASE)
    return int(match.group(1)) if match else None


def _entry_journal_trainee_count(entry: JournalMonitorEntry, registry_count: int = 0) -> int:
    if entry.trainees_status == "no_data":
        return 0
    if entry.trainees_status == "processed":
        expected_count = _expected_trainee_count_from_message(entry.trainees_message)
        if expected_count is not None:
            return expected_count
        return int(entry.trainee_count or registry_count or 0)
    if entry.trainees_status in {"pending", "failed"}:
        return 0
    return int(entry.trainee_count or registry_count or 0)


def _entry_needs_trainee_reimport(entry: JournalMonitorEntry, active_trainee_count: int | None = None) -> bool:
    expected_count = _expected_trainee_count_from_message(entry.trainees_message)
    if entry.trainees_status != "processed" or expected_count is None:
        return False
    if expected_count > int(entry.trainee_count or 0):
        return True
    return active_trainee_count is not None and active_trainee_count < expected_count


def _active_trainee_count_for_group(db: Session, branch_id: str, group_code: str | None) -> int:
    display_code = display_group_code(group_code)
    if not display_code:
        return 0
    return int(
        db.query(func.count(Trainee.id))
        .filter(
            Trainee.branch_id == branch_id,
            Trainee.group_code == display_code,
            Trainee.is_deleted.is_(False),
        )
        .scalar()
        or 0
    )


def _decode_service_account_json(raw_json: str | None = None) -> dict[str, Any]:
    raw_value = (raw_json if raw_json is not None else settings.google_drive_service_account_json).strip()
    if not raw_value:
        raise RuntimeError(SERVICE_ACCOUNT_SETUP_MESSAGE)

    try:
        if raw_value.startswith("{"):
            payload = json.loads(raw_value)
        else:
            try:
                payload = json.loads(base64.b64decode(raw_value).decode("utf-8"))
            except Exception:
                with Path(raw_value).open("r", encoding="utf-8") as handle:
                    payload = json.load(handle)
    except Exception as exc:
        raise RuntimeError("GOOGLE_DRIVE_SERVICE_ACCOUNT_JSON має бути JSON, base64(JSON) або шляхом до JSON-файлу") from exc

    if not payload.get("client_email") or not payload.get("private_key"):
        raise RuntimeError("JSON service account має містити client_email і private_key")
    return payload


def _get_service_account_access_token(raw_json: str | None = None) -> str:
    now = time.time()
    cache_key = raw_json or "__settings__"
    cached_token = _service_account_token_cache.get("access_token")
    if (
        cached_token
        and _service_account_token_cache.get("cache_key") == cache_key
        and float(_service_account_token_cache.get("expires_at") or 0) > now + 60
    ):
        return str(cached_token)

    account = _decode_service_account_json(raw_json)
    token_uri = account.get("token_uri") or GOOGLE_TOKEN_URI
    issued_at = int(now)
    claims = {
        "iss": account["client_email"],
        "scope": GOOGLE_DRIVE_SCOPE,
        "aud": token_uri,
        "iat": issued_at,
        "exp": issued_at + 3600,
    }
    assertion = jwt.encode(claims, account["private_key"], algorithm="RS256")
    body = urlencode(
        {
            "grant_type": "urn:ietf:params:oauth:grant-type:jwt-bearer",
            "assertion": assertion,
        }
    ).encode("utf-8")
    request = Request(
        token_uri,
        data=body,
        headers={"Content-Type": "application/x-www-form-urlencoded", "Accept": "application/json"},
        method="POST",
    )
    payload = json.loads(_read_drive_response(request).decode("utf-8"))
    access_token = payload.get("access_token")
    if not access_token:
        raise RuntimeError("Google OAuth не повернув access_token для service account")
    _service_account_token_cache["access_token"] = access_token
    _service_account_token_cache["cache_key"] = cache_key
    _service_account_token_cache["expires_at"] = now + int(payload.get("expires_in") or 3600)
    return str(access_token)


def _is_retryable_drive_error(exc: BaseException) -> bool:
    if isinstance(exc, HTTPError):
        return exc.code in {429, 500, 502, 503, 504}
    if isinstance(exc, URLError):
        return True
    return isinstance(exc, (TimeoutError, socket.timeout, ConnectionError))


def _read_drive_response(request_or_url: str | Request) -> bytes:
    last_error: Exception | None = None
    for attempt in range(GOOGLE_DRIVE_REQUEST_RETRY_ATTEMPTS):
        try:
            with urlopen(request_or_url, timeout=GOOGLE_DRIVE_REQUEST_TIMEOUT_SECONDS) as response:
                return response.read()
        except Exception as exc:
            last_error = exc
            is_last_attempt = attempt >= GOOGLE_DRIVE_REQUEST_RETRY_ATTEMPTS - 1
            if is_last_attempt or not _is_retryable_drive_error(exc):
                raise
            time.sleep(GOOGLE_DRIVE_REQUEST_RETRY_DELAY_SECONDS * (attempt + 1))
    if last_error:
        raise last_error
    raise RuntimeError("Google Drive не повернув відповідь")


def _drive_list_cache_key(kind: str, folder_id: str, service_account_json: str | None = None) -> str:
    credentials = service_account_json or settings.google_drive_service_account_json or settings.google_drive_api_key
    return f"drive:list:{kind}:{folder_id}:{hashed_cache_part(credentials)}"


def list_drive_child_folders(folder_id: str, service_account_json: str | None = None) -> list[dict[str, Any]]:
    effective_service_account_json = service_account_json or settings.google_drive_service_account_json
    use_service_account = bool(effective_service_account_json.strip())
    if not use_service_account and not settings.google_drive_api_key:
        raise RuntimeError(SERVICE_ACCOUNT_SETUP_MESSAGE)
    cache_key = _drive_list_cache_key("child-folders", folder_id, effective_service_account_json)
    cached = cache_get_json(cache_key)
    if isinstance(cached, list):
        return cached

    query = f"'{folder_id}' in parents and mimeType = '{GOOGLE_DRIVE_FOLDER_MIME}' and trashed = false"
    fields = "nextPageToken,files(id,name,mimeType,webViewLink,createdTime,modifiedTime)"
    page_token = ""
    folders: list[dict[str, Any]] = []
    access_token = _get_service_account_access_token(effective_service_account_json) if use_service_account else None
    while True:
        url = (
            "https://www.googleapis.com/drive/v3/files"
            f"?q={quote(query)}"
            f"&fields={quote(fields)}"
            "&pageSize=1000"
        )
        if not use_service_account:
            url += f"&key={quote(settings.google_drive_api_key)}"
        if page_token:
            url += f"&pageToken={quote(page_token)}"
        request_or_url: str | Request = url
        if access_token:
            request_or_url = Request(url, headers={"Authorization": f"Bearer {access_token}", "Accept": "application/json"})
        payload = json.loads(_read_drive_response(request_or_url).decode("utf-8"))
        for item in payload.get("files", []):
            folders.append(
                {
                    "id": item.get("id") or "",
                    "name": item.get("name") or "",
                    "url": item.get("webViewLink") or f"https://drive.google.com/drive/folders/{item.get('id')}",
                    "created_time": item.get("createdTime"),
                    "modified_time": item.get("modifiedTime"),
                }
            )
        page_token = payload.get("nextPageToken") or ""
        if not page_token:
            cache_set_json(cache_key, folders, GOOGLE_DRIVE_LIST_CACHE_TTL_SECONDS)
            return folders


def _drive_request_url(url: str, service_account_json: str | None = None) -> str | Request:
    effective_service_account_json = service_account_json or settings.google_drive_service_account_json
    if effective_service_account_json.strip():
        access_token = _get_service_account_access_token(effective_service_account_json)
        return Request(url, headers={"Authorization": f"Bearer {access_token}", "Accept": "application/json"})
    if settings.google_drive_api_key:
        separator = "&" if "?" in url else "?"
        return f"{url}{separator}key={quote(settings.google_drive_api_key)}"
    raise RuntimeError(SERVICE_ACCOUNT_SETUP_MESSAGE)


def list_drive_journal_workbook_files(folder_id: str, service_account_json: str | None = None) -> list[dict[str, Any]]:
    cache_key = _drive_list_cache_key("journal-workbooks", folder_id, service_account_json)
    cached = cache_get_json(cache_key)
    if isinstance(cached, list):
        return cached

    mime_filter = " or ".join(
        [
            f"mimeType = '{GOOGLE_DRIVE_SHEETS_MIME}'",
            f"mimeType = '{GOOGLE_DRIVE_XLSX_MIME}'",
            f"mimeType = '{GOOGLE_DRIVE_XLS_MIME}'",
        ]
    )
    query = f"'{folder_id}' in parents and ({mime_filter}) and trashed = false"
    fields = "nextPageToken,files(id,name,mimeType,webViewLink,createdTime,modifiedTime)"
    page_token = ""
    files: list[dict[str, Any]] = []
    while True:
        url = (
            "https://www.googleapis.com/drive/v3/files"
            f"?q={quote(query)}"
            f"&fields={quote(fields)}"
            "&pageSize=100"
        )
        if page_token:
            url += f"&pageToken={quote(page_token)}"
        payload = json.loads(_read_drive_response(_drive_request_url(url, service_account_json)).decode("utf-8"))
        files.extend(payload.get("files", []))
        page_token = payload.get("nextPageToken") or ""
        if not page_token:
            cache_set_json(cache_key, files, GOOGLE_DRIVE_LIST_CACHE_TTL_SECONDS)
            return files


def download_drive_file_bytes(
    file_id: str,
    mime_type: str | None = None,
    service_account_json: str | None = None,
) -> bytes:
    if mime_type == GOOGLE_DRIVE_SHEETS_MIME:
        url = (
            f"https://www.googleapis.com/drive/v3/files/{quote(file_id)}/export"
            f"?mimeType={quote(GOOGLE_DRIVE_XLSX_MIME)}"
        )
    elif mime_type == GOOGLE_DRIVE_DOCS_MIME:
        url = (
            f"https://www.googleapis.com/drive/v3/files/{quote(file_id)}/export"
            f"?mimeType={quote(GOOGLE_DRIVE_DOCX_MIME)}"
        )
    else:
        url = f"https://www.googleapis.com/drive/v3/files/{quote(file_id)}?alt=media"
    request_or_url = _drive_request_url(url, service_account_json)
    if isinstance(request_or_url, Request) and mime_type == GOOGLE_DRIVE_SHEETS_MIME:
        request_or_url.add_header("Accept", GOOGLE_DRIVE_XLSX_MIME)
    if isinstance(request_or_url, Request) and mime_type == GOOGLE_DRIVE_DOCS_MIME:
        request_or_url.add_header("Accept", GOOGLE_DRIVE_DOCX_MIME)
    return _read_drive_response(request_or_url)


def _find_disciplines_rows(workbook) -> list[list[Any]]:
    for sheet in workbook.worksheets:
        if sheet.title.strip().casefold() == "дисципліни":
            return [list(row) for row in sheet.iter_rows(values_only=True)]
    for sheet in workbook.worksheets:
        rows = [list(row) for row in sheet.iter_rows(values_only=True)]
        for raw_row in rows[:30]:
            headers = [_norm(value) for value in raw_row]
            subject_candidate = _find_header_column(headers, ("дисципл", "предмет", "назва"), ("стор", "год", "виклада", "піб"))
            hours_candidate = _find_header_column(headers, ("год", "кількість годин"))
            teacher_candidate = _find_header_column(headers, ("виклада", "піб", "прізвищ"))
            if hours_candidate is not None and teacher_candidate is not None:
                return rows
    raise JournalNoDataError("У файлі журналу не знайдено аркуш «Дисципліни»")


def _find_zv_rows(workbook) -> list[list[Any]]:
    for sheet in workbook.worksheets:
        if sheet.title.strip().casefold() == "зв":
            return [list(row) for row in sheet.iter_rows(values_only=True)]
    for sheet in workbook.worksheets:
        rows = [list(row) for row in sheet.iter_rows(values_only=True)]
        for raw_row in rows[:30]:
            if _zv_header_columns(raw_row)[0] is not None:
                return rows
    raise JournalNoDataError("У файлі журналу не знайдено аркуш «ЗВ»")


def _zv_header_columns(raw_row: list[Any]) -> tuple[int | None, int | None, int | None, int | None, int | None, int | None, int | None]:
    headers = [_norm(value) for value in raw_row]
    full_name_col = _find_header_column(headers, ("прізв", "піб", "слухач"), ("виклада",))
    if full_name_col is None:
        return None, None, None, None, None, None, None
    row_number_col = _find_header_column(headers, ("номер за поряд", "№", "п/п"), ("журнал", "з-снн", "догов"))
    journal_number_col = _find_header_column(headers, ("з-снн", "номер в журналі", "догов"))
    birth_date_col = _find_header_column(headers, ("дата народ",))
    tax_id_col = _find_header_column(headers, ("ідентиф", "рнокпп", "інн", "іпн"))
    address_col = _find_header_column(headers, ("адрес",))
    phone_col = _find_header_column(headers, ("тел",))
    supporting_columns = [birth_date_col, tax_id_col, address_col, phone_col, journal_number_col]
    if sum(1 for column in supporting_columns if column is not None) < 2:
        return None, None, None, None, None, None, None
    return full_name_col, row_number_col, birth_date_col, tax_id_col, address_col, phone_col, journal_number_col


def _looks_like_trainee_full_name(value: str) -> bool:
    text = _norm(value)
    if not text:
        return False
    normalized = text.casefold()
    blocked_fragments = (
        "прізв",
        "слухач",
        "п/п",
        "№",
        "номер",
        "загальні",
        "відомості",
    )
    if any(fragment in normalized for fragment in blocked_fragments):
        return False
    if not re.search(r"[A-Za-zА-Яа-яІіЇїЄєҐґ]", text):
        return False
    name_parts = [part for part in text.split(" ") if re.search(r"[A-Za-zА-Яа-яІіЇїЄєҐґ]", part)]
    return len(name_parts) >= 2


def _parse_table_row_number(value: Any) -> int | None:
    if isinstance(value, int):
        return value if value > 0 else None
    if isinstance(value, float) and value.is_integer():
        return int(value) if value > 0 else None
    text = _norm(value)
    match = re.match(r"^(\d+)\s*[.)]?$", text)
    if not match:
        return None
    number = int(match.group(1))
    return number if number > 0 else None


def _looks_like_table_teacher_name(value: Any) -> bool:
    text = _norm(value)
    if not text or re.search(r"\d", text):
        return False
    normalized = text.casefold()
    blocked_fragments = (
        "викладач",
        "майстер",
        "піб",
        "прізв",
        "назва",
        "дисципл",
        "предмет",
        "год",
        "стор",
        "загальний",
        "всього",
        "усього",
    )
    if any(fragment in normalized for fragment in blocked_fragments):
        return False
    name_parts = re.findall(r"[A-Za-zА-Яа-яЇїІіЄєҐґ]+", text)
    return len(name_parts) >= 2


def _last_teacher_cell_from_numbered_row(raw_row: list[Any], teacher_col: int | None) -> str:
    start_index = max((teacher_col or 0) + 1, 0)
    for value in reversed(raw_row[start_index:]):
        text = _norm(value)
        if _looks_like_table_teacher_name(text):
            return text
    return ""


def parse_journal_disciplines_xlsx(payload: bytes) -> list[dict[str, Any]]:
    workbook = load_workbook(BytesIO(payload), data_only=True, read_only=True)
    try:
        rows = _find_disciplines_rows(workbook)
    finally:
        workbook.close()

    header_index = -1
    subject_col: int | None = None
    hours_col: int | None = None
    pages_col: int | None = None
    teacher_col: int | None = None
    row_number_col: int | None = None
    for index, raw_row in enumerate(rows[:30]):
        headers = [_norm(value) for value in raw_row]
        subject_candidate = _find_header_column(headers, ("дисципл", "предмет", "назва"), ("стор", "год", "виклада", "піб"))
        hours_candidate = _find_header_column(headers, ("год", "кількість годин"))
        teacher_candidate = _find_header_column(headers, ("виклада", "піб", "прізвищ"))
        if hours_candidate is not None and teacher_candidate is not None:
            header_index = index
            subject_col = subject_candidate
            hours_col = hours_candidate
            pages_col = _find_header_column(headers, ("стор", "сторін"))
            teacher_col = teacher_candidate
            row_number_col = _find_header_column(headers, ("№", "номер", "п/п"), ("журнал", "з-снн", "догов"))
            break

    if header_index < 0 or hours_col is None or teacher_col is None:
        raise JournalNoDataError("На аркуші «Дисципліни» не знайдено колонки годин і викладача")

    parsed: dict[tuple[str, str], dict[str, Any]] = {}
    total_hours = 0.0
    incomplete_hours = 0.0
    for raw_row in rows[header_index + 1 :]:
        row_text = " ".join(_norm(value) for value in raw_row)
        if _is_total_hours_row(row_text):
            continue
        row_number = (
            _parse_table_row_number(raw_row[row_number_col])
            if row_number_col is not None and row_number_col < len(raw_row)
            else None
        )
        if row_number_col is not None and row_number is None:
            continue
        subject_name = (
            _norm(raw_row[subject_col] if subject_col < len(raw_row) else "")
            if subject_col is not None
            else ""
        ) or SUBJECTLESS_WORKLOAD_SUBJECT_NAME
        teacher_cell = _norm(raw_row[teacher_col] if teacher_col < len(raw_row) else "")
        if row_number is not None and not teacher_cell:
            teacher_cell = _last_teacher_cell_from_numbered_row(raw_row, teacher_col)
        hours = _parse_hours(raw_row[hours_col] if hours_col < len(raw_row) else "")
        pages = _norm(raw_row[pages_col] if pages_col is not None and pages_col < len(raw_row) else "")
        if hours <= 0:
            continue
        total_hours = round(total_hours + hours, 2)
        if not teacher_cell:
            incomplete_hours = round(incomplete_hours + hours, 2)
            continue
        for teacher_name in _split_teacher_cell(teacher_cell):
            key = (subject_name.casefold(), teacher_name.casefold())
            if key not in parsed:
                parsed[key] = {
                    "subject_name": subject_name[:255],
                    "hours": 0.0,
                    "pages": pages[:255] if pages else None,
                    "teacher_name": teacher_name,
                }
            parsed[key]["hours"] = round(float(parsed[key]["hours"]) + hours, 2)
            if pages and not parsed[key].get("pages"):
                parsed[key]["pages"] = pages[:255]

    if incomplete_hours > 0:
        raise JournalNoDataError(
            "На аркуші «Дисципліни» є години, але відсутні ПІБ викладачів",
            workload_hours=total_hours,
        )

    return list(parsed.values())


def parse_journal_zv_trainees_xlsx(payload: bytes, group_code: str | None = None, group_name: str | None = None) -> dict[str, Any]:
    workbook = load_workbook(BytesIO(payload), data_only=True, read_only=True)
    try:
        rows = _find_zv_rows(workbook)
    finally:
        workbook.close()

    header_index = -1
    full_name_col: int | None = None
    row_number_col: int | None = None
    birth_date_col: int | None = None
    tax_id_col: int | None = None
    address_col: int | None = None
    phone_col: int | None = None
    journal_number_col: int | None = None
    for index, raw_row in enumerate(rows[:30]):
        (
            full_name_candidate,
            row_number_candidate,
            birth_date_candidate,
            tax_id_candidate,
            address_candidate,
            phone_candidate,
            journal_number_candidate,
        ) = _zv_header_columns(raw_row)
        if full_name_candidate is not None:
            header_index = index
            full_name_col = full_name_candidate
            row_number_col = row_number_candidate
            journal_number_col = journal_number_candidate
            birth_date_col = birth_date_candidate
            tax_id_col = tax_id_candidate
            address_col = address_candidate
            phone_col = phone_candidate
            break

    if header_index < 0 or full_name_col is None:
        raise JournalNoDataError("На аркуші «ЗВ» не знайдено колонку ПІБ слухача")

    data: list[dict[str, Any]] = []
    for raw_row in rows[header_index + 1 :]:
        full_name = _norm(raw_row[full_name_col] if full_name_col < len(raw_row) else "")
        if not _looks_like_trainee_full_name(full_name):
            continue
        raw_address = raw_row[address_col] if address_col is not None and address_col < len(raw_row) else None
        raw_phone = raw_row[phone_col] if phone_col is not None and phone_col < len(raw_row) else None
        if address_col is not None and address_col == phone_col:
            address_value, phone_value = _split_address_phone_cell(raw_address)
        else:
            address_value = raw_address
            _phone_remainder, phone_value = _split_address_phone_cell(raw_phone)
        name_parts = full_name.split(" ")
        last_name = name_parts[0] if name_parts else ""
        first_name = name_parts[1] if len(name_parts) > 1 else ""
        middle_name = " ".join(name_parts[2:]) if len(name_parts) > 2 else ""
        data.append(
                {
                    "Номер за порядком": raw_row[row_number_col] if row_number_col is not None and row_number_col < len(raw_row) else None,
                    "Номер в журналі З-СНН": raw_row[journal_number_col] if journal_number_col is not None and journal_number_col < len(raw_row) else None,
                    "№ договору": raw_row[journal_number_col] if journal_number_col is not None and journal_number_col < len(raw_row) else None,
                    "Прізвище": last_name,
                "Ім'я": first_name,
                "По батькові": middle_name,
                "Дата народження": raw_row[birth_date_col] if birth_date_col is not None and birth_date_col < len(raw_row) else None,
                "Ідентифікаційний номер": raw_row[tax_id_col] if tax_id_col is not None and tax_id_col < len(raw_row) else None,
                "Домашня адреса": address_value,
                "Телефон": phone_value,
            }
        )
    if not data:
        raise JournalNoDataError("На аркуші «ЗВ» не знайдено рядків зі слухачами")

    return {
        "rows": len(data),
        "headers": list(data[0].keys()),
        "data": data,
        "sheet_name": "ЗВ",
        "default_group_code": group_code,
        "default_group_name": group_name,
    }


def _find_or_create_teacher(db: Session, branch_id: str, full_name: str) -> Teacher:
    last_name, first_name = _split_teacher_name(full_name)
    last_name = last_name[:120] or "Невідомий"
    first_name = first_name[:120] or "Викладач"
    incoming_identity = _normalize_teacher_identity_text(f"{last_name} {first_name}")
    for teacher in db.query(Teacher).filter(Teacher.branch_id == branch_id).all():
        existing_last = (teacher.last_name or "").strip()
        existing_identity = _normalize_teacher_identity_text(f"{teacher.last_name} {teacher.first_name}")
        incoming_raw = _normalize_teacher_identity_text(full_name)
        if existing_identity and incoming_raw and incoming_raw.startswith(existing_identity):
            return teacher
        if incoming_identity and existing_identity and incoming_identity.startswith(existing_identity):
            return teacher
        if existing_last.casefold() != last_name.casefold():
            existing_last_key = _normalize_teacher_identity_text(existing_last)
            incoming_raw = _normalize_teacher_identity_text(full_name)
            if not (existing_last_key and incoming_raw.startswith(existing_last_key)):
                continue
            first_name = re.sub(r"(?<=[a-zа-яіїєґ])(?=[A-ZА-ЯІЇЄҐ])", " ", _norm(full_name))[len(existing_last):].strip()[:120] or first_name
            if len(first_name) > len(teacher.first_name or ""):
                teacher.first_name = first_name
                db.add(teacher)
                db.flush()
            return teacher
        if existing_last.casefold() != last_name.casefold():
            continue
        existing_first = (teacher.first_name or "").strip().casefold()
        incoming_first = first_name.casefold()
        if existing_first == incoming_first:
            return teacher
        existing_initials = "".join(part[0].casefold() for part in _norm(teacher.first_name).split() if part)
        incoming_initials = "".join(part[0].casefold() for part in first_name.split() if part)
        if existing_initials and incoming_initials and existing_initials == incoming_initials:
            if len(first_name) > len(teacher.first_name or ""):
                teacher.first_name = first_name
                db.add(teacher)
                db.flush()
            return teacher

    teacher = Teacher(branch_id=branch_id, last_name=last_name, first_name=first_name, hourly_rate=0.0, is_active=True)
    db.add(teacher)
    db.flush()
    return teacher


def process_journal_workload_entry(
    db: Session,
    entry: JournalMonitorEntry,
    *,
    service_account_json: str | None = None,
    workbook_lister=None,
    workbook_downloader=None,
) -> dict[str, Any]:
    if workbook_lister is None:
        workbook_lister = list_drive_journal_workbook_files
    if workbook_downloader is None:
        workbook_downloader = download_drive_file_bytes
    files = _entry_workbook_files(entry, workbook_lister, service_account_json=service_account_json)
    if not files:
        entry.workload_source_names = []
        db.add(entry)
        db.flush()
        raise JournalMissingWorkbookError("У папці журналу не знайдено Google Sheet або Excel-файл")
    source_names = [
        display_name
        for display_name in (_workbook_display_name(str(file.get("name") or "")) for file in files)
        if display_name
    ]
    entry.workload_source_names = source_names
    db.add(entry)
    db.flush()
    rows: list[dict[str, Any]] = []
    errors: list[str] = []
    no_data_messages: list[str] = []
    no_data_hours = 0.0
    for workbook_file in sorted(files, key=lambda item: str(item.get("name") or "").casefold()):
        try:
            payload = workbook_downloader(
                str(workbook_file.get("id") or ""),
                mime_type=str(workbook_file.get("mimeType") or ""),
                service_account_json=service_account_json,
            )
            rows.extend(parse_journal_disciplines_xlsx(payload))
        except JournalNoDataError as exc:
            no_data_messages.append(str(exc))
            if exc.workload_hours is not None:
                no_data_hours = round(no_data_hours + exc.workload_hours, 2)
        except Exception as exc:
            errors.append(f"{workbook_file.get('name') or workbook_file.get('id')}: {exc}")
    if no_data_hours > 0 and not rows:
        detail = "; ".join(no_data_messages[:3])
        raise JournalNoDataError(
            f"Педнавантаження не імпортовано через неповні дані на аркуші «Дисципліни»: {detail}",
            workload_hours=no_data_hours,
        )
    if not rows:
        if errors and not no_data_messages:
            raise ValueError("; ".join(errors[:3]))
        detail = "; ".join(no_data_messages[:3])
        suffix = f": {detail}" if detail else ""
        raise JournalNoDataError(f"На аркушах «Дисципліни» не знайдено рядків з годинами викладачів{suffix}")

    db.query(JournalWorkloadEntry).filter(JournalWorkloadEntry.journal_monitor_entry_id == entry.id).delete(
        synchronize_session=False
    )
    aggregated: dict[tuple[int, str], dict[str, Any]] = {}
    for row in rows:
        teacher = _find_or_create_teacher(db, entry.branch_id, row["teacher_name"])
        hours = round(float(row["hours"]), 2)
        key = (teacher.id, row["subject_name"])
        if key not in aggregated:
            aggregated[key] = {"teacher": teacher, "subject_name": row["subject_name"], "hours": 0.0, "pages": row.get("pages")}
        aggregated[key]["hours"] = round(float(aggregated[key]["hours"]) + hours, 2)
        if row.get("pages") and not aggregated[key].get("pages"):
            aggregated[key]["pages"] = row.get("pages")
    total_hours = 0.0
    for item in aggregated.values():
        hours = float(item["hours"])
        total_hours += hours
        db.add(
            JournalWorkloadEntry(
                journal_monitor_entry_id=entry.id,
                branch_id=entry.branch_id,
                teacher_id=item["teacher"].id,
                subject_name=item["subject_name"],
                hours=hours,
                pages=item.get("pages"),
            )
        )
    entry.workload_status = "processed"
    skipped_messages = [*errors[:2], *no_data_messages[:2]]
    message_suffix = f"; частину файлів пропущено: {'; '.join(skipped_messages)}" if skipped_messages else ""
    entry.workload_message = f"Додано годин із журналу: {round(total_hours, 2)}{message_suffix}"
    entry.workload_processed_at = datetime.now(timezone.utc)
    entry.workload_hours = round(total_hours, 2)
    db.add(entry)
    db.flush()
    return {"entries": len(rows), "hours": round(total_hours, 2)}


def process_journal_trainees_entry(
    db: Session,
    entry: JournalMonitorEntry,
    *,
    service_account_json: str | None = None,
    workbook_lister=None,
    workbook_downloader=None,
) -> dict[str, Any]:
    if workbook_lister is None:
        workbook_lister = list_drive_journal_workbook_files
    if workbook_downloader is None:
        workbook_downloader = download_drive_file_bytes
    files = _entry_workbook_files(entry, workbook_lister, service_account_json=service_account_json)
    if not files:
        entry.trainees_source_names = []
        db.add(entry)
        db.flush()
        raise JournalMissingWorkbookError("У папці журналу не знайдено Google Sheet або Excel-файл")
    source_names = [
        display_name
        for display_name in (_workbook_display_name(str(file.get("name") or "")) for file in files)
        if display_name
    ]
    entry.trainees_source_names = source_names
    db.add(entry)
    db.flush()

    combined_data: list[dict[str, Any]] = []
    headers: list[str] = []
    no_data_messages: list[str] = []
    errors: list[str] = []
    for workbook_file in sorted(files, key=lambda item: str(item.get("name") or "").casefold()):
        try:
            payload = workbook_downloader(
                str(workbook_file.get("id") or ""),
                mime_type=str(workbook_file.get("mimeType") or ""),
                service_account_json=service_account_json,
            )
            parsed = parse_journal_zv_trainees_xlsx(payload, group_code=entry.group_code, group_name=entry.journal_name)
            if not headers:
                headers = list(parsed.get("headers") or [])
            combined_data.extend(parsed.get("data") or [])
        except JournalNoDataError as exc:
            no_data_messages.append(str(exc))
        except Exception as exc:
            errors.append(f"{workbook_file.get('name') or workbook_file.get('id')}: {exc}")

    if not combined_data:
        if errors and not no_data_messages:
            raise ValueError("; ".join(errors[:3]))
        message = no_data_messages[0] if no_data_messages else "На аркуші «ЗВ» не знайдено рядків зі слухачами"
        raise JournalNoDataError(message)

    combined_data = _clear_repeated_contract_numbers(combined_data)
    import_result = try_import_trainees(
        db,
        {
            "rows": len(combined_data),
            "headers": headers,
            "data": combined_data,
            "sheet_name": "ЗВ",
            "default_group_code": entry.group_code,
            "default_group_name": entry.journal_name,
        },
        entry.branch_id,
        update_existing_mode="overwrite",
        commit=False,
    )
    archived_missing = _archive_missing_group_trainees(db, entry.branch_id, entry.group_code, combined_data)
    changed_count = int(import_result.get("inserted") or 0) + int(import_result.get("updated_existing") or 0)
    changed_count += archived_missing
    total_seen = len(combined_data)
    entry.trainees_status = "processed"
    entry.trainees_message = f"Додано/оновлено слухачів із журналу: {total_seen}"
    if errors:
        entry.trainees_message = f"{entry.trainees_message}; частину файлів пропущено: {'; '.join(errors[:2])}"
    entry.trainees_processed_at = datetime.now(timezone.utc)
    entry.has_trainees = total_seen > 0
    entry.trainee_count = total_seen
    db.add(entry)
    db.flush()
    return {"entries": total_seen, "changed": changed_count, "import_result": import_result}


def process_next_journal_workload(
    db: Session,
    section: JournalMonitorSection,
    *,
    limit: int | None = 1,
    start_year: int = JOURNAL_WORKLOAD_START_YEAR,
    target_year: int | None = None,
    retry_failed: bool = False,
    entry_ids: set[int] | None = None,
) -> dict[str, Any]:
    from app.core.crypto import cipher

    processed = 0
    failed = 0
    skipped_year = 0
    handled = 0
    section_service_account_json = cipher.decrypt(section.service_account_json_encrypted)

    def status_priority(entry: JournalMonitorEntry) -> int:
        if entry.workload_status in {"pending", "needs_regeneration"}:
            return 0
        if retry_failed and entry.workload_status == "failed":
            return 1
        return 9

    entries = sorted(
        [entry for entry in section.entries if entry_ids is None or entry.id in entry_ids],
        key=lambda item: (
            status_priority(item),
            _infer_journal_year(item, section) or 9999,
            (item.group_code or "~~~~").casefold(),
            item.journal_name.casefold(),
        ),
    )
    for entry in entries:
        if limit is not None and handled >= limit:
            break
        if _entry_is_folder_audit_only(entry):
            entry.workload_status = "no_data"
            entry.workload_message = "У теці Google Drive не знайдено файлу журналу"
            entry.workload_processed_at = entry.workload_processed_at or datetime.now(timezone.utc)
            db.add(entry)
            continue
        if entry.workload_status in {"processed", "no_data", "skipped_year"}:
            if _journal_workbooks_modified_after(entry, entry.workload_processed_at, section_service_account_json):
                _requeue_entry_after_drive_change(db, entry)
            elif entry.workload_status != "no_data":
                continue
        if entry.workload_status == "failed" and not retry_failed:
            continue
        year = _infer_journal_year(entry, section)
        entry.workload_year = year
        if target_year is not None:
            eligible = year == target_year
        else:
            eligible = year is not None and year >= start_year
        if not eligible:
            if year is not None and year < start_year and entry.workload_status != "skipped_year":
                entry.workload_status = "skipped_year"
                entry.workload_message = f"Журнал {year} року пропущено поточним правилом обробки"
                db.add(entry)
                skipped_year += 1
            continue
        try:
            process_journal_workload_entry(db, entry, service_account_json=section_service_account_json)
            processed += 1
            handled += 1
        except JournalMissingWorkbookError:
            remove_journal_entries_from_project(db, [entry])
            handled += 1
            continue
        except JournalNoDataError as exc:
            db.query(JournalWorkloadEntry).filter(JournalWorkloadEntry.journal_monitor_entry_id == entry.id).delete(
                synchronize_session=False
            )
            entry.workload_status = "no_data"
            entry.workload_message = str(exc)[:500]
            entry.workload_processed_at = datetime.now(timezone.utc)
            entry.workload_hours = round(float(exc.workload_hours or 0.0), 2)
            entry.workload_source_names = entry.workload_source_names or []
            db.add(entry)
            failed += 1
            handled += 1
        except Exception as exc:
            entry.workload_status = "failed"
            entry.workload_message = str(exc)[:500]
            entry.workload_processed_at = datetime.now(timezone.utc)
            db.add(entry)
            failed += 1
            handled += 1
            continue

    db.flush()
    return {"processed": processed, "failed": failed, "skipped_year": skipped_year}


def process_journal_trainees_for_section(
    db: Session,
    section: JournalMonitorSection,
    *,
    limit: int | None = 1,
    target_year: int | None = None,
    retry_failed: bool = True,
    entry_ids: set[int] | None = None,
) -> dict[str, Any]:
    from app.core.crypto import cipher

    processed = 0
    no_data = 0
    failed = 0
    handled = 0
    section_service_account_json = cipher.decrypt(section.service_account_json_encrypted)

    def status_priority(entry: JournalMonitorEntry) -> int:
        if entry.trainees_status == "pending":
            return 0
        if retry_failed and entry.trainees_status == "failed":
            return 1
        return 9

    entries = sorted(
        [entry for entry in section.entries if entry_ids is None or entry.id in entry_ids],
        key=lambda item: (
            status_priority(item),
            _infer_journal_year(item, section) or 9999,
            (item.group_code or "~~~~").casefold(),
            item.journal_name.casefold(),
        ),
    )
    for entry in entries:
        if limit is not None and handled >= limit:
            break
        if _entry_is_folder_audit_only(entry):
            entry.trainees_status = "no_data"
            entry.trainees_message = "У теці Google Drive не знайдено файлу журналу"
            entry.trainees_processed_at = entry.trainees_processed_at or datetime.now(timezone.utc)
            db.add(entry)
            continue
        if not entry.group_code:
            continue
        active_trainee_count = _active_trainee_count_for_group(db, entry.branch_id, entry.group_code)
        if entry.trainees_status == "processed" and not _entry_needs_trainee_reimport(entry, active_trainee_count):
            if _journal_workbooks_modified_after(entry, entry.trainees_processed_at, section_service_account_json):
                _requeue_entry_after_drive_change(db, entry)
            else:
                continue
        if entry.trainees_status == "no_data":
            if _journal_workbooks_modified_after(entry, entry.trainees_processed_at, section_service_account_json):
                _requeue_entry_after_drive_change(db, entry)
            else:
                continue
        if entry.trainees_status == "failed" and not retry_failed:
            continue
        entry_year = _infer_journal_year(entry, section)
        if target_year is not None and entry_year is not None and entry_year != target_year:
            continue
        try:
            process_journal_trainees_entry(db, entry, service_account_json=section_service_account_json)
            processed += 1
            handled += 1
        except JournalMissingWorkbookError:
            remove_journal_entries_from_project(db, [entry])
            handled += 1
            continue
        except JournalNoDataError as exc:
            archived_missing = _archive_missing_group_trainees(db, entry.branch_id, entry.group_code, [])
            entry.trainees_status = "no_data"
            archive_suffix = f"; архівовано відсутніх слухачів: {archived_missing}" if archived_missing else ""
            entry.trainees_message = f"{str(exc)}{archive_suffix}"[:500]
            entry.trainees_processed_at = datetime.now(timezone.utc)
            entry.trainees_source_names = entry.trainees_source_names or []
            entry.has_trainees = False
            entry.trainee_count = 0
            db.add(entry)
            no_data += 1
            handled += 1
        except Exception as exc:
            entry.trainees_status = "failed"
            entry.trainees_message = str(exc)[:500]
            entry.trainees_processed_at = datetime.now(timezone.utc)
            db.add(entry)
            failed += 1
            handled += 1
    db.flush()
    return {"processed": processed, "no_data": no_data, "failed": failed}


def requeue_journal_workload_for_year(db: Session, section: JournalMonitorSection, year: int, *, force: bool = False) -> int:
    changed = 0
    for entry in section.entries:
        entry_year = _infer_journal_year(entry, section)
        entry.workload_year = entry_year
        if entry_year != year:
            continue
        if force or entry.workload_status in {"failed", "needs_regeneration", "skipped_year"}:
            entry.workload_status = "pending"
            entry.workload_message = "Поставлено в чергу повторної обробки"
            entry.workload_processed_at = None
            entry.workload_hours = 0.0
            entry.workload_source_names = None
            db.query(JournalWorkloadEntry).filter(JournalWorkloadEntry.journal_monitor_entry_id == entry.id).delete(
                synchronize_session=False
            )
            db.add(entry)
            changed += 1
    db.flush()
    return changed


def requeue_journal_trainees_for_year(db: Session, section: JournalMonitorSection, year: int, *, force: bool = False) -> int:
    changed = 0
    for entry in section.entries:
        entry_year = _infer_journal_year(entry, section)
        if entry_year is not None and entry_year != year:
            continue
        if not entry.group_code:
            continue
        if not force and entry.trainees_status == "processed" and entry.has_trainees:
            continue
        if force or entry.trainees_status in {"pending", "failed", "no_data", "processed"}:
            entry.trainees_status = "pending"
            entry.trainees_message = "Поставлено в чергу повторної обробки слухачів"
            entry.trainees_processed_at = None
            entry.trainees_source_names = None
            db.add(entry)
            changed += 1
    db.flush()
    return changed


def requeue_selected_journal_entries(
    db: Session,
    section: JournalMonitorSection,
    entry_ids: list[int],
) -> tuple[int, int]:
    requested_ids = set(_unique_ints(entry_ids))
    workload_changed = 0
    trainees_changed = 0
    for entry in section.entries:
        if entry.id is None or entry.id not in requested_ids:
            continue
        entry.workload_status = "pending"
        entry.workload_message = "Поставлено в чергу ручного опрацювання"
        entry.workload_processed_at = None
        entry.workload_hours = 0.0
        entry.workload_source_names = None
        db.query(JournalWorkloadEntry).filter(JournalWorkloadEntry.journal_monitor_entry_id == entry.id).delete(
            synchronize_session=False
        )
        db.add(entry)
        workload_changed += 1
        if entry.group_code:
            entry.trainees_status = "pending"
            entry.trainees_message = "Поставлено в чергу ручного опрацювання"
            entry.trainees_processed_at = None
            entry.trainees_source_names = None
            db.add(entry)
            trainees_changed += 1
    db.flush()
    return workload_changed, trainees_changed


def lock_journal_monitor_section(
    db: Session,
    section: JournalMonitorSection | int,
    *,
    skip_locked: bool = False,
) -> JournalMonitorSection | None:
    section_id = section if isinstance(section, int) else section.id
    if section_id is None:
        return section if isinstance(section, JournalMonitorSection) else None
    query = (
        db.query(JournalMonitorSection)
        .filter(JournalMonitorSection.id == section_id)
    )
    return query.with_for_update(skip_locked=skip_locked).first()


def process_journal_monitor_section_step(
    db: Session,
    section: JournalMonitorSection,
    *,
    process_workload: bool = True,
    process_trainees: bool = True,
) -> dict[str, Any]:
    locked_section = lock_journal_monitor_section(db, section)
    if locked_section is not None:
        section = locked_section
    message_parts: list[str] = []
    result: dict[str, Any] = {
        "workload": {"processed": 0, "failed": 0, "skipped_year": 0},
        "trainees": {"processed": 0, "no_data": 0, "failed": 0},
    }
    if process_workload and section.workload_auto_enabled:
        workload_result = process_next_journal_workload(
            db,
            section,
            limit=None,
            target_year=section.workload_auto_year,
            retry_failed=True,
        )
        result["workload"] = workload_result
        if workload_result.get("processed") or workload_result.get("failed") or workload_result.get("skipped_year"):
            message_parts.append(
                f"педнавантаження: опрацьовано {workload_result['processed']}, "
                f"помилок {workload_result['failed']}"
            )
        db.flush()
        db.refresh(section)
    if process_trainees:
        trainees_result = process_journal_trainees_for_section(
            db,
            section,
            limit=1,
            target_year=section.workload_auto_year if section.workload_auto_enabled else None,
            retry_failed=True,
        )
        result["trainees"] = trainees_result
        if trainees_result.get("processed") or trainees_result.get("failed") or trainees_result.get("no_data"):
            message_parts.append(
                f"слухачі: опрацьовано {trainees_result['processed']}, "
                f"н/даних {trainees_result['no_data']}, помилок {trainees_result['failed']}"
            )
        db.flush()
        db.refresh(section)
        groups_by_code, schedule_counts, trainee_counts = _group_maps(db, section.branch_id)
        for entry in section.entries:
            _refresh_entry_project_state(db, entry, groups_by_code, schedule_counts, trainee_counts)
    if message_parts:
        message = _clip_monitor_message("; ".join(message_parts))
        section.last_processing_message = message
        section.last_sync_message = message
    db.flush()
    db.refresh(section)
    return result


def collect_monitor_stats(entries: list[JournalMonitorEntry]) -> dict[str, int]:
    stats = {
        "total": len(entries),
        "complete": 0,
        "schedule_only": 0,
        "trainees_only": 0,
        "not_processed": 0,
        "unknown_code": 0,
        "workload_only": 0,
        "workload_and_trainees": 0,
        "workload_trainees_schedule": 0,
    }
    for entry in entries:
        has_workload = entry.workload_status == "processed"
        has_journal_trainees = _entry_journal_trainee_count(entry) > 0
        effective_status = _status(entry.has_schedule, has_journal_trainees, entry.group_code)
        if has_workload and has_journal_trainees and entry.has_schedule:
            stats["workload_trainees_schedule"] += 1
            continue
        if has_workload and has_journal_trainees:
            stats["workload_and_trainees"] += 1
            continue
        if has_workload and not entry.has_schedule and not has_journal_trainees:
            stats["workload_only"] += 1
            continue
        if not has_workload and effective_status in stats:
            stats[effective_status] += 1
    return stats


def _daily_activity_entry_payload(entry: JournalMonitorEntry) -> dict[str, Any]:
    return {
        "id": entry.id,
        "drive_file_id": entry.drive_file_id,
        "drive_folder_id": entry.drive_folder_id,
        "drive_url": entry.drive_url,
        "journal_name": entry.journal_name,
        "group_code": entry.group_code,
        "created_at": _as_aware_utc(entry.drive_created_at),
        "change_started_at": _as_aware_utc(entry.drive_change_started_at),
        "modified_at": _as_aware_utc(entry.drive_modified_at),
    }


def collect_daily_journal_activity(section: JournalMonitorSection, now: datetime | None = None) -> dict[str, Any]:
    cutoff_at = _journal_daily_cutoff(now)
    entries = list(section.entries)
    created = [
        _daily_activity_entry_payload(entry)
        for entry in entries
        if (created_at := _as_aware_utc(entry.drive_created_at)) is not None and created_at >= cutoff_at
    ]
    changed = [
        _daily_activity_entry_payload(entry)
        for entry in entries
        if (change_started_at := _as_aware_utc(entry.drive_change_started_at)) is not None and change_started_at >= cutoff_at
    ]
    created.sort(key=lambda item: (item["created_at"] or datetime.max.replace(tzinfo=timezone.utc), str(item["journal_name"]).casefold()))
    changed.sort(
        key=lambda item: (
            item["change_started_at"] or datetime.max.replace(tzinfo=timezone.utc),
            str(item["journal_name"]).casefold(),
        )
    )
    return {
        "cutoff_at": cutoff_at,
        "created_count": len(created),
        "changed_count": len(changed),
        "created": created,
        "changed": changed,
    }


def section_to_response_payload(section: JournalMonitorSection, include_entries: bool = False) -> dict[str, Any]:
    entries = sorted(section.entries, key=lambda item: ((item.group_code or "~~~~").casefold(), item.journal_name.casefold()))
    payload = {
        "id": section.id,
        "name": section.name,
        "folder_url": section.folder_url,
        "folder_id": section.folder_id,
        "is_active": section.is_active,
        "workload_auto_enabled": section.workload_auto_enabled,
        "workload_auto_year": section.workload_auto_year,
        "has_service_account_credentials": bool(section.service_account_json_encrypted),
        "last_synced_at": section.last_synced_at,
        "last_sync_status": section.last_sync_status,
        "last_sync_message": section.last_sync_message,
        "last_processing_message": section.last_processing_message,
        "priority_queue_size": len(_section_priority_entry_ids(section)),
        "stats": collect_monitor_stats(entries),
        "daily_activity": collect_daily_journal_activity(section),
    }
    if include_entries:
        payload["entries"] = [entry_to_response_payload(entry) for entry in entries]
    return payload


def _entry_workload_teachers(entry: JournalMonitorEntry) -> list[dict[str, Any]]:
    if entry.workload_status != "processed":
        return []
    totals: dict[int, dict[str, Any]] = {}
    for workload_entry in entry.workload_entries:
        teacher = workload_entry.teacher
        teacher_id = int(workload_entry.teacher_id)
        if teacher_id not in totals:
            totals[teacher_id] = {
                "teacher_id": teacher_id,
                "teacher_name": _short_teacher_display_name(teacher),
                "hours": 0.0,
            }
        totals[teacher_id]["hours"] = round(float(totals[teacher_id]["hours"]) + float(workload_entry.hours or 0.0), 2)
    return sorted(totals.values(), key=lambda item: str(item["teacher_name"]).casefold())


def entry_to_response_payload(entry: JournalMonitorEntry) -> dict[str, Any]:
    journal_trainee_count = _entry_journal_trainee_count(entry)
    journal_has_trainees = journal_trainee_count > 0
    return {
        "id": entry.id,
        "drive_file_id": entry.drive_file_id,
        "drive_folder_id": entry.drive_folder_id,
        "drive_url": entry.drive_url,
        "journal_name": entry.journal_name,
        "group_code": entry.group_code,
        "matched_group_id": entry.matched_group_id,
        "has_group": entry.has_group,
        "has_schedule": entry.has_schedule,
        "has_trainees": journal_has_trainees,
        "schedule_lessons": entry.schedule_lessons,
        "trainee_count": journal_trainee_count,
        "processing_status": _status(entry.has_schedule, journal_has_trainees, entry.group_code),
        "workload_status": entry.workload_status,
        "workload_message": entry.workload_message,
        "workload_processed_at": entry.workload_processed_at,
        "workload_year": entry.workload_year,
        "workload_hours": entry.workload_hours,
        "workload_teachers": _entry_workload_teachers(entry),
        "workload_source_names": _visible_source_names(entry, entry.workload_source_names),
        "trainees_status": entry.trainees_status,
        "trainees_message": entry.trainees_message,
        "trainees_processed_at": entry.trainees_processed_at,
        "trainees_source_names": _visible_source_names(entry, entry.trainees_source_names),
        "drive_created_at": entry.drive_created_at,
        "drive_modified_at": entry.drive_modified_at,
        "drive_change_started_at": entry.drive_change_started_at,
        "last_seen_at": entry.last_seen_at,
    }


def _status(has_schedule: bool, has_trainees: bool, group_code: str | None) -> str:
    if not group_code:
        return "unknown_code"
    if has_schedule and has_trainees:
        return "complete"
    if has_schedule:
        return "schedule_only"
    if has_trainees:
        return "trainees_only"
    return "not_processed"


def _group_maps(db: Session, branch_id: str) -> tuple[dict[str, Group], dict[str, int], dict[str, int]]:
    groups = db.query(Group).filter(Group.branch_id == branch_id).all()
    groups_by_code = {normalize_group_code(group.code): group for group in groups}

    schedule_counts: dict[str, int] = {}
    schedule_rows = (
        db.query(Group.code, func.count(ScheduleSlot.id))
        .join(ScheduleSlot, ScheduleSlot.group_id == Group.id)
        .filter(Group.branch_id == branch_id)
        .group_by(Group.code)
        .all()
    )
    for code, count in schedule_rows:
        schedule_counts[normalize_group_code(code)] = int(count or 0)

    trainee_counts: dict[str, int] = {}
    trainee_rows = (
        db.query(Trainee.group_code, func.count(Trainee.id))
        .filter(
            Trainee.branch_id == branch_id,
            Trainee.is_deleted.is_(False),
            Trainee.group_code.is_not(None),
            Trainee.group_code != "",
        )
        .group_by(Trainee.group_code)
        .all()
    )
    for code, count in trainee_rows:
        trainee_counts[normalize_group_code(code)] = trainee_counts.get(normalize_group_code(code), 0) + int(count or 0)

    return groups_by_code, schedule_counts, trainee_counts


def _refresh_entry_project_state(
    db: Session,
    entry: JournalMonitorEntry,
    groups_by_code: dict[str, Group] | None = None,
    schedule_counts: dict[str, int] | None = None,
    trainee_counts: dict[str, int] | None = None,
) -> None:
    if groups_by_code is None or schedule_counts is None or trainee_counts is None:
        groups_by_code, schedule_counts, trainee_counts = _group_maps(db, entry.branch_id)
    normalized_code = normalize_group_code(entry.group_code)
    matched_group = groups_by_code.get(normalized_code) if normalized_code else None
    schedule_lessons = schedule_counts.get(normalized_code, 0)
    trainee_count = trainee_counts.get(normalized_code, 0)
    entry.matched_group_id = matched_group.id if matched_group else None
    entry.has_group = matched_group is not None
    entry.has_schedule = schedule_lessons > 0
    journal_trainee_count = _entry_journal_trainee_count(entry, trainee_count)
    entry.has_trainees = journal_trainee_count > 0
    entry.schedule_lessons = schedule_lessons
    entry.trainee_count = journal_trainee_count
    entry.processing_status = _status(entry.has_schedule, entry.has_trainees, entry.group_code)


def _journal_event_actor_name(actor_name: str | None, actor_source: str) -> str:
    name = (actor_name or "").strip()
    if name:
        return name[:255]
    return "Автооновлення" if actor_source == "auto" else "Система"


def _latest_journal_drive_events(db: Session, section_id: int) -> dict[tuple[str, str], JournalMonitorEvent]:
    events = (
        db.query(JournalMonitorEvent)
        .filter(JournalMonitorEvent.section_id == section_id)
        .order_by(JournalMonitorEvent.detected_at.asc(), JournalMonitorEvent.id.asc())
        .all()
    )
    return {(event.object_type, event.drive_file_id): event for event in events}


def _journal_event_changed(
    latest: JournalMonitorEvent,
    *,
    journal_name: str,
    drive_url: str | None,
    drive_mime_type: str | None,
    group_code: str | None,
    drive_modified_at: datetime | None,
) -> bool:
    latest_modified_at = _as_aware_utc(latest.drive_modified_at)
    next_modified_at = _as_aware_utc(drive_modified_at)
    if latest.action == "deleted":
        return True
    if latest_modified_at is not None and next_modified_at is not None and next_modified_at > latest_modified_at:
        return True
    return (
        latest.journal_name != journal_name
        or (latest.drive_url or None) != (drive_url or None)
        or (latest.drive_mime_type or None) != (drive_mime_type or None)
        or display_group_code(latest.group_code) != display_group_code(group_code)
    )


def _record_journal_drive_seen_event(
    db: Session,
    *,
    section: JournalMonitorSection,
    event_state: dict[tuple[str, str], JournalMonitorEvent],
    seen_event_keys: set[tuple[str, str]],
    object_type: str,
    drive_file_id: str,
    drive_folder_id: str | None,
    journal_name: str,
    group_code: str | None,
    drive_url: str | None,
    drive_mime_type: str | None,
    drive_created_at: datetime | None,
    drive_modified_at: datetime | None,
    detected_at: datetime,
    actor_user_id: int | None,
    actor_name: str | None,
    actor_source: str,
) -> None:
    key = (object_type, drive_file_id)
    seen_event_keys.add(key)
    latest = event_state.get(key)
    next_created_at = _as_aware_utc(drive_created_at)
    next_modified_at = _as_aware_utc(drive_modified_at)
    if latest is None or latest.action == "deleted":
        action = "created"
        occurred_at = next_created_at or next_modified_at or detected_at
    elif _journal_event_changed(
        latest,
        journal_name=journal_name,
        drive_url=drive_url,
        drive_mime_type=drive_mime_type,
        group_code=group_code,
        drive_modified_at=next_modified_at,
    ):
        action = "changed"
        occurred_at = next_modified_at or detected_at
    else:
        return
    event = JournalMonitorEvent(
        section_id=section.id,
        branch_id=section.branch_id,
        object_type=object_type,
        action=action,
        drive_file_id=drive_file_id,
        drive_folder_id=drive_folder_id,
        drive_mime_type=drive_mime_type,
        drive_url=drive_url,
        journal_name=journal_name,
        group_code=group_code,
        actor_user_id=actor_user_id,
        actor_name=_journal_event_actor_name(actor_name, actor_source),
        source=actor_source,
        drive_created_at=next_created_at,
        drive_modified_at=next_modified_at,
        occurred_at=occurred_at,
        detected_at=detected_at,
    )
    db.add(event)
    event_state[key] = event


def _record_deleted_journal_drive_events(
    db: Session,
    *,
    section: JournalMonitorSection,
    event_state: dict[tuple[str, str], JournalMonitorEvent],
    seen_event_keys: set[tuple[str, str]],
    detected_at: datetime,
    actor_user_id: int | None,
    actor_name: str | None,
    actor_source: str,
) -> None:
    for key, latest in list(event_state.items()):
        if latest.section_id != section.id or latest.action == "deleted" or key in seen_event_keys:
            continue
        event = JournalMonitorEvent(
            section_id=section.id,
            branch_id=section.branch_id,
            object_type=latest.object_type,
            action="deleted",
            drive_file_id=latest.drive_file_id,
            drive_folder_id=latest.drive_folder_id,
            drive_mime_type=latest.drive_mime_type,
            drive_url=latest.drive_url,
            journal_name=latest.journal_name,
            group_code=latest.group_code,
            actor_user_id=actor_user_id,
            actor_name=_journal_event_actor_name(actor_name, actor_source),
            source=actor_source,
            drive_created_at=_as_aware_utc(latest.drive_created_at),
            drive_modified_at=_as_aware_utc(latest.drive_modified_at),
            occurred_at=detected_at,
            detected_at=detected_at,
        )
        db.add(event)
        event_state[key] = event


def sync_journal_monitor_section(
    db: Session,
    section: JournalMonitorSection,
    folder_lister=None,
    workbook_lister=None,
    process_workload: bool = True,
    process_trainees: bool = True,
    actor_user_id: int | None = None,
    actor_name: str | None = None,
    actor_source: str = "auto",
) -> JournalMonitorSection:
    now = datetime.now(timezone.utc)
    from app.core.crypto import cipher

    if folder_lister is None:
        folder_lister = list_drive_child_folders
    if workbook_lister is None:
        workbook_lister = list_drive_journal_workbook_files
    locked_section = lock_journal_monitor_section(db, section)
    if locked_section is not None:
        section = locked_section
    section_service_account_json = cipher.decrypt(section.service_account_json_encrypted)
    folders = folder_lister(section.folder_id, service_account_json=section_service_account_json)
    groups_by_code, schedule_counts, trainee_counts = _group_maps(db, section.branch_id)
    seen_drive_ids: set[str] = set()
    daily_cutoff_at = _journal_daily_cutoff(now)
    event_state = _latest_journal_drive_events(db, section.id)
    seen_event_keys: set[tuple[str, str]] = set()
    section_entries = (
        db.query(JournalMonitorEntry)
        .filter(JournalMonitorEntry.section_id == section.id)
        .all()
    )

    entries_by_drive_id = {entry.drive_file_id: entry for entry in section_entries}
    for folder in folders:
        drive_id = str(folder.get("id") or "").strip()
        name = str(folder.get("name") or "").strip() or drive_id
        if not drive_id:
            continue
        folder_group_code = extract_group_code(name)
        folder_created_at = _as_aware_utc(_parse_datetime(str(folder.get("created_time") or "")))
        folder_modified_at = _as_aware_utc(_parse_datetime(str(folder.get("modified_time") or "")))
        _record_journal_drive_seen_event(
            db,
            section=section,
            event_state=event_state,
            seen_event_keys=seen_event_keys,
            object_type="folder",
            drive_file_id=drive_id,
            drive_folder_id=None,
            journal_name=name,
            group_code=folder_group_code,
            drive_url=str(folder.get("url") or f"https://drive.google.com/drive/folders/{drive_id}"),
            drive_mime_type=GOOGLE_DRIVE_FOLDER_MIME,
            drive_created_at=folder_created_at,
            drive_modified_at=folder_modified_at,
            detected_at=now,
            actor_user_id=actor_user_id,
            actor_name=actor_name,
            actor_source=actor_source,
        )
        try:
            workbook_files = workbook_lister(drive_id, service_account_json=section_service_account_json)
        except Exception:
            for existing_entry in section_entries:
                if existing_entry.drive_folder_id == drive_id or (
                    existing_entry.drive_folder_id is None and existing_entry.drive_file_id == drive_id
                ):
                    seen_drive_ids.add(existing_entry.drive_file_id)
            continue
        if not workbook_files:
            entry = entries_by_drive_id.get(drive_id)
            has_imported_journal_data = bool(
                entry
                and entry.drive_folder_id is None
                and (
                    entry.workload_status == "processed"
                    or entry.trainees_status == "processed"
                    or entry.workload_hours > 0
                    or entry.trainee_count > 0
                )
            )
            if has_imported_journal_data:
                continue
            seen_drive_ids.add(drive_id)
            if not entry:
                entry = JournalMonitorEntry(
                    section_id=section.id,
                    branch_id=section.branch_id,
                    drive_file_id=drive_id,
                    journal_name=name,
                )
                db.add(entry)
                entries_by_drive_id[drive_id] = entry
                section_entries.append(entry)
            entry.drive_file_id = drive_id
            entry.drive_folder_id = drive_id
            entry.drive_mime_type = GOOGLE_DRIVE_FOLDER_MIME
            entry.drive_url = str(folder.get("url") or f"https://drive.google.com/drive/folders/{drive_id}")
            entry.journal_name = name
            entry.group_code = folder_group_code
            entry.workload_status = "no_data"
            entry.workload_message = "У теці Google Drive не знайдено файлу журналу"
            entry.workload_processed_at = entry.workload_processed_at or now
            entry.workload_hours = 0.0
            entry.workload_source_names = []
            entry.trainees_status = "no_data"
            entry.trainees_message = "У теці Google Drive не знайдено файлу журналу"
            entry.trainees_processed_at = entry.trainees_processed_at or now
            entry.trainees_source_names = []
            entry.drive_created_at = folder_created_at
            entry.drive_modified_at = folder_modified_at
            _refresh_entry_project_state(db, entry, groups_by_code, schedule_counts, trainee_counts)
            entry.last_seen_at = now
            continue
        legacy_folder_entry = entries_by_drive_id.get(drive_id)
        for workbook_file in workbook_files:
            workbook_id = str(workbook_file.get("id") or "").strip()
            if not workbook_id:
                continue
            seen_drive_ids.add(workbook_id)
            workbook_name = _journal_name_from_workbook(workbook_file, name, drive_id)
            group_code = extract_group_code(workbook_name) or folder_group_code
            next_created_at = _as_aware_utc(_parse_datetime(str(workbook_file.get("createdTime") or ""))) or _as_aware_utc(
                _parse_datetime(folder.get("created_time"))
            )
            next_modified_at = _drive_activity_modified_at(folder, [workbook_file])
            _record_journal_drive_seen_event(
                db,
                section=section,
                event_state=event_state,
                seen_event_keys=seen_event_keys,
                object_type="workbook",
                drive_file_id=workbook_id,
                drive_folder_id=drive_id,
                journal_name=workbook_name,
                group_code=group_code,
                drive_url=str(workbook_file.get("webViewLink") or f"https://drive.google.com/file/d/{workbook_id}/view"),
                drive_mime_type=str(workbook_file.get("mimeType") or ""),
                drive_created_at=next_created_at,
                drive_modified_at=next_modified_at,
                detected_at=now,
                actor_user_id=actor_user_id,
                actor_name=actor_name,
                actor_source=actor_source,
            )
            entry = entries_by_drive_id.get(workbook_id)
            created_entry = False
            if not entry and legacy_folder_entry and len(workbook_files) == 1:
                entry = legacy_folder_entry
                entries_by_drive_id.pop(drive_id, None)
                entries_by_drive_id[workbook_id] = entry
            if not entry:
                entry = JournalMonitorEntry(
                    section_id=section.id,
                    branch_id=section.branch_id,
                    drive_file_id=workbook_id,
                    journal_name=workbook_name,
                )
                db.add(entry)
                entries_by_drive_id[workbook_id] = entry
                section_entries.append(entry)
                created_entry = True
            entry.drive_file_id = workbook_id
            old_group_code = display_group_code(entry.group_code)
            next_group_code = display_group_code(group_code)
            renamed_group = old_group_code and next_group_code and normalize_group_code(old_group_code) != normalize_group_code(next_group_code)
            workbook_changed_after_workload = _workbook_files_modified_after([workbook_file], entry.workload_processed_at)
            workbook_changed_after_trainees = _workbook_files_modified_after([workbook_file], entry.trainees_processed_at)

            if renamed_group:
                hide_groups_for_deleted_journal_entries(db, [entry])
                archive_trainees_for_deleted_journal_entries(db, [entry])

            entry.drive_folder_id = drive_id
            entry.drive_mime_type = str(workbook_file.get("mimeType") or "")
            entry.drive_url = str(workbook_file.get("webViewLink") or f"https://drive.google.com/file/d/{workbook_id}/view")
            entry.journal_name = workbook_name
            entry.group_code = group_code
            _refresh_entry_project_state(db, entry, groups_by_code, schedule_counts, trainee_counts)
            if next_created_at is not None:
                entry.drive_created_at = next_created_at
            entry.drive_modified_at = next_modified_at
            _update_daily_drive_change_start(
                entry,
                cutoff_at=daily_cutoff_at,
                drive_created_at=next_created_at,
                folder=folder,
                workbook_files=[workbook_file],
            )
            entry.last_seen_at = now
            if created_entry:
                entry.workload_message = "Новий журнал у Google Drive. Очікує опрацювання"
                if entry.group_code:
                    entry.trainees_message = "Новий журнал у Google Drive. Очікує опрацювання"
            if renamed_group or workbook_changed_after_workload or workbook_changed_after_trainees:
                _requeue_entry_after_drive_change(
                    db,
                    entry,
                    requeue_workload=workbook_changed_after_workload,
                    requeue_trainees=renamed_group or workbook_changed_after_trainees,
                )

    db.flush()
    db.expire(section, ["entries"])
    db.refresh(section)
    created_groups = ensure_groups_for_journal_entries(db, section)
    if created_groups:
        groups_by_code, schedule_counts, trainee_counts = _group_maps(db, section.branch_id)
        for entry in section.entries:
            _refresh_entry_project_state(db, entry, groups_by_code, schedule_counts, trainee_counts)

    current_entries = (
        db.query(JournalMonitorEntry)
        .filter(JournalMonitorEntry.section_id == section.id)
        .all()
    )
    _record_deleted_journal_drive_events(
        db,
        section=section,
        event_state=event_state,
        seen_event_keys=seen_event_keys,
        detected_at=now,
        actor_user_id=actor_user_id,
        actor_name=actor_name,
        actor_source=actor_source,
    )
    removed_entries = [entry for entry in current_entries if entry.drive_file_id not in seen_drive_ids]
    if removed_entries:
        remove_journal_entries_from_project(db, removed_entries)

    section.last_synced_at = now
    section.last_sync_status = "success"
    section.last_sync_message = f"Оновлено журналів: {len(seen_drive_ids)}"
    db.flush()
    db.refresh(section)
    process_journal_monitor_section_step(
        db,
        section,
        process_workload=process_workload,
        process_trainees=process_trainees,
    )
    db.flush()
    db.refresh(section)
    return section


def _workload_background_priority(
    entry: JournalMonitorEntry,
    section: JournalMonitorSection,
    target_year: int | None,
    *,
    retry_failed: bool = True,
    manual_selected: bool = False,
) -> int | None:
    year = _infer_journal_year(entry, section)
    if target_year is not None:
        eligible = year == target_year
    else:
        eligible = year is not None and year >= JOURNAL_WORKLOAD_START_YEAR
    if not eligible:
        return None
    if entry.workload_status in {"pending", "needs_regeneration"}:
        return 0
    if retry_failed and entry.workload_status == "failed":
        return 1
    if manual_selected and entry.workload_status == "no_data":
        return 2
    return None


def _trainees_background_priority(
    db: Session,
    entry: JournalMonitorEntry,
    section: JournalMonitorSection,
    target_year: int | None,
    *,
    retry_failed: bool = True,
    manual_selected: bool = False,
) -> int | None:
    if not entry.group_code:
        return None
    entry_year = _infer_journal_year(entry, section)
    if target_year is not None and entry_year is not None and entry_year != target_year:
        return None
    if entry.trainees_status == "pending":
        return 0
    if retry_failed and entry.trainees_status == "failed":
        return 1
    if entry.trainees_status == "processed":
        active_count = _active_trainee_count_for_group(db, entry.branch_id, entry.group_code)
        if _entry_needs_trainee_reimport(entry, active_count):
            return 1
    return None


def _next_background_journal_entry(
    db: Session,
    section: JournalMonitorSection,
    target_year: int | None,
    *,
    entry_ids: set[int] | None = None,
    retry_failed: bool = False,
) -> JournalMonitorEntry | None:
    candidates: list[tuple[int, int, str, str, JournalMonitorEntry]] = []
    for entry in section.entries:
        if entry_ids is not None and entry.id not in entry_ids:
            continue
        manual_selected = bool(entry_ids is not None and entry.id in entry_ids)
        priorities = [
            priority
            for priority in (
                _workload_background_priority(
                    entry,
                    section,
                    target_year,
                    retry_failed=retry_failed,
                    manual_selected=manual_selected,
                ),
                _trainees_background_priority(
                    db,
                    entry,
                    section,
                    target_year,
                    retry_failed=retry_failed,
                    manual_selected=manual_selected,
                ),
            )
            if priority is not None
        ]
        if not priorities:
            continue
        candidates.append(
            (
                min(priorities),
                _infer_journal_year(entry, section) or 9999,
                (entry.group_code or "~~~~").casefold(),
                entry.journal_name.casefold(),
                entry,
            )
        )
    return min(candidates, default=None, key=lambda item: item[:4])[4] if candidates else None


def process_journal_monitor_background_step(
    db: Session,
    section: JournalMonitorSection,
    *,
    folder_lister=list_drive_child_folders,
    target_year: int | None = None,
    sync_before: bool = True,
    workload_limit: int | None = 1,
    trainees_limit: int | None = 1,
    actor_user_id: int | None = None,
    actor_name: str | None = None,
    actor_source: str = "auto",
) -> dict[str, Any]:
    locked_section = lock_journal_monitor_section(db, section, skip_locked=True)
    if locked_section is None and getattr(section, "id", None) is not None:
        return {
            "workload": {"processed": 0, "failed": 0, "skipped_year": 0},
            "trainees": {"processed": 0, "no_data": 0, "failed": 0},
            "skipped": "locked",
        }
    if locked_section is not None:
        section = locked_section
    section_id = section.id
    sync_warning: str | None = None
    sync_deferred = False
    retry_failed = False
    priority_entry_ids = set(_section_priority_entry_ids(section))
    initial_target_year = (
        section.priority_queue_year
        if priority_entry_ids
        else (target_year if target_year is not None else section.workload_auto_year)
    )
    has_actionable_backlog = (
        _next_background_journal_entry(
            db,
            section,
            initial_target_year,
            entry_ids=priority_entry_ids or None,
            retry_failed=bool(priority_entry_ids),
        )
        is not None
    )
    if sync_before and not priority_entry_ids and not has_actionable_backlog:
        for attempt in range(2):
            try:
                section = sync_journal_monitor_section(
                    db,
                    section,
                    folder_lister=folder_lister,
                    process_workload=False,
                    process_trainees=False,
                    actor_user_id=actor_user_id,
                    actor_name=actor_name,
                    actor_source=actor_source,
                )
                retry_failed = True
                break
            except Exception as exc:
                db.rollback()
                section = db.get(JournalMonitorSection, section_id)
                if section is None:
                    raise
                if attempt == 0 and _is_transient_postgres_lock_error(exc):
                    continue
                if _is_transient_postgres_lock_error(exc):
                    _mark_sync_temporarily_busy(section)
                    sync_deferred = True
                    db.add(section)
                    db.flush()
                    break
                sync_warning = f"Синхронізацію Drive пропущено: {exc}"
                section.last_sync_status = "failed"
                section.last_sync_message = _clip_monitor_message(sync_warning)
                db.add(section)
                db.flush()
                break
    if priority_entry_ids:
        valid_entry_ids = {entry.id for entry in section.entries if entry.id is not None}
        filtered_entry_ids = sorted(entry_id for entry_id in priority_entry_ids if entry_id in valid_entry_ids)
        if filtered_entry_ids != _section_priority_entry_ids(section):
            _set_section_priority_queue(section, filtered_entry_ids, queue_year=section.priority_queue_year)
            db.add(section)
            db.flush()
        priority_entry_ids = set(filtered_entry_ids)
    effective_target_year = (
        section.priority_queue_year
        if priority_entry_ids
        else (target_year if target_year is not None else section.workload_auto_year)
    )
    entry_ids: set[int] | None = priority_entry_ids or None
    if workload_limit == 1 and trainees_limit == 1:
        target_entry = _next_background_journal_entry(
            db,
            section,
            effective_target_year,
            entry_ids=entry_ids,
            retry_failed=retry_failed or bool(priority_entry_ids),
        )
        entry_ids = {target_entry.id} if target_entry and target_entry.id is not None else set()
    workload_result = process_next_journal_workload(
        db,
        section,
        limit=workload_limit,
        target_year=effective_target_year,
        retry_failed=retry_failed or bool(priority_entry_ids),
        entry_ids=entry_ids,
    )
    db.flush()
    db.refresh(section)
    trainees_result = process_journal_trainees_for_section(
        db,
        section,
        limit=trainees_limit,
        target_year=effective_target_year,
        retry_failed=retry_failed or bool(priority_entry_ids),
        entry_ids=entry_ids,
    )
    db.flush()
    db.refresh(section)
    groups_by_code, schedule_counts, trainee_counts = _group_maps(db, section.branch_id)
    for entry in section.entries:
        _refresh_entry_project_state(db, entry, groups_by_code, schedule_counts, trainee_counts)
    processing_message = (
        "Фонове опрацювання: "
        f"педнавантаження {workload_result['processed']}/{workload_result['failed']}, "
        f"слухачі {trainees_result['processed']}/{trainees_result['no_data']}/{trainees_result['failed']}"
    )
    if not sync_deferred:
        section.last_sync_message = _clip_monitor_message(
            f"{sync_warning}; {processing_message}" if sync_warning else processing_message
        )
    if priority_entry_ids:
        entry_map = {entry.id: entry for entry in section.entries if entry.id is not None}
        remaining_ids = [
            entry_id
            for entry_id in _section_priority_entry_ids(section)
            if (entry := entry_map.get(entry_id)) is not None
            and (
                entry.workload_status in {"pending", "needs_regeneration"}
                or (entry.group_code and entry.trainees_status == "pending")
            )
        ]
        if remaining_ids:
            _set_section_priority_queue(
                section,
                remaining_ids,
                queue_year=section.priority_queue_year,
                message=f"Пріоритетна черга: залишилось {len(remaining_ids)} журналів",
            )
        else:
            _clear_section_priority_queue(
                section,
                message="Пріоритетну чергу завершено. Повернулись до відстеження змін у журналах.",
            )
    else:
        section.last_processing_message = _clip_monitor_message(processing_message)
    if sync_warning:
        section.last_sync_message = _clip_monitor_message(sync_warning)
    db.add(section)
    db.flush()
    db.refresh(section)
    return {"workload": workload_result, "trainees": trainees_result}


def collect_export_rows(
    section: JournalMonitorSection,
    entries: list[JournalMonitorEntry] | None = None,
) -> list[dict[str, Any]]:
    source_entries = list(section.entries) if entries is None else entries
    sorted_entries = sorted(source_entries, key=lambda item: ((item.group_code or "~~~~").casefold(), item.journal_name.casefold()))
    return [
        {
            "Розділ": section.name,
            "Номер групи": entry.group_code or "",
            "Назва папки журналу": entry.journal_name,
            "Статус опрацювання": format_processing_status(entry.processing_status),
            "Статус педнавантаження": format_workload_status(entry.workload_status),
            "Годин із журналу": entry.workload_hours,
            "Файли журналів": "; ".join(_visible_source_names(entry, entry.workload_source_names)),
            "Статус слухачів": format_trainees_status(entry.trainees_status),
            "Файли ЗВ": "; ".join(_visible_source_names(entry, entry.trainees_source_names)),
            "Рік педнавантаження": entry.workload_year or "",
            "Повідомлення педнавантаження": entry.workload_message or "",
            "Повідомлення слухачів": entry.trainees_message or "",
            "Є група в системі": "Так" if entry.has_group else "Ні",
            "Є розклад": "Так" if entry.has_schedule else "Ні",
            "Занять у розкладі": entry.schedule_lessons,
            "Є слухачі": "Так" if entry.has_trainees else "Ні",
            "Кількість слухачів": entry.trainee_count,
            "Посилання Drive": entry.drive_url or "",
            "Остання синхронізація": section.last_synced_at.isoformat() if section.last_synced_at else "",
        }
        for entry in sorted_entries
    ]


def format_processing_status(value: str) -> str:
    return {
        "complete": "Опрацьовано: розклад і слухачі",
        "schedule_only": "Опрацьовано тільки розклад",
        "trainees_only": "Опрацьовано тільки слухачі",
        "not_processed": "Не опрацьовано",
        "unknown_code": "Не визначено номер групи",
    }.get(value, value)


def format_workload_status(value: str) -> str:
    return {
        "pending": "Очікує обробки",
        "processed": "Педнавантаження додано",
        "failed": "Помилка обробки",
        "skipped_year": "Пропущено за роком",
        "needs_regeneration": "Потребує повторної обробки",
        "no_data": "Н/даних",
    }.get(value, value)


def format_trainees_status(value: str) -> str:
    return {
        "pending": "Очікує обробки",
        "processed": "Додано",
        "failed": "Помилка обробки",
        "no_data": "Н/даних",
    }.get(value, value)


def filter_journal_monitor_entries(
    entries: list[JournalMonitorEntry],
    *,
    query: str | None = None,
    status: str | None = None,
    workload: str | None = None,
    has_schedule: bool | None = None,
    has_trainees: bool | None = None,
) -> list[JournalMonitorEntry]:
    normalized_query = (query or "").strip().casefold()
    filtered = entries
    if normalized_query:
        filtered = [
            entry
            for entry in filtered
            if normalized_query in (entry.group_code or "").casefold()
            or normalized_query in (entry.journal_name or "").casefold()
        ]
    if status:
        filtered = [entry for entry in filtered if entry.processing_status == status]
    if workload == "workload_only":
        filtered = [
            entry
            for entry in filtered
            if entry.workload_status == "processed" and not entry.has_schedule and not entry.has_trainees
        ]
    elif workload == "with_workload":
        filtered = [entry for entry in filtered if entry.workload_status == "processed"]
    elif workload == "without_workload":
        filtered = [entry for entry in filtered if entry.workload_status != "processed"]
    if has_schedule is not None:
        filtered = [entry for entry in filtered if entry.has_schedule is has_schedule]
    if has_trainees is not None:
        filtered = [entry for entry in filtered if entry.has_trainees is has_trainees]
    return filtered


def save_journal_monitor_export(
    section: JournalMonitorSection,
    export_format: str,
    *,
    query: str | None = None,
    status: str | None = None,
    workload: str | None = None,
    has_schedule: bool | None = None,
    has_trainees: bool | None = None,
) -> tuple[str, str, str]:
    if export_format not in EXPORT_FORMATS:
        raise ValueError("Підтримуються формати xlsx, pdf, docx, csv")

    source_entries = filter_journal_monitor_entries(
        list(section.entries),
        query=query,
        status=status,
        workload=workload,
        has_schedule=has_schedule,
        has_trainees=has_trainees,
    )
    rows = collect_export_rows(section, source_entries)
    safe_name = re.sub(r"[^0-9A-Za-zА-Яа-яІіЇїЄєҐґ_-]+", "_", section.name).strip("_") or "journals"
    timestamp = datetime.now(timezone.utc).strftime("%Y%m%d%H%M%S")

    if export_format == "pdf":
        path, _doc_type = save_report_file(rows, "journal_monitor", "pdf")
        return path, f"{safe_name}_{timestamp}.pdf", "application/pdf"

    temp_dir = Path(tempfile.gettempdir()) / "suptc_exports"
    temp_dir.mkdir(parents=True, exist_ok=True)
    out_file = temp_dir / f"{safe_name}_{timestamp}.{export_format}"
    headers = list(rows[0].keys()) if rows else ["Дані"]

    if export_format == "csv":
        with out_file.open("w", encoding="utf-8-sig", newline="") as handle:
            writer = csv.DictWriter(handle, fieldnames=headers)
            writer.writeheader()
            for row in rows:
                writer.writerow(row)
        return str(out_file), out_file.name, "text/csv; charset=utf-8"

    if export_format == "xlsx":
        workbook = Workbook()
        sheet = workbook.active
        sheet.title = "Журнали"
        sheet.append(headers)
        for row in rows:
            sheet.append([row.get(header) for header in headers])
        workbook.save(out_file)
        return str(out_file), out_file.name, "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"

    document = DocxDocument()
    document.add_heading(f"Моніторинг журналів: {section.name}", level=1)
    table = document.add_table(rows=1, cols=len(headers))
    header_cells = table.rows[0].cells
    for index, header in enumerate(headers):
        header_cells[index].text = header
    for row in rows:
        cells = table.add_row().cells
        for index, header in enumerate(headers):
            cells[index].text = str(row.get(header, ""))
    document.save(out_file)
    return str(out_file), out_file.name, "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
